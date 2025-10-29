import os, sys, json, re
import hashlib
import secrets
from datetime import datetime, timedelta, date
from typing import List
import fitz  # PyMuPDF (kept to preserve functionality if used in templates/utilities)
import pytesseract  # OCR (kept to preserve functionality if used elsewhere)
from PIL import Image  # Image handling (kept)
from flask import Flask, render_template, request, redirect, url_for, session, send_from_directory, send_file, flash, abort, jsonify, Response
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from flask_sqlalchemy import SQLAlchemy
from extensions import db
from sqlalchemy import func, or_, and_, text, inspect
from sqlalchemy.exc import OperationalError
from pywebpush import webpush, WebPushException
from docx import Document
from pdf_templates import create_pdf
from reportlab.graphics.barcode import qr as rl_qr
from reportlab.graphics.shapes import Drawing
from reportlab.graphics import renderPM
import requests
import time
from b2sdk.v2 import InMemoryAccountInfo, B2Api

# Optional override to load the 'consulting' package from a specific directory.
# If the provided path points to the 'consulting' directory itself, we add its parent
# to sys.path so that `import consulting...` resolves correctly.
CONSULTING_DIR_OVERRIDE = (
    os.environ.get("CONSULTING_DIR")
    or r"C:\\Users\\User\\Documents\\GitHub\\valeu\\erp-valuation\\consulting"
)
if CONSULTING_DIR_OVERRIDE:
    normalized_path = os.path.normpath(CONSULTING_DIR_OVERRIDE)
    try:
        if os.path.isdir(normalized_path):
            base_name = os.path.basename(normalized_path).lower()
            candidate_path = (
                os.path.dirname(normalized_path)
                if base_name == "consulting"
                else normalized_path
            )
            if candidate_path not in sys.path:
                sys.path.insert(0, candidate_path)
    except Exception:
        # Silently ignore invalid paths to avoid breaking app startup
        pass

from consulting.projects.models import ConsultingProject
from consulting.clients.models import Client
from consulting.projects.forms import PROJECT_TYPES

# ---------------- Ø¥Ø¹Ø¯Ø§Ø¯ Flask ----------------
app = Flask(__name__)
app.secret_key = "secret_key"

# ---------------- Ø¥Ø¹Ø¯Ø§Ø¯ Ø§Ù„Ù…Ù„ÙØ§Øª ----------------
UPLOAD_FOLDER = os.path.join(app.root_path, "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# ---------------- Backblaze B2 ----------------
# ÙŠÙØ¶Ù„ Ø¶Ø¨Ø· Ø¨ÙŠØ§Ù†Ø§Øª Ø§Ù„Ø¯Ø®ÙˆÙ„ Ø¹Ø¨Ø± Ù…ØªØºÙŠØ±Ø§Øª Ø§Ù„Ø¨ÙŠØ¦Ø©: B2_KEY_ID Ùˆ B2_APPLICATION_KEY
app.config["B2_KEY_ID"] = os.environ.get("B2_KEY_ID")
app.config["B2_APPLICATION_KEY"] = os.environ.get("B2_APPLICATION_KEY")
# ÙŠÙ…ÙƒÙ† ØªØ­Ø¯ÙŠØ¯ Ø§Ù„Ø¨ÙƒØª Ø¥Ù…Ø§ Ø¹Ø¨Ø± Ø§Ù„Ù…Ø¹Ø±Ù‘Ù Ø£Ùˆ Ø§Ù„Ø§Ø³Ù…. Ù„Ø§ Ù†Ø¶Ø¹ Ù‚ÙŠÙ…Ø© Ø§ÙØªØ±Ø§Ø¶ÙŠØ© ØµÙ„Ø¨Ø©.
app.config["B2_BUCKET_ID"] = os.environ.get("B2_BUCKET_ID")
app.config["B2_BUCKET_NAME"] = os.environ.get("B2_BUCKET_NAME") or os.environ.get("B2_BUCKET")

def get_b2_api() -> B2Api:
    key_id = app.config.get("B2_KEY_ID")
    app_key = app.config.get("B2_APPLICATION_KEY")
    if not key_id or not app_key:
        raise RuntimeError("B2 credentials (B2_KEY_ID/B2_APPLICATION_KEY) are not configured")
    info = InMemoryAccountInfo()
    api = B2Api(info)
    api.authorize_account("production", key_id, app_key)
    return api

def get_b2_bucket():
    api = get_b2_api()
    bucket_id = app.config.get("B2_BUCKET_ID")
    bucket_name = app.config.get("B2_BUCKET_NAME")

    # Ø£ÙˆÙ„ÙˆÙŠØ©: Ø¥Ø°Ø§ Ø¹Ø±Ù‘Ù Ø§Ù„Ù…Ø³ØªØ®Ø¯Ù… Ø§Ù„Ù…Ø¹Ø±Ù‘Ù Ù†Ø¨Ø­Ø« Ø¨Ù‡ØŒ ÙˆØ¥Ù„Ø§ Ù†Ø­Ø§ÙˆÙ„ Ø¨Ø§Ù„Ø§Ø³Ù…
    if bucket_id:
        try:
            # Ù…ØªÙˆÙØ± ÙÙŠ b2sdk v2
            return api.get_bucket_by_id(bucket_id)
        except Exception:
            # Ø§Ø­ØªÙŠØ§Ø·ÙŠÙ‹Ø§: Ø§Ø¨Ø­Ø« Ø¶Ù…Ù† Ø§Ù„Ù‚ÙˆØ§Ø¦Ù…
            for b in api.list_buckets():
                if getattr(b, "id_", None) == bucket_id:
                    return b
            raise RuntimeError("B2 bucket not found for configured B2_BUCKET_ID")

    if bucket_name:
        try:
            return api.get_bucket_by_name(bucket_name)
        except Exception:
            for b in api.list_buckets():
                if getattr(b, "name", None) == bucket_name:
                    return b
            raise RuntimeError("B2 bucket not found for configured B2_BUCKET_NAME/B2_BUCKET")

    raise RuntimeError("B2 bucket is not configured. Set B2_BUCKET_ID or B2_BUCKET_NAME/B2_BUCKET")

# ØªÙˆÙ„ÙŠØ¯ Ø±Ø§Ø¨Ø· ØªÙ†Ø²ÙŠÙ„ Ø¹Ø§Ù… Ù…Ø¨Ø§Ø´Ø± Ù„Ù…Ù„Ù Ø¯Ø§Ø®Ù„ Backblaze B2 (ÙŠØªØ·Ù„Ø¨ Ø£Ù† ÙŠÙƒÙˆÙ† Ø§Ù„Ø¨ÙƒØª Ø¹Ø§Ù…Ù‹Ø§)
def build_b2_public_url(file_name: str) -> str | None:
    try:
        if not file_name:
            return None
        api = get_b2_api()
        # Ø§Ù„Ø­ØµÙˆÙ„ Ø¹Ù„Ù‰ base download url Ù…Ù† Ù…Ø¹Ù„ÙˆÙ…Ø§Øª Ø§Ù„Ø­Ø³Ø§Ø¨ (ØªØ®ØªÙ„Ù Ø­Ø³Ø¨ Ù†Ø³Ø®Ø© b2sdk)
        download_base = None
        try:
            download_base = api.account_info.get_download_url()  # type: ignore[attr-defined]
        except Exception:
            try:
                download_base = api.session.account_info.get_download_url()  # type: ignore[attr-defined]
            except Exception:
                download_base = None

        bucket_name = app.config.get("B2_BUCKET_NAME")
        if not bucket_name:
            try:
                bucket = get_b2_bucket()
                bucket_name = getattr(bucket, "name", None)
            except Exception:
                bucket_name = None

        if not download_base or not bucket_name:
            return None

        from urllib.parse import quote
        # Ù…Ù„Ø§Ø­Ø¸Ø©: ÙŠÙØªØ±Ø¶ Ø£Ù† Ø§Ù„Ø¨ÙƒØª Ø¹Ø§Ù…. Ø¥Ù† ÙƒØ§Ù† Ø®Ø§ØµÙ‹Ø§ ÙØ³ÙŠØ­ØªØ§Ø¬ Ø±Ø§Ø¨Ø·Ù‹Ø§ Ù…ÙˆÙ‚Ù‘ØªÙ‹Ø§ (Ø®Ø§Ø±Ø¬ Ù†Ø·Ø§Ù‚ Ù‡Ø°Ø§ Ø§Ù„Ø·Ù„Ø¨)
        return f"{download_base}/file/{bucket_name}/{quote(file_name)}"
    except Exception:
        return None

# Ø§Ø¬Ø¹Ù„ Ø§Ù„Ø¯Ø§Ù„Ø© Ù…ØªØ§Ø­Ø© Ø¯Ø§Ø®Ù„ Ù‚ÙˆØ§Ù„Ø¨ Jinja Ù…Ø¨Ø§Ø´Ø±Ø©Ù‹
@app.context_processor
def inject_template_helpers():
    return {
        "build_b2_public_url": build_b2_public_url,
    }

# ---------------- Ø¥Ø¹Ø¯Ø§Ø¯ Ù…ÙØ§ØªÙŠØ­ Web Push (VAPID) ----------------
# ÙŠÙ…ÙƒÙ† Ø¶Ø¨Ø·Ù‡Ø§ Ø¹Ø¨Ø± Ù…ØªØºÙŠØ±Ø§Øª Ø§Ù„Ø¨ÙŠØ¦Ø© VAPID_PUBLIC_KEY / VAPID_PRIVATE_KEY / VAPID_SUBJECT
app.config["VAPID_PUBLIC_KEY"] = os.environ.get(
    "VAPID_PUBLIC_KEY",
    app.config.get("VAPID_PUBLIC_KEY") or "BFNeZpjEro8pwFxR1H20twlTd2pL5MZtWrDATu4ME2RcbzhN"
)
app.config["VAPID_PRIVATE_KEY"] = os.environ.get("VAPID_PRIVATE_KEY", app.config.get("VAPID_PRIVATE_KEY"))
app.config["VAPID_CLAIMS"] = {
    "sub": os.environ.get("VAPID_SUBJECT", "mailto:your-email@example.com")
}

# ---------------- Ø¥Ø¹Ø¯Ø§Ø¯ Ù‚Ø§Ø¹Ø¯Ø© Ø§Ù„Ø¨ÙŠØ§Ù†Ø§Øª ----------------
# ÙŠØ¯Ø¹Ù… DATABASE_URL (Ù…Ø«Ù„Ø§Ù‹ PostgreSQL Ø¹Ù„Ù‰ Render)ØŒ ÙˆØ¥Ù„Ø§ ÙŠØ³ØªØ®Ø¯Ù… SQLite Ø¯Ø§Ø®Ù„ instance/erp.db
try:
    os.makedirs(app.instance_path, exist_ok=True)
except Exception:
    pass

default_sqlite_path = os.path.join(app.instance_path, "erp.db")
default_sqlite_uri = f"sqlite:///{default_sqlite_path}"

# Ø¯Ø¹Ù… Ø±Ø¨Ø· PostgreSQL Ø¹Ø¨Ø± Ù…ØªØºÙŠØ± Ø§Ù„Ø¨ÙŠØ¦Ø© DATABASE_URL Ø¨Ù…Ø§ ÙŠØªÙˆØ§ÙÙ‚ Ù…Ø¹ SQLAlchemy
database_url = os.environ.get("DATABASE_URL", default_sqlite_uri)
if database_url.startswith("postgres://"):
    # ØªØ­ÙˆÙŠÙ„ Ø§Ù„ØµÙŠØºØ© Ø§Ù„Ù‚Ø¯ÙŠÙ…Ø© Ø¥Ù„Ù‰ Ù…Ø­Ø±Ùƒ SQLAlchemy psycopg (psycopg v3)
    database_url = database_url.replace("postgres://", "postgresql+psycopg://", 1)
elif database_url.startswith("postgresql://"):
    # ÙØ±Ø¶ Ø§Ø³ØªØ®Ø¯Ø§Ù… Ø¨Ø±Ù†Ø§Ù…Ø¬ ØªØ´ØºÙŠÙ„ psycopg v3 Ø­ØªÙ‰ Ù„Ùˆ ÙƒØ§Ù†Øª Ø§Ù„ØµÙŠØºØ© Ø­Ø¯ÙŠØ«Ø© Ø¨Ø¯ÙˆÙ† ØªØ­Ø¯ÙŠØ¯ Ø¯Ø±Ø§ÙŠÙØ±
    database_url = database_url.replace("postgresql://", "postgresql+psycopg://", 1)
elif database_url.startswith("postgresql+psycopg2://"):
    # ØªØ±Ù‚ÙŠØ© Ø£ÙŠ ØªÙˆØ¬ÙŠÙ‡ Ù‚Ø¯ÙŠÙ… Ø¥Ù„Ù‰ psycopg v3
    database_url = database_url.replace("postgresql+psycopg2://", "postgresql+psycopg://", 1)

# ÙØ±Ø¶ SSL Ø¹Ù„Ù‰ Render ÙˆÙ…Ø§ Ø´Ø§Ø¨Ù‡ Ø¥Ø°Ø§ Ù„Ù… ÙŠÙØ°ÙƒØ± ØµØ±Ø§Ø­Ø©Ù‹
if database_url.startswith("postgresql") and "sslmode=" not in database_url:
    connector = "&" if "?" in database_url else "?"
    database_url = f"{database_url}{connector}sslmode=require"

# ØªÙ…ÙƒÙŠÙ† keepalives Ùˆ connect_timeout Ù„ØªÙ‚Ù„ÙŠÙ„ Ø§Ù†Ù‚Ø·Ø§Ø¹Ø§Øª Ø§Ù„Ø§ØªØµØ§Ù„ Ø¹Ù„Ù‰ Postgres
if database_url.startswith("postgresql"):
    params_to_add = []
    if "keepalives=" not in database_url:
        params_to_add.append("keepalives=1")
    if "keepalives_idle=" not in database_url:
        params_to_add.append("keepalives_idle=30")
    if "keepalives_interval=" not in database_url:
        params_to_add.append("keepalives_interval=10")
    if "keepalives_count=" not in database_url:
        params_to_add.append("keepalives_count=3")
    if "connect_timeout=" not in database_url:
        params_to_add.append("connect_timeout=10")
    if params_to_add:
        connector = "&" if "?" in database_url else "?"
        database_url = f"{database_url}{connector}{'&'.join(params_to_add)}"

app.config["SQLALCHEMY_DATABASE_URI"] = database_url
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_pre_ping": True,
    "pool_recycle": 270,
    "pool_size": int(os.environ.get("SQL_POOL_SIZE", "5")),
    "max_overflow": int(os.environ.get("SQL_MAX_OVERFLOW", "10")),
    "pool_timeout": int(os.environ.get("SQL_POOL_TIMEOUT", "30")),
}
db.init_app(app)

# ---------------- Register Blueprints (Consulting) ----------------
# Ù…ÙˆØ¯ÙŠÙˆÙ„ Ø§Ù„Ø¹Ù…Ù„Ø§Ø¡ Ù„Ù‚Ø³Ù… Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª
from consulting.clients.routes import clients_bp
app.register_blueprint(clients_bp)

# Ù…ÙˆØ¯ÙŠÙˆÙ„ Ø§Ù„Ù…Ø´Ø§Ø±ÙŠØ¹ Ù„Ù‚Ø³Ù… Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª
from consulting.projects.routes import projects_bp
app.register_blueprint(projects_bp)

# Ù…ÙˆØ¯ÙŠÙˆÙ„ Ø§Ù„Ø¹Ù‚ÙˆØ¯ Ù„Ù‚Ø³Ù… Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª
from consulting.contracts.routes import contracts_bp
app.register_blueprint(contracts_bp)

# Ù…ÙˆØ¯ÙŠÙˆÙ„ Ø§Ù„Ù…Ø³ØªÙ†Ø¯Ø§Øª Ù„Ù‚Ø³Ù… Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª
from consulting.documents.routes import documents_bp
app.register_blueprint(documents_bp)

# Ù…ÙˆØ¯ÙŠÙˆÙ„ Ø§Ù„Ù…ÙˆØ§Ø±Ø¯ Ø§Ù„Ø¨Ø´Ø±ÙŠØ© (Ù…Ù‡Ù†Ø¯Ø³ÙˆÙ† ÙˆÙ…Ù‡Ø§Ù…) Ù„Ù‚Ø³Ù… Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª
from consulting.hr.routes import hr_bp
app.register_blueprint(hr_bp)

# Ù…ÙˆØ¯ÙŠÙˆÙ„ Ø§Ù„ÙÙˆØ§ØªÙŠØ± ÙˆØ§Ù„Ù…Ø­Ø§Ø³Ø¨Ø© Ù„Ù‚Ø³Ù… Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª
from consulting.invoices.routes import invoices_bp
app.register_blueprint(invoices_bp)

# Ù„ÙˆØ­Ø© Ø§Ù„Ù…ØªØ§Ø¨Ø¹Ø© Ù„Ù‚Ø³Ù… Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª
from consulting.dashboard.routes import dashboard_bp
app.register_blueprint(dashboard_bp)

# ---------------- Service Worker at root scope ----------------
@app.route('/service-worker.js')
def serve_service_worker():
    try:
        sw_path = os.path.join(app.root_path, 'static', 'service-worker.js')
        # Ensure correct content type and caching so browser picks updates
        response = send_file(sw_path, mimetype='application/javascript', max_age=0)
        response.headers['Service-Worker-Allowed'] = '/'
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        return response
    except Exception:
        return abort(404)

# ---------------- Ø¨Ø«/Ø¥Ø´Ø§Ø±Ø© ØªØ­Ø¯ÙŠØ« Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø§Øª (Ù†Ø³Ø®Ø© Ø¨Ø³ÙŠØ·Ø©) ----------------
TRANSACTIONS_VERSION = 0

def bump_transactions_version() -> None:
    global TRANSACTIONS_VERSION
    try:
        TRANSACTIONS_VERSION = (TRANSACTIONS_VERSION + 1) % 1_000_000_000
    except Exception:
        # fallback ÙÙŠ Ø­Ø§Ù„ Ø­Ø¯Ø« overflow ØºÙŠØ± Ù…ØªÙˆÙ‚Ø¹
        TRANSACTIONS_VERSION = int(datetime.utcnow().timestamp())

@app.route("/api/transactions/version")
def api_transactions_version():
    return jsonify({"version": TRANSACTIONS_VERSION, "ts": int(datetime.utcnow().timestamp())})

# -------- ØªØ¬Ø²Ø¦Ø© Ø§Ù„Ù…Ù„ÙØ§Øª Ù„Ù„ØªØ­Ù‚Ù‚ Ù…Ù† Ø³Ù„Ø§Ù…ØªÙ‡Ø§ --------

# ---------------- ØªÙˆÙ„ÙŠØ¯ Ø±Ù‚Ù… ÙØ§ØªÙˆØ±Ø© ÙØ±ÙŠØ¯ ----------------
def generate_unique_invoice_number(prefix: str = "INV", kind: str | None = None) -> str:
    """ÙŠÙˆÙ„Ù‘Ø¯ Ø±Ù‚Ù… ÙØ§ØªÙˆØ±Ø© ÙØ±ÙŠØ¯Ù‹Ø§ Ø¹Ù„Ù‰ Ù…Ø³ØªÙˆÙ‰ Ø§Ù„Ù†Ø¸Ø§Ù… Ø¨Ø´ÙƒÙ„ Ù…ØªØ³Ù„Ø³Ù„ Ø³Ù†ÙˆÙŠÙ‹Ø§.

    Ø§Ù„ØªÙ†Ø³ÙŠÙ‚: {prefix}-{YYYY}{optional-kind}-{NNNNN}
    Ø£Ù…Ø«Ù„Ø©: INV-2025-00001 Ø£Ùˆ INV-2025-CUST-00042
    ÙŠØ¶Ù…Ù† Ø¹Ø¯Ù… Ø§Ù„ØªÙƒØ±Ø§Ø± Ø¹Ø¨Ø± Ø¬Ù…ÙŠØ¹ Ø£Ù†ÙˆØ§Ø¹ Ø§Ù„ÙÙˆØ§ØªÙŠØ± Ø¨Ø§Ø³ØªØ®Ø¯Ø§Ù… Ø¬Ø¯ÙˆÙ„ invoice_sequence.
    """
    current_year = datetime.utcnow().year
    # Ø§Ø­ØµÙ„/Ø£Ù†Ø´Ø¦ ØµÙ Ø§Ù„Ø³Ù†Ø© Ø§Ù„Ø­Ø§Ù„ÙŠØ©
    seq = InvoiceSequence.query.filter_by(year=current_year).first()
    if not seq:
        seq = InvoiceSequence(year=current_year, last_number=0)
        db.session.add(seq)
        db.session.commit()

    # Ø²Ø¯ Ø§Ù„Ø±Ù‚Ù… Ø§Ù„Ù…ØªØ³Ù„Ø³Ù„
    seq.last_number = int(seq.last_number or 0) + 1
    db.session.commit()

    serial = f"{seq.last_number:05d}"
    if kind:
        return f"{prefix}-{current_year}-{kind}-{serial}"
    return f"{prefix}-{current_year}-{serial}"
def compute_file_sha256(file_path: str) -> str:
    """Ø¥Ø±Ø¬Ø§Ø¹ Ø¨ØµÙ…Ø© SHA-256 Ù„Ù…Ù„Ù ÙƒØ¨ÙŠØ± Ø¨Ø·Ø±ÙŠÙ‚Ø© ÙØ¹Ù‘Ø§Ù„Ø© Ø¨Ø§Ù„Ø°Ø§ÙƒØ±Ø©."""
    sha256 = hashlib.sha256()
    with open(file_path, 'rb') as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b''):
            sha256.update(chunk)
    return sha256.hexdigest()

# -------- Ø®ØªÙ… PDF Ø¨Ø´ÙƒÙ„ Ø¨Ø³ÙŠØ· Ø¹Ù„Ù‰ Ø§Ù„Ø®Ø§Ø¯Ù… --------
def stamp_pdf_with_seal(input_path: str, title: str, lines: List[str]) -> None:
    """ÙŠØ¶ÙŠÙ Ø®ØªÙ…Ù‹Ø§ Ù†ØµÙŠÙ‹Ø§ Ø¨Ø³ÙŠØ·Ù‹Ø§ Ø¹Ù„Ù‰ ÙƒÙ„ ØµÙØ­Ø© Ù…Ù† Ù…Ù„Ù PDF.

    - ÙŠØ±Ø³Ù… ØµÙ†Ø¯ÙˆÙ‚Ù‹Ø§ ÙÙŠ Ø§Ù„Ø²Ø§ÙˆÙŠØ© Ø§Ù„Ø¹Ù„ÙˆÙŠØ© Ø§Ù„ÙŠÙ…Ù†Ù‰
    - ÙŠÙƒØªØ¨ Ø¹Ù†ÙˆØ§Ù† Ø§Ù„Ø®ØªÙ… ÙˆØ¹Ø¯Ø© Ø£Ø³Ø·Ø± Ù…Ø¹Ù„ÙˆÙ…Ø§Øª
    ØªØ­ÙØ¸ Ø§Ù„Ù†ØªÙŠØ¬Ø© ÙÙˆÙ‚ Ù†ÙØ³ Ø§Ù„Ù…Ù„Ù.
    """
    try:
        doc = fitz.open(input_path)
        for page in doc:
            page_rect = page.rect
            # ØµÙ†Ø¯ÙˆÙ‚ Ø§Ù„Ø®ØªÙ… ÙÙŠ Ø£Ø¹Ù„Ù‰ ÙŠÙ…ÙŠÙ† Ø§Ù„ØµÙØ­Ø©
            margin = 20
            box_width = 220
            box_height = 120
            rect = fitz.Rect(
                page_rect.x1 - margin - box_width,
                margin,
                page_rect.x1 - margin,
                margin + box_height,
            )

            # Ø®Ù„ÙÙŠØ© ÙˆØµÙ†Ø¯ÙˆÙ‚
            page.draw_rect(rect, color=(0.8, 0.1, 0.1), fill=(1, 1, 1), width=1)

            # Ù†Øµ Ø§Ù„Ø®ØªÙ…
            content = title.strip()
            if lines:
                content += "\n" + "\n".join(str(x) for x in lines if x)

            # Ø¥Ø¯Ø±Ø§Ø¬ Ø§Ù„Ù†Øµ Ø¯Ø§Ø®Ù„ Ø§Ù„ØµÙ†Ø¯ÙˆÙ‚
            page.insert_textbox(
                rect.inflate(-8),
                content,
                fontsize=9,
                fontname="helv",
                color=(0, 0, 0),
                align=1,  # ÙˆØ³Ø·
            )

        doc.save(input_path, incremental=False, deflate=True)
        doc.close()
    except Exception:
        # ÙÙŠ Ø­Ø§Ù„ Ø­Ø¯ÙˆØ« Ø®Ø·Ø£ Ø¨Ø§Ù„Ø®ØªÙ…ØŒ Ù†ÙƒØªÙÙŠ Ø¨Ù…Ù„Ù Ø§Ù„Ø£ØµÙ„ Ø¯ÙˆÙ† Ø¥ÙŠÙ‚Ø§Ù Ø§Ù„Ø¹Ù…Ù„ÙŠØ©
        try:
            doc.close()
        except Exception:
            pass

# -------- ØªÙˆÙ„ÙŠØ¯ ØµÙˆØ±Ø© QR ÙƒÙ€ PNG (Ø¨Ø§ÙŠØªØ³) --------
def generate_qr_png_bytes(text: str, size: int = 100) -> bytes:
    """ÙŠÙ†Ø´Ø¦ ØµÙˆØ±Ø© QR ÙÙŠ Ø§Ù„Ø°Ø§ÙƒØ±Ø© ÙˆÙŠØ¹ÙŠØ¯Ù‡Ø§ ÙƒÙ€ PNG bytes.

    ÙŠØ­Ø§ÙˆÙ„ Ø£ÙˆÙ„Ø§Ù‹ Ø¨Ø§Ø³ØªØ®Ø¯Ø§Ù… ReportLab. ÙˆØ¥Ù† ÙØ´Ù„ØŒ ÙŠØ³ØªØ®Ø¯Ù… Ø®Ø¯Ù…Ø© Ø¹Ø§Ù…Ø© ÙƒØ­Ù„ Ø§Ø­ØªÙŠØ§Ø·ÙŠ.
    """
    try:
        widget = rl_qr.QrCodeWidget(text)
        bounds = widget.getBounds()
        width = bounds[2] - bounds[0]
        height = bounds[3] - bounds[1]
        scale = max(size / float(width), size / float(height))
        drawing = Drawing(width * scale, height * scale)
        widget.scale(scale, scale)
        drawing.add(widget)
        png_bytes = renderPM.drawToString(drawing, fmt='PNG')
        return png_bytes
    except Exception:
        # Ø§Ø­ØªÙŠØ§Ø·ÙŠ: ØªÙˆÙ„ÙŠØ¯ Ù…Ù† Ø®Ø¯Ù…Ø© Ø¹Ø§Ù…Ø©
        try:
            url = f"https://api.qrserver.com/v1/create-qr-code/?size={size}x{size}&data=" + requests.utils.quote(text, safe="")
            r = requests.get(url, timeout=10)
            r.raise_for_status()
            return r.content
        except Exception:
            # ÙƒØ­Ù„ Ø£Ø®ÙŠØ±ØŒ Ø£Ø¹ÙØ¯ Ø¨Ø§ÙŠØªØ³ ÙØ§Ø±ØºØ©
            return b""

# -------- Ø®ØªÙ… PDF ÙˆØ¥Ø¯Ø±Ø§Ø¬ QR ÙŠØ´ÙŠØ± Ø¥Ù„Ù‰ /file?hash=<hash> --------
def stamp_pdf_with_qr(input_path: str, hash_value: str) -> None:
    """ÙŠØ¶ÙŠÙ Ø¹Ù„Ø§Ù…Ø© Ù†ØµÙŠØ© ÙˆQR Ù„Ù„ØµÙØ­Ø© Ø§Ù„Ø£ÙˆÙ„Ù‰ ÙˆÙŠÙƒØªØ¨ Ù…Ù‚ØªØ·Ù Ø§Ù„Ø¨ØµÙ…Ø©.

    - QR ÙŠØ´ÙŠØ± Ø¥Ù„Ù‰ /file?hash=<hash_value>
    - Ù†Øµ Ù…Ø®ØªØµØ± Ù„Ù„Ø¨ØµÙ…Ø© ÙŠØ¸Ù‡Ø± ÙÙŠ Ø£Ø³ÙÙ„ Ø§Ù„ÙŠØ³Ø§Ø±
    ØªØ­ÙØ¸ Ø§Ù„Ù†ØªÙŠØ¬Ø© ÙÙˆÙ‚ Ù†ÙØ³ Ø§Ù„Ù…Ù„Ù.
    """
    try:
        doc = fitz.open(input_path)
        qr_link = url_for("file_by_hash", hash=hash_value, _external=True)
        qr_png = generate_qr_png_bytes(text=qr_link, size=100)
        for page_index, page in enumerate(doc):
            page_rect = page.rect
            # Ù†ØµÙˆØµ Ø³ÙÙ„ÙŠØ© ÙŠØ³Ø§Ø±
            try:
                page.insert_text(
                    fitz.Point(20, 35),
                    f"Hash: {hash_value[:10]}...",
                    fontsize=8,
                    fontname="helv",
                    color=(0, 0, 0),
                )
                page.insert_text(
                    fitz.Point(20, 20),
                    "Ù†Ø³Ø®Ø© Ø£ØµÙ„ÙŠØ© Ù„Ù„Ø¨Ù†Ùƒ",
                    fontsize=12,
                    fontname="helv",
                    color=(0, 0, 0),
                )
            except Exception:
                pass

            # QR ÙÙŠ Ø£Ø³ÙÙ„ ÙŠÙ…ÙŠÙ† Ø§Ù„ØµÙØ­Ø© Ø§Ù„Ø£ÙˆÙ„Ù‰ ÙÙ‚Ø·
            if page_index == 0 and qr_png:
                try:
                    qr_size = 100
                    margin = 20
                    rect = fitz.Rect(
                        page_rect.x1 - margin - qr_size,
                        margin,
                        page_rect.x1 - margin,
                        margin + qr_size,
                    )
                    page.insert_image(rect, stream=qr_png)
                except Exception:
                    pass

        doc.save(input_path, incremental=False, deflate=True)
        doc.close()
    except Exception:
        try:
            doc.close()
        except Exception:
            pass

# -------- Ø£Ø¯ÙˆØ§Øª Ù…Ø³Ø§Ø¹Ø¯Ø© Ù„Ù„Ø£Ø±Ù‚Ø§Ù… (ØªÙ‚Ø¨Ù„ Ø£Ø±Ù‚Ø§Ù… Ø¹Ø±Ø¨ÙŠØ© ÙˆÙÙˆØ§ØµÙ„) --------
def parse_float_input(value) -> float:
    """ØªØ­ÙˆÙŠÙ„ Ù…Ø¯Ø®Ù„ Ù†ØµÙŠ Ø¥Ù„Ù‰ Ø±Ù‚Ù… Ø¹Ø´Ø±ÙŠ Ù…Ø¹ Ø¯Ø¹Ù… Ø§Ù„Ø£Ø±Ù‚Ø§Ù… Ø§Ù„Ø¹Ø±Ø¨ÙŠØ© ÙˆØ§Ù„ÙÙˆØ§ØµÙ„.

    Ø£Ù…Ø«Ù„Ø© Ù…Ø¯Ø¹ÙˆÙ…Ø©:
    "12,345.67"  "12.345,67"  "Ù¡Ù¢Ù£Ù¤Ù¥Ù«Ù¦Ù§"  "Ù¡Ù¢Ù¬Ù£Ù¤Ù¥Ù«Ù¦Ù§"
    ØªØ¹Ø§Ø¯ 0.0 ÙÙŠ Ø­Ø§Ù„ Ø§Ù„ÙØ´Ù„.
    """
    if value is None:
        return 0.0
    try:
        s = str(value).strip()
        if not s:
            return 0.0
        # Ø®Ø±ÙŠØ·Ø© Ø§Ù„Ø£Ø±Ù‚Ø§Ù… Ø§Ù„Ø¹Ø±Ø¨ÙŠØ© Ø¥Ù„Ù‰ Ø§Ù„Ø¥Ù†Ø¬Ù„ÙŠØ²ÙŠØ©
        arabic_to_ascii = str.maketrans("Ù Ù¡Ù¢Ù£Ù¤Ù¥Ù¦Ù§Ù¨Ù©", "0123456789")
        s = s.translate(arabic_to_ascii)
        # Ø±Ù…ÙˆØ² Ø§Ù„ÙÙˆØ§ØµÙ„ Ø§Ù„Ø¹Ø±Ø¨ÙŠØ©
        arabic_thousands = "\u066C"  # Ù¬
        arabic_decimal   = "\u066B"  # Ù«
        # Ø¥Ø²Ø§Ù„Ø© Ø§Ù„Ù…Ø³Ø§ÙØ§Øª ÙˆØ£Ù†ÙˆØ§Ø¹ Ø§Ù„Ù…Ø³Ø§ÙØ§Øª ØºÙŠØ± Ø§Ù„Ù‚Ø§Ø¨Ù„Ø© Ù„Ù„ÙƒØ³Ø±
        s = s.replace(" ", "").replace("\u00A0", "").replace("\u202F", "")
        # Ø£Ø²Ù„ ÙÙˆØ§ØµÙ„ Ø§Ù„Ø¢Ù„Ø§Ù (Ø¹Ø±Ø¨ÙŠØ© Ø£Ùˆ Ø¥Ù†Ø¬Ù„ÙŠØ²ÙŠØ©)
        s = s.replace(",", "").replace(arabic_thousands, "")
        # ÙˆØ­Ù‘Ø¯ Ø§Ù„ÙØ§ØµÙ„Ø© Ø§Ù„Ø¹Ø´Ø±ÙŠØ© Ø¥Ù„Ù‰ Ù†Ù‚Ø·Ø©
        s = s.replace(arabic_decimal, ".")
        return float(s)
    except Exception:
        return 0.0


# ---------------- Ø§Ù„Ù†Ù…Ø§Ø°Ø¬ ----------------
class Branch(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    # Ø§Ù„Ù‚Ø³Ù…/Ø§Ù„Ù‚Ø·Ø§Ø¹ Ø§Ù„Ø®Ø§Øµ Ø¨Ø§Ù„ÙØ±Ø¹ (Ù…Ø«Ù„Ø§Ù‹: valuation | consultations | finance)
    department = db.Column(db.String(50), nullable=True)
    users = db.relationship("User", backref="branch", lazy=True)
    transactions = db.relationship("Transaction", backref="branch", lazy=True)
    # Ø£Ù‚Ø³Ø§Ù… Ù…ØªØ¹Ø¯Ø¯Ø© Ù…Ø±ØªØ¨Ø·Ø© Ø¨Ø§Ù„ÙØ±Ø¹ (Ø¨Ø¯Ù„Ø§Ù‹ Ù…Ù† Ø¹Ù…ÙˆØ¯ ÙˆØ§Ø­Ø¯ Ù‚Ø¯ÙŠÙ… department)
    sections = db.relationship("BranchSection", backref="branch", lazy=True, cascade="all, delete-orphan")

class BranchSection(db.Model):
    __tablename__ = "branch_section"
    id = db.Column(db.Integer, primary_key=True)
    branch_id = db.Column(db.Integer, db.ForeignKey("branch.id"), nullable=False, index=True)
    name = db.Column(db.String(50), nullable=False)  # valuation | consultations
    # ğŸ§‘â€ğŸ’¼ Ø§Ù„Ù…Ø³ØªØ®Ø¯Ù…ÙˆÙ† Ø§Ù„Ù…Ù†ØªÙ…ÙˆÙ† Ù„Ù‡Ø°Ø§ Ø§Ù„Ù‚Ø³Ù…
    users = db.relationship("User", backref="section", lazy=True)

class Bank(db.Model):
    __tablename__ = "bank"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False, unique=True)

    # Ø¹Ù„Ø§Ù‚Ø© ÙˆØ§Ø­Ø¯Ø© ÙÙ‚Ø·ØŒ ÙˆÙ…Ø§ Ù†ÙƒØ±Ø±Ù‡Ø§ ÙÙŠ Transaction
    transactions = db.relationship("Transaction", backref="bank", lazy=True)

class User(db.Model):
    __tablename__ = "user"
    id       = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)
    role     = db.Column(db.String(20), nullable=False)  # manager/employee/visit/engineer/finance
    branch_id = db.Column(db.Integer, db.ForeignKey('branch.id'), nullable=True)
    # ğŸ†• Ø±Ø¨Ø· Ø§Ù„Ù…ÙˆØ¸Ù Ø¨Ù‚Ø³Ù… Ø¯Ø§Ø®Ù„ Ø§Ù„ÙØ±Ø¹ (Ø§Ø®ØªÙŠØ§Ø±ÙŠ)
    section_id = db.Column(db.Integer, db.ForeignKey('branch_section.id'), nullable=True, index=True)

class Transaction(db.Model):
    __tablename__ = "transaction"
    id              = db.Column(db.Integer, primary_key=True)
    client          = db.Column(db.String(100))
    employee        = db.Column(db.String(50))
    date            = db.Column(db.DateTime, default=datetime.utcnow)
    status          = db.Column(db.String(30), default="Ù…Ø¹Ù„Ù‚Ø©")
    fee             = db.Column(db.Float, default=0)
    land_value      = db.Column(db.Float, default=0)
    building_value  = db.Column(db.Float, default=0)
    total_estimate  = db.Column(db.Float, default=0)
    files           = db.Column(db.Text)
    # Ù…Ù„ÙØ§Øª Ø£ÙØ±Ø³Ù„Øª Ù„Ù„Ø¨Ù†Ùƒ Ø¨ÙˆØ§Ø³Ø·Ø© Ø§Ù„Ù…ÙˆØ¸Ù (Ù‚Ø§Ø¦Ù…Ø© Ø£Ø³Ù…Ø§Ø¡ Ù…ÙØµÙˆÙ„Ø© Ø¨ÙÙˆØ§ØµÙ„)
    bank_sent_files = db.Column(db.Text)
    area            = db.Column(db.Float, default=0)
    building_area   = db.Column(db.Float, default=0)
    building_age    = db.Column(db.Integer, default=0)
    report_file     = db.Column(db.String(200))
    report_number   = db.Column(db.String(50))
    sent_to_engineer_at = db.Column(db.DateTime, nullable=True)
    engineer_report = db.Column(db.Text, nullable=True)  # ØªÙ‚Ø±ÙŠØ± Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³
    transaction_type = db.Column(db.String(50), default="real_estate")  
    vehicle_type  = db.Column(db.String(100))
    vehicle_model = db.Column(db.String(100))
    vehicle_year  = db.Column(db.String(20))
    type = db.Column(db.String(50))          # Ù†ÙˆØ¹ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© (Ø¹Ù‚Ø§Ø±ØŒ Ø³ÙŠØ§Ø±Ø© â€¦)
    valuation_amount = db.Column(db.Float)   # Ù…Ø¨Ù„Øº Ø§Ù„ØªØ«Ù…ÙŠÙ†
    state = db.Column(db.String(100), nullable=True)   # Ø§Ù„ÙˆÙ„Ø§ÙŠØ©
    region = db.Column(db.String(100), nullable=True)  # Ø§Ù„Ù…Ù†Ø·Ù‚Ø©
    
    # ğŸ‘‡ Ù‡Ù†Ø§ ÙÙ‚Ø· Ù…ÙØªØ§Ø­ Ø®Ø§Ø±Ø¬ÙŠ ÙŠØ±Ø¨Ø· Ø¨Ø§Ù„Ø¬Ø¯ÙˆÙ„ Bank
    bank_id = db.Column(db.Integer, db.ForeignKey("bank.id"), nullable=True)
    # ğŸ‘‡ Ø§Ø³Ù… ÙØ±Ø¹ Ø§Ù„Ø¨Ù†Ùƒ Ø§Ù„Ù…Ø±ØªØ¨Ø· Ø¨Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø©
    bank_branch = db.Column(db.String(120), nullable=True)
    # ğŸ‘‡ Ø§Ø³Ù… Ù…ÙˆØ¸Ù Ø§Ù„Ø¨Ù†Ùƒ Ø§Ù„Ø°ÙŠ Ø¬Ù„Ø¨/Ù‚Ø¯Ù‘Ù… Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø©
    bank_employee_name = db.Column(db.String(120), nullable=True)
    # ğŸ‘‡ Ø§Ù„Ù…ÙˆØ¸Ù Ø§Ù„Ø°ÙŠ Ø¬Ù„Ø¨ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© (Ù…Ù† Ø¯Ø§Ø®Ù„ Ø´Ø±ÙƒØªÙ†Ø§)
    brought_by = db.Column(db.String(120), nullable=True)
    # ğŸ‘‡ Ø§Ù„Ø´Ø®Øµ Ø§Ù„Ø°ÙŠ Ù‚Ø§Ù… Ø¨Ø§Ù„Ø²ÙŠØ§Ø±Ø©
    visited_by = db.Column(db.String(120), nullable=True)

    price = db.Column(db.Float, nullable=True)   # Ø³Ø¹Ø± Ø§Ù„ØªØ«Ù…ÙŠÙ† (Ø§Ø®ØªÙŠØ§Ø±ÙŠ)

    assigned_to = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=True)
    created_by  = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    branch_id   = db.Column(db.Integer, db.ForeignKey("branch.id"), nullable=False)

    payment_status  = db.Column(db.String(20), default="ØºÙŠØ± Ù…Ø¯ÙÙˆØ¹Ø©")

    # Ø¨ØµÙ…Ø© Ø§Ù„ØªÙ‚Ø±ÙŠØ± (SHA-256) Ù„Ù„ØªØ­Ù‚Ù‚ Ù…Ù† Ø¹Ø¯Ù… Ø§Ù„Ø¹Ø¨Ø«
    report_sha256 = db.Column(db.String(64), nullable=True)
    # Ø±Ø§Ø¨Ø· Ù…Ø´Ø§Ø±ÙƒØ© Ø¹Ø§Ù… (Token)
    public_share_token = db.Column(db.String(128), nullable=True)
    # Ù…Ø¹Ù„ÙˆÙ…Ø§Øª Backblaze B2 Ù„Ù…Ù„Ù Ø§Ù„ØªÙ‚Ø±ÙŠØ±
    report_b2_file_name = db.Column(db.String(255), nullable=True)
    report_b2_file_id = db.Column(db.String(255), nullable=True)

    payments = db.relationship("Payment", backref="transaction", lazy=True)


class NotificationSubscription(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    subscription_json = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)



class LandPrice(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    state = db.Column(db.String(120))      # Ø§Ù„ÙˆÙ„Ø§ÙŠØ©
    region = db.Column(db.String(120))     # Ø§Ù„Ù…Ù†Ø·Ù‚Ø©
    bank_id = db.Column(db.Integer, db.ForeignKey("bank.id"))
    price_per_meter = db.Column(db.Float)  # Ø³Ø¹Ø± Ø§Ù„Ù…ØªØ±
    last_updated = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

class Payment(db.Model):
    __tablename__ = "payment"
    id             = db.Column(db.Integer, primary_key=True)
    transaction_id = db.Column(db.Integer, db.ForeignKey('transaction.id'))
    amount         = db.Column(db.Float, default=0)
    date_received  = db.Column(db.DateTime, default=datetime.utcnow)
    received_by    = db.Column(db.String(50))
    method         = db.Column(db.String(20))   # ÙƒØ§Ø´ / ØªØ­ÙˆÙŠÙ„
    receipt_file   = db.Column(db.String(200))  # ØµÙˆØ±Ø© Ø£Ùˆ Ù…Ù„Ù Ø§Ù„Ø¥ÙŠØµØ§Ù„
    branch_id      = db.Column(db.Integer, db.ForeignKey("branch.id"), nullable=True)


class ReportTemplate(db.Model):
    __tablename__ = "report_template"
    id = db.Column(db.Integer, primary_key=True)
    template_type = db.Column(db.String(50), nullable=False)  # real_estate / vehicle
    content = db.Column(db.Text, nullable=True)
    title = db.Column(db.String(150), nullable=True)
    file = db.Column(db.String(255), nullable=True)  # Ù…Ø³Ø§Ø± Ù…Ù„Ù DOCX Ø§Ù„Ù…Ø±ÙÙˆØ¹ Ø¥Ù† ÙˆÙØ¬Ø¯


# ØªØ³Ù„Ø³Ù„ Ø±Ù‚Ù…ÙŠ Ø¹Ø§Ù… Ù„Ù„ÙÙˆØ§ØªÙŠØ± Ø¨Ø­Ø³Ø¨ Ø§Ù„Ø³Ù†Ø©
class InvoiceSequence(db.Model):
    __tablename__ = "invoice_sequence"
    id = db.Column(db.Integer, primary_key=True)
    year = db.Column(db.Integer, unique=True, nullable=False)
    last_number = db.Column(db.Integer, nullable=False, default=0)

# ÙÙˆØ§ØªÙŠØ± Ø§Ù„Ø¨Ù†Ùƒ Ø¨Ù…Ø±Ø§Ø­Ù„Ù‡Ø§
class BankInvoice(db.Model):
    __tablename__ = "bank_invoice"
    id = db.Column(db.Integer, primary_key=True)
    bank_id = db.Column(db.Integer, db.ForeignKey("bank.id"), nullable=False)
    transaction_id = db.Column(db.Integer, db.ForeignKey("transaction.id"), nullable=True)
    amount = db.Column(db.Float, default=0)
    issued_at = db.Column(db.DateTime, nullable=True)     # Ø§Ù„Ù…Ø±Ø­Ù„Ø© 1: Ø¥ØµØ¯Ø§Ø±
    delivered_at = db.Column(db.DateTime, nullable=True)  # Ø§Ù„Ù…Ø±Ø­Ù„Ø© 2: ØªØ³Ù„ÙŠÙ…
    received_at = db.Column(db.DateTime, nullable=True)   # Ø§Ù„Ù…Ø±Ø­Ù„Ø© 3: Ø§Ø³ØªÙ„Ø§Ù… Ø§Ù„Ù…Ø¨Ù„Øº
    note = db.Column(db.String(255))
    # Ø±Ù‚Ù… ÙØ§ØªÙˆØ±Ø© Ù…ÙˆØ­Ù‘Ø¯ Ø¹Ø¨Ø± Ø§Ù„Ù†Ø¸Ø§Ù…
    invoice_number = db.Column(db.String(50), unique=True, nullable=True)

class Quote(db.Model):
    __tablename__ = "quote"
    id = db.Column(db.Integer, primary_key=True)
    bank_id = db.Column(db.Integer, db.ForeignKey("bank.id"), nullable=False)
    transaction_id = db.Column(db.Integer, db.ForeignKey("transaction.id"), nullable=True)
    amount = db.Column(db.Float, default=0)
    valid_until = db.Column(db.DateTime, nullable=True)
    note = db.Column(db.String(255))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    created_by = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=True)

class CustomerInvoice(db.Model):
    __tablename__ = "customer_invoice"
    id = db.Column(db.Integer, primary_key=True)
    customer_name = db.Column(db.String(150), nullable=False)
    amount = db.Column(db.Float, default=0)
    issued_at = db.Column(db.DateTime, default=datetime.utcnow)
    transaction_id = db.Column(db.Integer, db.ForeignKey("transaction.id"), nullable=True)
    note = db.Column(db.String(255))
    created_by = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=True)
    # Ø±Ù‚Ù… ÙØ§ØªÙˆØ±Ø© Ù…ÙˆØ­Ù‘Ø¯ Ø¹Ø¨Ø± Ø§Ù„Ù†Ø¸Ø§Ù…
    invoice_number = db.Column(db.String(50), unique=True, nullable=True)

class CustomerQuote(db.Model):
    __tablename__ = "customer_quote"
    id = db.Column(db.Integer, primary_key=True)
    customer_name = db.Column(db.String(150), nullable=False)
    amount = db.Column(db.Float, default=0)
    valid_until = db.Column(db.DateTime, nullable=True)
    transaction_id = db.Column(db.Integer, db.ForeignKey("transaction.id"), nullable=True)
    note = db.Column(db.String(255))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    created_by = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=True)

# ---------------- ØªÙ†Ø¸ÙŠÙ Ø¹Ø±ÙˆØ¶ Ø§Ù„Ø£Ø³Ø¹Ø§Ø± Ù…Ù†ØªÙ‡ÙŠØ© Ø§Ù„ØµÙ„Ø§Ø­ÙŠØ© ----------------
# Ù…Ù„Ø§Ø­Ø¸Ø©: Ù†Ø³ØªØ®Ø¯Ù… Ø®Ø§Ù†Ø© global Ù„ØªÙ‚Ù„ÙŠÙ„ ØªÙƒØ±Ø§Ø± Ø§Ù„ØªÙ†ÙÙŠØ° Ø¹Ù„Ù‰ ÙƒÙ„ Ø·Ù„Ø¨
last_quotes_purge_at = None
PURGE_INTERVAL_SECONDS = 600  # 10 Ø¯Ù‚Ø§Ø¦Ù‚

def purge_expired_quotes() -> None:
    """ÙŠØ­Ø°Ù ÙƒØ§ÙØ© Ø¹Ø±ÙˆØ¶ Ø§Ù„Ø£Ø³Ø¹Ø§Ø± Ø§Ù„ØªÙŠ Ø§Ù†ØªÙ‡Ù‰ ØªØ§Ø±ÙŠØ® ØµÙ„Ø§Ø­ÙŠØªÙ‡Ø§.

    ÙŠÙØ¹ØªØ¨Ø± Ø§Ù„Ø¹Ø±Ø¶ Ù…Ù†ØªÙ‡ÙŠÙ‹Ø§ Ø¥Ø°Ø§ ÙƒØ§Ù†Øª Ù„Ù‡ Ù‚ÙŠÙ…Ø© valid_until ÙˆØ£ØµØ¨Ø­Øª Ø£Ù‚Ù„ Ù…Ù† Ø§Ù„Ø¢Ù† (UTC).
    ØªÙ†Ø·Ø¨Ù‚ Ø§Ù„Ø¹Ù…Ù„ÙŠØ© Ø¹Ù„Ù‰ Ø¬Ø¯ÙˆÙ„ÙŠ Quote Ùˆ CustomerQuote.
    """
    now_utc = datetime.utcnow()
    try:
        deleted_quotes = Quote.query \
            .filter(Quote.valid_until != None, Quote.valid_until < now_utc) \
            .delete(synchronize_session=False)

        deleted_customer_quotes = CustomerQuote.query \
            .filter(CustomerQuote.valid_until != None, CustomerQuote.valid_until < now_utc) \
            .delete(synchronize_session=False)

        if (deleted_quotes or 0) > 0 or (deleted_customer_quotes or 0) > 0:
            db.session.commit()
    except Exception:
        # ÙÙŠ Ø­Ø§Ù„ Ø­Ø¯ÙˆØ« Ø£ÙŠ Ø®Ø·Ø£ØŒ Ù†Ø±Ø¬Ø¹ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ù„Ø­Ø§Ù„ØªÙ‡Ø§ Ø§Ù„Ø³Ø§Ø¨Ù‚Ø©
        db.session.rollback()

@app.before_request
def auto_purge_expired_quotes():
    global last_quotes_purge_at
    now_utc = datetime.utcnow()
    if last_quotes_purge_at is None or (now_utc - last_quotes_purge_at).total_seconds() >= PURGE_INTERVAL_SECONDS:
        purge_expired_quotes()
        last_quotes_purge_at = now_utc

class ValuationMemory(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    state = db.Column(db.String(100), nullable=False)   # Ø§Ù„ÙˆÙ„Ø§ÙŠØ©
    region = db.Column(db.String(100), nullable=False)  # Ø§Ù„Ù…Ù†Ø·Ù‚Ø©
    bank_id = db.Column(db.Integer, nullable=False)     # Ø§Ù„Ø¨Ù†Ùƒ
    price_per_meter = db.Column(db.Float, nullable=False)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

class Expense(db.Model):
    __tablename__ = "expense"
    id = db.Column(db.Integer, primary_key=True)
    description = db.Column(db.String(200))
    amount = db.Column(db.Float, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    file   = db.Column(db.String(200))
    branch_id = db.Column(db.Integer, db.ForeignKey("branch.id"))
    branch = db.relationship("Branch", backref="expenses")

class BranchDocument(db.Model):
    __tablename__ = "branch_document"
    id = db.Column(db.Integer, primary_key=True)
    branch_id = db.Column(db.Integer, db.ForeignKey("branch.id"), nullable=False)
    title = db.Column(db.String(200), nullable=False)
    doc_type = db.Column(db.String(100), nullable=True)
    file = db.Column(db.String(255), nullable=True)
    # Backblaze B2 identifiers for the uploaded file (if stored on B2)
    b2_file_name = db.Column(db.String(255), nullable=True)
    b2_file_id = db.Column(db.String(255), nullable=True)
    issued_at = db.Column(db.DateTime, nullable=True)
    expires_at = db.Column(db.DateTime, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    branch = db.relationship("Branch", backref="documents")

# âœ… Ù…Ø³ØªÙ†Ø¯Ø§Øª Ø¹Ø§Ù…Ø© Ù…Ø±Ø³Ù„Ø© Ø¥Ù„Ù‰ Ø§Ù„Ø¨Ù†ÙˆÙƒ (ØºÙŠØ± Ù…Ø±ØªØ¨Ø·Ø© Ø¨Ù…Ø¹Ø§Ù…Ù„Ø©)
class BankDocument(db.Model):
    __tablename__ = "bank_document"
    id = db.Column(db.Integer, primary_key=True)
    bank_id = db.Column(db.Integer, db.ForeignKey("bank.id"), nullable=False)
    title = db.Column(db.String(200), nullable=False)
    message = db.Column(db.Text, nullable=True)
    doc_type = db.Column(db.String(100), nullable=True)  # Ø±Ø³Ø§Ù„Ø©ØŒ Ø³ÙŠØ±Ø© Ø°Ø§ØªÙŠØ©ØŒ ...
    file = db.Column(db.String(255), nullable=True)
    # Backblaze B2 identifiers for the uploaded file (if stored on B2)
    b2_file_name = db.Column(db.String(255), nullable=True)
    b2_file_id = db.Column(db.String(255), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    created_by = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=True)
    branch_id = db.Column(db.Integer, db.ForeignKey("branch.id"), nullable=True)

# âœ… Ø¬Ø¯ÙˆÙ„ Ø¨Ø³ÙŠØ· Ù„Ø­ÙØ¸ Ø§Ù„Ø¹Ù…Ù„Ø§Ø¡ (Ø§Ø³Ù… ÙˆØ±Ù‚Ù…)
class Customer(db.Model):
    __tablename__ = "customer"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(150), nullable=False)
    phone = db.Column(db.String(50), nullable=False)

# âœ… Ø­ÙØ¸ Ù…Ù„ÙØ§Øª Ù‚ÙˆØ§Ù„Ø¨ Ø§Ù„ÙˆÙˆØ±Ø¯ Ù„Ù„ÙÙˆØ§ØªÙŠØ± ÙˆØ¹Ø±ÙˆØ¶ Ø§Ù„Ø£Ø³Ø¹Ø§Ø±
class TemplateDoc(db.Model):
    __tablename__ = "template_doc"
    id = db.Column(db.Integer, primary_key=True)
    doc_type = db.Column(db.String(50), nullable=False)  # invoice | quote
    filename = db.Column(db.String(255), nullable=False)  # Ø§Ø³Ù… Ø§Ù„Ù…Ù„Ù Ø¯Ø§Ø®Ù„ uploads
    uploaded_at = db.Column(db.DateTime, default=datetime.utcnow)
    # ğŸ†• Ù‚Ø§Ù„Ø¨ Ø®Ø§Øµ Ø¨ÙØ±Ø¹ Ù…Ø¹ÙŠÙ† (Ø§Ø®ØªÙŠØ§Ø±ÙŠ). Ù„Ùˆ ÙƒØ§Ù†Øª NULL ÙÙ‡Ùˆ Ù‚Ø§Ù„Ø¨ Ø¹Ø§Ù…
    branch_id = db.Column(db.Integer, db.ForeignKey("branch.id"), nullable=True)

# ================= Consulting Department Models =================
# Status and types constants (kept simple strings for SQLite compatibility)
CONSULTATION_STATUSES = [
    "Pending",
    "In Progress",
    "Completed",
]

CONSULTATION_TYPES = [
    "Architectural",
    "Structural",
    "Mechanical",
    "Electrical",
    "Other",
]


class Project(db.Model):
    __tablename__ = "project"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False, unique=True)
    description = db.Column(db.Text, nullable=True)
    client_id = db.Column(db.Integer, db.ForeignKey("customer.id"), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    client = db.relationship("Customer", backref=db.backref("projects", lazy=True))


class Consultation(db.Model):
    __tablename__ = "consultation"
    id = db.Column(db.Integer, primary_key=True)
    project_id = db.Column(
        db.Integer,
        db.ForeignKey("consulting_project.id"),
        nullable=True,
        index=True,
    )
    client_id = db.Column(
        db.Integer,
        db.ForeignKey("consulting_client.id"),
        nullable=True,
        index=True,
    )
    consultant_name = db.Column(db.String(150), nullable=True)
    consultation_type = db.Column(db.String(50), nullable=False)
    description = db.Column(db.Text, nullable=True)
    status = db.Column(db.String(20), nullable=False, default="Pending", index=True)
    start_date = db.Column(db.Date, nullable=True, index=True)
    end_date = db.Column(db.Date, nullable=True, index=True)
    cost = db.Column(db.Float, nullable=True)
    created_by = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=True)
    consultant_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    # Ø§Ø±Ø¨Ø· Ø¨Ù†Ù…Ø§Ø°Ø¬ Ù‚Ø³Ù… Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª (ConsultingProject / Client)
    project = db.relationship(
        ConsultingProject, backref=db.backref("consultations", lazy=True)
    )
    client = db.relationship(Client, backref=db.backref("consultations", lazy=True))
    creator = db.relationship("User", foreign_keys=[created_by])
    consultant = db.relationship("User", foreign_keys=[consultant_id])

def replace_placeholders_in_docx(doc: Document, replacements: dict) -> None:
    # ÙŠØ¯Ø¹Ù… Ø§Ù„Ø§Ø³ØªØ¨Ø¯Ø§Ù„ Ø­ØªÙ‰ Ù„Ùˆ ÙˆÙØ¬Ø¯Øª Ù…Ø³Ø§ÙØ§Øª/Ø¹Ù„Ø§Ù…Ø§Øª RTL Ø¯Ø§Ø®Ù„ Ø§Ù„Ø£Ù‚ÙˆØ§Ø³
    # Ù†Ø¨Ù†ÙŠ Ø®Ø±ÙŠØ·Ø© Ø¨Ø§Ù„Ø§Ø³Ù… Ø¨Ø¯ÙˆÙ† Ø§Ù„Ø£Ù‚ÙˆØ§Ø³ ÙˆØ¨Ø­Ø±ÙˆÙ ÙƒØ¨ÙŠØ±Ø©
    import re
    zero_width = "\u200c\u200d\u200e\u200f\u202a\u202b\u202c\u202d\u202e\u2066\u2067\u2068\u2069"
    def strip_braces_and_controls(s: str) -> str:
        s = str(s)
        s = s.replace("{", "").replace("}", "")
        s = re.sub(rf"[{zero_width}]", "", s)
        return s

    # Ø§Ø¨Ù†Ù Ù‚Ø§Ù…ÙˆØ³Ù‹Ø§ Ø¨Ù…ØªØºÙŠØ±Ø§Øª Ù…ØªØ¹Ø¯Ø¯Ø© Ø§Ù„Ø£Ø´ÙƒØ§Ù„ Ù„Ù†ÙØ³ Ø§Ù„Ù…ÙØªØ§Ø­
    token_to_value = {}
    for k, v in replacements.items():
        base = strip_braces_and_controls(k).strip()
        variants = set()
        variants.add(base)
        variants.add(base.replace(" ", ""))
        variants.add(base.replace("_", " "))
        variants.add(base.replace("_", ""))
        for var in variants:
            token_to_value[var.upper()] = str(v)
    # Ù†Ù…Ø· ÙŠÙ„ØªÙ‚Ø· { TOKEN } Ù…Ø¹ Ø§Ø­ØªÙ…Ø§Ù„ÙŠØ© ÙˆØ¬ÙˆØ¯ Ù…Ø³Ø§ÙØ§Øª/Ø¹Ù„Ø§Ù…Ø§Øª Ø§ØªØ¬Ø§Ù‡ Ø¯Ø§Ø®Ù„ Ø§Ù„Ø£Ù‚ÙˆØ§Ø³
    # ÙŠØ¯Ø¹Ù… {TOKEN} Ùˆ {{TOKEN}} ÙˆÙŠØ³ØªØ«Ù†ÙŠ Ø§Ù„Ø£Ù‚ÙˆØ§Ø³ Ø¯Ø§Ø®Ù„ Ø§Ù„Ø§Ø³Ù…
    # Ù…Ù„Ø§Ø­Ø¸Ø©: Ù†ØªØ¬Ù†Ø¨ Ø§Ø³ØªØ®Ø¯Ø§Ù… f-string Ù‡Ù†Ø§ Ø¨Ø³Ø¨Ø¨ Ø§Ù„Ø£Ù‚ÙˆØ§Ø³ Ø§Ù„Ù…ØªØ¯Ø§Ø®Ù„Ø© ÙÙŠ regex
    token_pattern = re.compile(
        r"\{{{1,2}}[\s" + zero_width + r"]*([^{}]+?)[\s" + zero_width + r"]*\}}{{{1,2}}"
    )

    def replace_in_paragraph(paragraph) -> None:
        combined_text = "".join(run.text for run in paragraph.runs) or paragraph.text
        if not combined_text:
            return
        def _repl(m):
            raw_name = strip_braces_and_controls(m.group(1)).strip().upper()
            return (
                token_to_value.get(raw_name)
                or token_to_value.get(raw_name.replace(" ", ""))
                or token_to_value.get(raw_name.replace(" ", "_"))
                or token_to_value.get(raw_name.replace("_", ""))
                or m.group(0)
            )
        new_text = token_pattern.sub(_repl, combined_text)
        # ØªÙ…Ø±ÙŠØ± Ø§Ø­ØªÙŠØ§Ø·ÙŠ: Ø§Ø³ØªØ¨Ø¯Ø§Ù„ Ù…Ø¨Ø§Ø´Ø± Ù„Ø£ÙŠ Ù…ÙØ§ØªÙŠØ­ Ù…Ù‚Ø¯Ù‘ÙÙ…Ø© ÙƒÙ…Ø§ Ù‡ÙŠ
        if new_text == combined_text:
            for raw_key, raw_val in replacements.items():
                if raw_key and isinstance(raw_key, str) and raw_key in new_text:
                    new_text = new_text.replace(raw_key, str(raw_val))
        if new_text != combined_text:
            paragraph.text = new_text

    def replace_in_table(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)
                for nested in getattr(cell, "tables", []):
                    replace_in_table(nested)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        replace_in_table(table)

    for section in getattr(doc, "sections", []):
        header = getattr(section, "header", None)
        if header:
            for p in header.paragraphs:
                replace_in_paragraph(p)
            for t in header.tables:
                replace_in_table(t)
        footer = getattr(section, "footer", None)
        if footer:
            for p in footer.paragraphs:
                replace_in_paragraph(p)
            for t in footer.tables:
                replace_in_table(t)

def ensure_template_doc_branch_column():
    # Ø¥Ø¶Ø§ÙØ© Ø§Ù„Ø¹Ù…ÙˆØ¯ branch_id Ø¥Ù† Ù„Ù… ÙŠÙƒÙ† Ù…ÙˆØ¬ÙˆØ¯Ù‹Ø§ (SQLite: ALTER TABLE ADD COLUMN)
    try:
        if not column_exists("template_doc", "branch_id"):
            db.session.execute(text("ALTER TABLE template_doc ADD COLUMN branch_id INTEGER"))
            db.session.commit()
    except Exception:
        db.session.rollback()

def ensure_branch_department_column():
    """Ø¶Ù…Ø§Ù† ÙˆØ¬ÙˆØ¯ Ø¹Ù…ÙˆØ¯ department ÙÙŠ Ø¬Ø¯ÙˆÙ„ Ø§Ù„ÙØ±ÙˆØ¹ (Ù„Ù„ØªÙˆØ§ÙÙ‚ Ù…Ø¹ Ù‚ÙˆØ§Ø¹Ø¯ Ø¨ÙŠØ§Ù†Ø§Øª Ù‚Ø¯ÙŠÙ…Ø©)."""
    try:
        if not column_exists("branch", "department"):
            db.session.execute(text("ALTER TABLE branch ADD COLUMN department VARCHAR(50)"))
            db.session.commit()
    except Exception:
        db.session.rollback()

def ensure_branch_sections_from_department():
    """Ø¥Ù†Ø´Ø§Ø¡ Ø³Ø¬Ù„Ø§Øª Ø£Ù‚Ø³Ø§Ù… Ù„Ù„ÙØ±Ø¹ Ù…Ù† Ø§Ù„Ù‚ÙŠÙ…Ø© Ø§Ù„Ù‚Ø¯ÙŠÙ…Ø© branch.department (Ù…Ø±Ø© ÙˆØ§Ø­Ø¯Ø©).

    Ù†Ø·Ø¨Ø¹ Ø§Ù„Ø£Ù‚Ø³Ø§Ù… Ø§Ù„Ù…Ù‚Ø¨ÙˆÙ„Ø© ÙÙ‚Ø· (valuation | consultations). Ø§Ù„Ù…Ø§Ù„ÙŠØ© Ø«Ø§Ø¨ØªØ© Ù„Ù„ÙØ±Ø¹
    ÙˆÙ„Ø§ ØªØ­ØªØ§Ø¬ Ù„ØªØ³Ø¬ÙŠÙ„ ÙƒÙ‚Ø³Ù… Ù…Ù†ÙØµÙ„.
    """
    try:
        # ØªØ£ÙƒØ¯ Ø£Ù† Ø¬Ø¯ÙˆÙ„ Ø§Ù„Ø£Ù‚Ø³Ø§Ù… Ù…ÙˆØ¬ÙˆØ¯
        db.session.execute(text("SELECT 1 FROM branch_section LIMIT 1"))
    except Exception:
        # ÙÙŠ Ø­Ø§Ù„ Ù„Ù… ÙŠØªÙ… Ø¥Ù†Ø´Ø§Ø¤Ù‡ Ø¨Ø¹Ø¯ØŒ Ù„Ø¹Ù„ create_all Ø³ÙŠÙ‚ÙˆÙ… Ø¨Ø°Ù„Ùƒ
        try:
            db.create_all()
        except Exception:
            pass

    try:
        branches = Branch.query.all()
        for b in branches:
            dept = (getattr(b, "department", None) or "").strip().lower()
            if not dept:
                continue
            normalized = None
            if dept in ("consultations", "consultation", "consulting", "Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª"):
                normalized = "consultations"
            elif dept in ("valuation", "Ø§Ù„ØªØ«Ù…ÙŠÙ†", "ØªØ«Ù…ÙŠÙ†"):
                normalized = "valuation"
            # Ø§Ù„Ù…Ø§Ù„ÙŠØ© Ù‚Ø³Ù… Ø«Ø§Ø¨Øª Ø¹Ù„Ù‰ Ù…Ø³ØªÙˆÙ‰ Ø§Ù„ÙØ±Ø¹ØŒ Ù†ØªØ¬Ø§Ù‡Ù„Ù‡ Ù‡Ù†Ø§
            if not normalized:
                continue
            exists = BranchSection.query.filter_by(branch_id=b.id, name=normalized).first()
            if not exists:
                db.session.add(BranchSection(branch_id=b.id, name=normalized))
        db.session.commit()
    except Exception:
        db.session.rollback()

def get_template_filename(doc_type: str, branch_id: int | None = None) -> str | None:
    # ÙŠÙØ¶Ù‘Ù„ Ø§Ù„Ù‚Ø§Ù„Ø¨ Ø§Ù„Ø®Ø§Øµ Ø¨Ø§Ù„ÙØ±Ø¹ Ø¥Ù† ÙˆØ¬Ø¯ØŒ Ø«Ù… ÙŠØ¹ÙˆØ¯ Ù„Ù„Ù‚Ø§Ù„Ø¨ Ø§Ù„Ø¹Ø§Ù…
    ensure_template_doc_branch_column()
    base_q = TemplateDoc.query.filter(TemplateDoc.doc_type == doc_type)
    if branch_id is not None:
        rec = base_q.filter(TemplateDoc.branch_id == branch_id).order_by(TemplateDoc.uploaded_at.desc()).first()
        if rec:
            return rec.filename
    rec = base_q.filter(TemplateDoc.branch_id == None).order_by(TemplateDoc.uploaded_at.desc()).first()
    return rec.filename if rec else None

# ---------------- Ø¯ÙˆØ§Ù„ Ù…Ø³Ø§Ø¹Ø¯Ø© ----------------
def save_price(state, region, bank, price):
    record = ValuationMemory.query.filter_by(
        state=state, region=region, bank_id=bank
    ).first()
    if record:
        record.price_per_meter = price
        record.updated_at = datetime.utcnow()
    else:
        record = ValuationMemory(state=state, region=region, bank_id=bank, price_per_meter=price)
        db.session.add(record)
    db.session.commit()




def send_notification(user_id, title, body):
    subs = NotificationSubscription.query.filter_by(user_id=user_id).all()
    vapid_private = app.config.get("VAPID_PRIVATE_KEY") or os.environ.get("VAPID_PRIVATE_KEY")
    vapid_claims = app.config.get("VAPID_CLAIMS", {"sub": "mailto:your-email@example.com"})
    if not vapid_private:
        # Ù„Ø§ ÙŠÙˆØ¬Ø¯ Ù…ÙØªØ§Ø­ Ø®Ø§Øµ Ù„Ù„Ø¥Ø±Ø³Ø§Ù„ØŒ Ù†ØªØ¬Ø§ÙˆØ² Ø­ØªÙ‰ Ù„Ø§ Ù†ÙØ´Ù„ Ø§Ù„ØªØ·Ø¨ÙŠÙ‚
        return
    for sub in subs:
        try:
            webpush(
                subscription_info=json.loads(sub.subscription_json),
                data=json.dumps({"title": title, "body": body}),
                vapid_private_key=vapid_private,
                vapid_claims=vapid_claims
            )
        except WebPushException as e:
            print("âŒ Ø¥Ø´Ø¹Ø§Ø± ÙØ´Ù„:", e)


@app.route('/notify_me')
def notify_me():
    if not session.get("user_id"):
        return {"error": "Unauthorized"}, 401
    try:
        send_notification(session["user_id"], "ğŸ”” Ø§Ø®ØªØ¨Ø§Ø± Ø§Ù„Ø¥Ø´Ø¹Ø§Ø±Ø§Øª", "Ù‡Ø°Ù‡ Ø±Ø³Ø§Ù„Ø© ØªØ¬Ø±ÙŠØ¨ÙŠØ©")
        return {"success": True}
    except Exception as e:
        return {"success": False, "error": str(e)}, 500





def get_last_price(state, region, bank):
    record = ValuationMemory.query.filter_by(
        state=state, region=region, bank_id=bank
    ).order_by(ValuationMemory.updated_at.desc()).first()
    return record.price_per_meter if record else None


# ÙØ­Øµ ÙˆØ¬ÙˆØ¯ Ø¹Ù…ÙˆØ¯ Ø¯Ø§Ø®Ù„ Ø¬Ø¯ÙˆÙ„ (Ù„Ù…Ø´Ø§ÙƒÙ„ Ø§Ù„Ø¥ØµØ¯Ø§Ø±Ø§Øª Ø§Ù„Ù‚Ø¯ÙŠÙ…Ø©)
def column_exists(table_name: str, column_name: str) -> bool:
    try:
        inspector = inspect(db.engine)
        columns = inspector.get_columns(table_name)
        column_names = {col.get("name") for col in columns}
        return column_name in column_names
    except Exception:
        return False

# ---------------- ÙÙÙ„ØªØ± Ø¬ÙŠÙ†Ø¬Ø§: "ÙƒÙ… Ù…Ø¶Ù‰" Ø¨Ø§Ù„Ø¹Ø±Ø¨ÙŠØ© ----------------
@app.template_filter('ago')
def naturaltime_ar(dt):
    if not dt:
        return "Ù„Ù… ÙŠØªÙ… Ø§Ù„Ø¥Ø±Ø³Ø§Ù„"
    delta = datetime.utcnow() - dt
    seconds = int(delta.total_seconds())
    if seconds < 60:
        return "Ù…Ù†Ø° Ø«ÙˆØ§Ù†Ù"
    minutes = seconds // 60
    if minutes < 60:
        if minutes == 1:
            return "Ù…Ù†Ø° Ø¯Ù‚ÙŠÙ‚Ø©"
        elif minutes == 2:
            return "Ù…Ù†Ø° Ø¯Ù‚ÙŠÙ‚ØªÙŠÙ†"
        elif 3 <= minutes <= 10:
            return f"Ù…Ù†Ø° {minutes} Ø¯Ù‚Ø§Ø¦Ù‚"
        else:
            return f"Ù…Ù†Ø° {minutes} Ø¯Ù‚ÙŠÙ‚Ø©"
    hours = minutes // 60
    if hours < 24:
        if hours == 1:
            return "Ù…Ù†Ø° Ø³Ø§Ø¹Ø©"
        elif hours == 2:
            return "Ù…Ù†Ø° Ø³Ø§Ø¹ØªÙŠÙ†"
        elif 3 <= hours <= 10:
            return f"Ù…Ù†Ø° {hours} Ø³Ø§Ø¹Ø§Øª"
        else:
            return f"Ù…Ù†Ø° {hours} Ø³Ø§Ø¹Ø©"
    days = hours // 24
    if days < 30:
        if days == 1:
            return "Ù…Ù†Ø° ÙŠÙˆÙ…"
        elif days == 2:
            return "Ù…Ù†Ø° ÙŠÙˆÙ…ÙŠÙ†"
        else:
            return f"Ù…Ù†Ø° {days} Ø£ÙŠØ§Ù…"
    months = days // 30
    if months < 12:
        if months == 1:
            return "Ù…Ù†Ø° Ø´Ù‡Ø±"
        elif months == 2:
            return "Ù…Ù†Ø° Ø´Ù‡Ø±ÙŠÙ†"
        else:
            return f"Ù…Ù†Ø° {months} Ø£Ø´Ù‡Ø±"
    years = months // 12
    if years == 1:
        return "Ù…Ù†Ø° Ø³Ù†Ø©"
    elif years == 2:
        return "Ù…Ù†Ø° Ø³Ù†ØªÙŠÙ†"
    else:
        return f"Ù…Ù†Ø° {years} Ø³Ù†ÙˆØ§Øª"

# ---------------- Ø­Ø§Ù„Ø© Ù…Ø³ØªÙ†Ø¯ Ø­Ø³Ø¨ ØªØ§Ø±ÙŠØ® Ø§Ù„Ø§Ù†ØªÙ‡Ø§Ø¡ ----------------
def document_status(doc):
    try:
        exp = getattr(doc, "expires_at", None)
        if not exp:
            return "Ø¨Ø¯ÙˆÙ† Ø§Ù†ØªÙ‡Ø§Ø¡"
        delta_days = (exp - datetime.utcnow()).days
        if delta_days < 0:
            return "Ù…Ù†ØªÙ‡ÙŠ"
        if delta_days <= 30:
            return "Ù‚Ø±ÙŠØ¨ Ø§Ù„Ø§Ù†ØªÙ‡Ø§Ø¡"
        return "Ø³Ø§Ø±ÙŠ"
    except Exception:
        return "Ø¨Ø¯ÙˆÙ† Ø§Ù†ØªÙ‡Ø§Ø¡"

# ---------------- Ø§Ù„Ù…Ø³Ø§Ø±Ø§Øª ----------------
@app.route("/")
def index():
    if "user_id" not in session:
        return redirect(url_for("login"))

    role = session.get("role")
    if role == "manager":
        return redirect(url_for("manager_dashboard"))
    elif role == "employee":
        # ØªÙˆØ¬ÙŠÙ‡ Ø§Ù„Ù…ÙˆØ¸Ù Ø­Ø³Ø¨ Ø§Ù„Ù‚Ø³Ù… Ø§Ù„Ù…Ø¹ÙŠÙ‘Ù† Ù„Ù‡ Ø¥Ù† ÙˆØ¬Ø¯ØŒ ÙˆØ¥Ù„Ø§ Ø­Ø³Ø¨ Ù‚Ø³Ù… Ø§Ù„ÙØ±Ø¹ØŒ ÙˆØ¥Ù„Ø§ Ù„ÙˆØ­Ø© Ø§Ù„Ù…ÙˆØ¸Ù Ø§Ù„Ø§ÙØªØ±Ø§Ø¶ÙŠØ©
        try:
            ensure_branch_department_column()
            user = User.query.get(session.get("user_id"))
            # Ø£ÙˆÙ„ÙˆÙŠØ©: Ù‚Ø³Ù… Ø§Ù„Ù…ÙˆØ¸Ù
            if user and getattr(user, "section_id", None):
                sec = BranchSection.query.get(user.section_id)
                if sec:
                    return _redirect_to_section(sec.name)
            # fallback: Ù‚Ø³Ù… Ø§Ù„ÙØ±Ø¹ Ø§Ù„Ù‚Ø¯ÙŠÙ…
            if user and user.branch_id:
                b = Branch.query.get(user.branch_id)
                dept = (b.department or "").lower() if b else ""
                if dept in ("consultations", "consultation", "consulting", "Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª"):
                    return redirect(url_for("consultations_list"))
                if dept in ("finance", "financial", "Ø§Ù„Ù…Ø§Ù„ÙŠØ©"):
                    return redirect(url_for("finance_dashboard"))
        except Exception:
            pass
        return redirect(url_for("employee_dashboard"))
    elif role == "engineer":
        return redirect(url_for("engineer_dashboard"))
    elif role == "finance":
        return redirect(url_for("finance_dashboard"))
    return redirect(url_for("login"))

# ================= Consulting Department Routes =================
def _parse_date(value: str | None) -> date | None:
    try:
        if not value:
            return None
        return datetime.strptime(value, "%Y-%m-%d").date()
    except Exception:
        return None


@app.route("/consultations")
def consultations_list():
    if session.get("role") not in ["manager", "finance", "consultant"]:
        return redirect(url_for("login"))

    # Filters
    q = (request.args.get("q") or "").strip()
    status = (request.args.get("status") or "").strip()
    ctype = (request.args.get("type") or "").strip()
    project = (request.args.get("project") or "").strip()
    client = (request.args.get("client") or "").strip()
    page = int(request.args.get("page") or 1)
    per_page = min(int(request.args.get("per_page") or 20), 100)

    query = Consultation.query
    if status:
        query = query.filter(Consultation.status == status)
    if ctype:
        query = query.filter(Consultation.consultation_type == ctype)
    if project:
        if project.isdigit():
            query = (
                query.join(ConsultingProject, isouter=True)
                .filter(Consultation.project_id == int(project))
            )
        else:
            query = (
                query.join(ConsultingProject, isouter=True)
                .filter(ConsultingProject.name.ilike(f"%{project}%"))
            )
    if client:
        if client.isdigit():
            query = (
                query.join(Client, isouter=True)
                .filter(Consultation.client_id == int(client))
            )
        else:
            query = (
                query.join(Client, isouter=True)
                .filter(Client.name.ilike(f"%{client}%"))
            )
    if q:
        query = query.filter(or_(
            Consultation.consultant_name.ilike(f"%{q}%"),
            Consultation.description.ilike(f"%{q}%"),
        ))

    total = query.count()
    consultations = (
        query.order_by(Consultation.created_at.desc())
        .offset((page - 1) * per_page)
        .limit(per_page)
        .all()
    )

    return render_template(
        "consultations.html",
        consultations=consultations,
        total=total,
        page=page,
        per_page=per_page,
        q=q,
        status=status,
        ctype=ctype,
        project=project,
        client=client,
        statuses=CONSULTATION_STATUSES,
        types=CONSULTATION_TYPES,
    )


def _find_or_create_customer(name: str | None, phone: str | None) -> Client | None:
    name = (name or "").strip()
    phone = (phone or "").strip()
    if not name and not phone:
        return None
    existing = None
    try:
        if phone:
            existing = Client.query.filter_by(phone=phone).first()
    except Exception:
        existing = None
    if existing:
        if name and existing.name != name:
            existing.name = name
            try:
                db.session.commit()
            except Exception:
                db.session.rollback()
        return existing
    # Ø§Ù„Ù†ÙˆØ¹ Ù…Ø·Ù„ÙˆØ¨ ÙÙŠ Ù†Ù…ÙˆØ°Ø¬ Ø¹Ù…Ù„Ø§Ø¡ Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª
    c = Client(name=name or "-", phone=phone or "-", type="ÙØ±Ø¯")
    db.session.add(c)
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
        try:
            db.session.flush()
        except Exception:
            pass
    return c


def _find_or_create_project(project_name: str | None, client_obj: Client | None) -> ConsultingProject | None:
    name = (project_name or "").strip()
    if not name or not client_obj:
        # Ù…Ø´Ø±ÙˆØ¹ Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª ÙŠØªØ·Ù„Ø¨ Ø¹Ù…ÙŠÙ„Ù‹Ø§
        return None
    p = (
        ConsultingProject.query
        .filter(func.lower(ConsultingProject.name) == func.lower(name))
        .first()
    )
    if p:
        return p
    # Ø§Ù„Ù†ÙˆØ¹ Ù…Ø·Ù„ÙˆØ¨ ÙÙŠ Ù†Ù…ÙˆØ°Ø¬ Ù…Ø´Ø§Ø±ÙŠØ¹ Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª
    p = ConsultingProject(
        name=name,
        client_id=client_obj.id,
        type=(PROJECT_TYPES[0] if PROJECT_TYPES else "ØªØµÙ…ÙŠÙ… Ù…Ø¹Ù…Ø§Ø±ÙŠ"),
    )
    db.session.add(p)
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
        try:
            db.session.flush()
        except Exception:
            pass
    return p


@app.route("/consultations/new", methods=["GET", "POST"])
def consultations_new():
    role = session.get("role")
    if role not in ["manager", "finance", "consultant"]:
        return redirect(url_for("login"))

    if request.method == "POST":
        consultation_type = (request.form.get("consultation_type") or "").strip() or "Architectural"
        status = (request.form.get("status") or "").strip() or "Pending"
        description = (request.form.get("description") or "").strip()
        consultant_name = (request.form.get("consultant_name") or "").strip()
        client_name = (request.form.get("client_name") or "").strip()
        client_phone = (request.form.get("client_phone") or "").strip()
        project_name = (request.form.get("project_name") or "").strip()
        start_date = _parse_date(request.form.get("start_date"))
        end_date = _parse_date(request.form.get("end_date"))
        # Cost can only be set by manager/finance
        cost = None
        if role in ["manager", "finance"]:
            try:
                cost = float(request.form.get("cost") or 0)
            except Exception:
                cost = 0.0

        if consultation_type not in CONSULTATION_TYPES:
            consultation_type = "Other"
        if status not in CONSULTATION_STATUSES:
            status = "Pending"

        # link client and project
        client_obj = _find_or_create_customer(client_name, client_phone)
        project_obj = _find_or_create_project(project_name, client_obj)

        c = Consultation(
            consultation_type=consultation_type,
            status=status,
            description=description,
            consultant_name=consultant_name or (session.get("username") if role == "consultant" else None),
            client_id=(client_obj.id if client_obj else None),
            project_id=(project_obj.id if project_obj else None),
            start_date=start_date,
            end_date=end_date,
            cost=cost,
            created_by=session.get("user_id"),
            consultant_id=(session.get("user_id") if role == "consultant" else None),
        )
        db.session.add(c)
        db.session.commit()

        # Notify manager
        try:
            mgr = User.query.filter_by(role="manager").first()
            if mgr:
                send_notification(mgr.id, "New Consultation", f"New consultation #{c.id} added")
        except Exception:
            pass

        flash("âœ… Consultation created", "success")
        return redirect(url_for("consultations_detail", cid=c.id))

    return render_template(
        "consultation_form.html",
        consultation=None,
        statuses=CONSULTATION_STATUSES,
        types=CONSULTATION_TYPES,
        mode="new",
    )


@app.route("/consultations/<int:cid>")
def consultations_detail(cid: int):
    if session.get("role") not in ["manager", "finance", "consultant"]:
        return redirect(url_for("login"))
    c = Consultation.query.get_or_404(cid)
    # prefetch project and client (from consulting models)
    project = c.project if c.project_id else None
    client = c.client if c.client_id else None
    return render_template(
        "consultation_detail.html",
        c=c,
        project=project,
        client=client,
        statuses=CONSULTATION_STATUSES,
        types=CONSULTATION_TYPES,
    )


@app.route("/consultations/<int:cid>/edit", methods=["GET", "POST"])
def consultations_edit(cid: int):
    role = session.get("role")
    if role not in ["manager", "finance", "consultant"]:
        return redirect(url_for("login"))
    c = Consultation.query.get_or_404(cid)

    if request.method == "POST":
        prev_status = c.status
        # role-based updates
        if role == "manager":
            # full control
            c.consultant_name = (request.form.get("consultant_name") or c.consultant_name)
            c.consultation_type = request.form.get("consultation_type") or c.consultation_type
            c.description = request.form.get("description") or c.description
            new_status = request.form.get("status") or c.status
            c.status = new_status if new_status in CONSULTATION_STATUSES else c.status
            c.start_date = _parse_date(request.form.get("start_date")) or c.start_date
            c.end_date = _parse_date(request.form.get("end_date")) or c.end_date
            try:
                c.cost = float(request.form.get("cost") or c.cost or 0)
            except Exception:
                pass
            # client/project
            client_name = (request.form.get("client_name") or "").strip()
            client_phone = (request.form.get("client_phone") or "").strip()
            project_name = (request.form.get("project_name") or "").strip()
            cl = _find_or_create_customer(client_name, client_phone) if (client_name or client_phone) else None
            if cl:
                c.client_id = cl.id
            pr = _find_or_create_project(project_name, cl or (Client.query.get(c.client_id) if c.client_id else None)) if project_name else None
            if pr:
                c.project_id = pr.id
        elif role == "finance":
            # can edit cost/status
            new_status = request.form.get("status") or c.status
            c.status = new_status if new_status in CONSULTATION_STATUSES else c.status
            try:
                c.cost = float(request.form.get("cost") or c.cost or 0)
            except Exception:
                pass
        elif role == "consultant":
            # can edit description and status
            c.description = request.form.get("description") or c.description
            new_status = request.form.get("status") or c.status
            c.status = new_status if new_status in CONSULTATION_STATUSES else c.status
        else:
            return redirect(url_for("login"))

        db.session.commit()

        try:
            if prev_status != c.status:
                mgr = User.query.filter_by(role="manager").first()
                if mgr:
                    send_notification(mgr.id, "Consultation status updated", f"Consultation #{c.id} â†’ {c.status}")
        except Exception:
            pass

        flash("âœ… Consultation updated", "success")
        return redirect(url_for("consultations_detail", cid=c.id))

    project = c.project if c.project_id else None
    client = c.client if c.client_id else None
    return render_template(
        "consultation_form.html",
        consultation=c,
        project=project,
        client=client,
        statuses=CONSULTATION_STATUSES,
        types=CONSULTATION_TYPES,
        mode="edit",
    )


@app.route("/consultations/<int:cid>/delete", methods=["POST"])
def consultations_delete(cid: int):
    if session.get("role") != "manager":
        return redirect(url_for("login"))
    c = Consultation.query.get_or_404(cid)
    db.session.delete(c)
    db.session.commit()
    flash("ğŸ—‘ï¸ Consultation deleted", "success")
    return redirect(url_for("consultations_list"))


@app.route("/consultations/export.csv")
def consultations_export_csv():
    if session.get("role") not in ["manager", "finance"]:
        return redirect(url_for("login"))
    # reuse filters
    q = (request.args.get("q") or "").strip()
    status = (request.args.get("status") or "").strip()
    ctype = (request.args.get("type") or "").strip()
    project = (request.args.get("project") or "").strip()
    client = (request.args.get("client") or "").strip()

    query = Consultation.query
    if status:
        query = query.filter(Consultation.status == status)
    if ctype:
        query = query.filter(Consultation.consultation_type == ctype)
    if project:
        if project.isdigit():
            query = (
                query.join(ConsultingProject, isouter=True)
                .filter(Consultation.project_id == int(project))
            )
        else:
            query = (
                query.join(ConsultingProject, isouter=True)
                .filter(ConsultingProject.name.ilike(f"%{project}%"))
            )
    if client:
        if client.isdigit():
            query = (
                query.join(Client, isouter=True)
                .filter(Consultation.client_id == int(client))
            )
        else:
            query = (
                query.join(Client, isouter=True)
                .filter(Client.name.ilike(f"%{client}%"))
            )
    if q:
        query = query.filter(or_(
            Consultation.consultant_name.ilike(f"%{q}%"),
            Consultation.description.ilike(f"%{q}%"),
        ))

    rows = query.order_by(Consultation.created_at.desc()).all()
    import csv
    from io import StringIO

    buf = StringIO()
    writer = csv.writer(buf)
    writer.writerow([
        "ID", "Project", "Client", "Type", "Status", "Start", "End", "Cost", "Consultant", "Created At",
    ])
    for r in rows:
        pr = ConsultingProject.query.get(r.project_id) if r.project_id else None
        cl = Client.query.get(r.client_id) if r.client_id else None
        writer.writerow([
            r.id,
            (pr.name if pr else ""),
            (cl.name if cl else ""),
            r.consultation_type or "",
            r.status or "",
            (r.start_date.isoformat() if r.start_date else ""),
            (r.end_date.isoformat() if r.end_date else ""),
            (f"{float(r.cost or 0):.2f}"),
            r.consultant_name or "",
            (r.created_at.strftime("%Y-%m-%d %H:%M") if r.created_at else ""),
        ])
    buf.seek(0)
    return Response(
        buf.getvalue(),
        mimetype="text/csv; charset=utf-8",
        headers={
            "Content-Disposition": "attachment; filename=consultations.csv"
        },
    )


@app.route("/consultations/<int:cid>/invoice", methods=["POST"])
def consultations_invoice(cid: int):
    if session.get("role") not in ["finance", "manager"]:
        return redirect(url_for("login"))
    c = Consultation.query.get_or_404(cid)
    amount = float(c.cost or 0)
    if amount <= 0:
        flash("âš ï¸ Consultation has no cost to invoice", "warning")
        return redirect(url_for("consultations_detail", cid=c.id))
    client = Client.query.get(c.client_id) if c.client_id else None
    inv = CustomerInvoice(
        customer_name=(client.name if client else (c.consultant_name or "Consultation")),
        amount=amount,
        transaction_id=None,
        note=f"Consultation #{c.id} - {c.consultation_type}",
        created_by=session.get("user_id"),
    )
    db.session.add(inv)
    db.session.commit()
    flash("âœ… Invoice created", "success")
    return redirect(url_for("print_customer_invoice_html", invoice_id=inv.id))

# ---------------- ØªØ³Ø¬ÙŠÙ„ Ø§Ù„Ø¯Ø®ÙˆÙ„ ----------------
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = (request.form.get("username") or "").strip()
        password = (request.form.get("password") or "").strip()
        # Ù…Ø­Ø§ÙˆÙ„Ø§Øª Ø¨Ø³ÙŠØ·Ø© Ù„Ø¥Ø¹Ø§Ø¯Ø© Ø§Ù„Ù…Ø­Ø§ÙˆÙ„Ø© ÙÙŠ Ø­Ø§Ù„ Ø§Ù†Ù‚Ø·Ø§Ø¹ Ø§ØªØµØ§Ù„ Ù‚Ø§Ø¹Ø¯Ø© Ø§Ù„Ø¨ÙŠØ§Ù†Ø§Øª Ø¨Ø´ÙƒÙ„ Ø¹Ø§Ø¨Ø±
        user = None
        last_error = None
        for attempt_index, sleep_seconds in enumerate([0.0, 0.3, 0.8], start=1):
            try:
                if sleep_seconds:
                    time.sleep(sleep_seconds)
                user = User.query.filter_by(username=username).first()
                last_error = None
                break
            except OperationalError as op_err:
                # Ø¥Ù†Ù‡Ø§Ø¡ Ø£ÙŠ Ù…Ø¹Ø§Ù…Ù„Ø© Ù…Ø¹Ù„Ù‘Ù‚Ø© ÙˆØ¥Ø¹Ø§Ø¯Ø© Ø§Ù„Ù…Ø­Ø§ÙˆÙ„Ø©
                try:
                    db.session.rollback()
                except Exception:
                    pass
                last_error = op_err
            except Exception as unhandled_err:
                last_error = unhandled_err
                break
        if last_error is not None and user is None:
            flash("âš ï¸ ØªØ¹Ø°Ù‘Ø± Ø§Ù„Ø§ØªØµØ§Ù„ Ø¨Ù‚Ø§Ø¹Ø¯Ø© Ø§Ù„Ø¨ÙŠØ§Ù†Ø§Øª Ù…Ø¤Ù‚ØªÙ‹Ø§. Ø§Ù„Ø±Ø¬Ø§Ø¡ Ø§Ù„Ù…Ø­Ø§ÙˆÙ„Ø© Ø¨Ø¹Ø¯ Ù„Ø­Ø¸Ø§Øª.", "warning")
            return render_template("login.html"), 503
        # Ø¯Ø¹Ù… Ø§Ù„ØªØ­Ù‚Ù‚ Ø§Ù„Ø¹Ø§Ù… Ù„Ø¬Ù…ÙŠØ¹ Ø§Ù„Ø®ÙˆØ§Ø±Ø²Ù…ÙŠØ§Øª (pbkdf2, scrypt, ...)
        # Ù…Ø¹ Ù…Ø³Ø§Ø± ØªÙˆØ§ÙÙ‚ Ù„Ù„Ù†Øµ Ø§Ù„ØµØ±ÙŠØ­ Ø§Ù„Ù‚Ø¯ÙŠÙ… Ø«Ù… Ø§Ù„ØªØ±Ù‚ÙŠØ© Ù„Ù„ØªØ¬Ø²Ø¦Ø©
        is_valid = False
        if user and user.password:
            # Ø£ÙˆÙ„Ù‹Ø§ Ø­Ø§ÙˆÙ„ Ø§Ù„ØªØ­Ù‚Ù‚ ÙƒÙƒÙ„Ù…Ø© Ù…Ø±ÙˆØ± Ù…ÙØ¬Ø²Ù‘Ø£Ø© (Werkzeug ÙŠØ¯Ø¹Ù… Ø¹Ø¯Ø© Ø®ÙˆØ§Ø±Ø²Ù…ÙŠØ§Øª ØªÙ„Ù‚Ø§Ø¦ÙŠÙ‹Ø§)
            try:
                is_valid = check_password_hash(user.password, password)
            except Exception:
                is_valid = False

            # Ù…Ø³Ø§Ø± Ø®Ù„ÙÙŠ: Ø¥Ù† ÙƒØ§Ù†Øª Ù…Ø®Ø²Ù‘Ù†Ø© ÙƒÙ†Øµ ØµØ±ÙŠØ­ Ù‚Ø¯ÙŠÙ…Ù‹Ø§
            if not is_valid and user.password == password:
                is_valid = True
                try:
                    user.password = generate_password_hash(password)
                    db.session.commit()
                except Exception:
                    db.session.rollback()

        if user and is_valid:
            session["user_id"] = user.id
            session["role"] = user.role
            session["username"] = user.username  # Ù†Ø­ØªØ§Ø¬Ù‡ Ù„Ù„ØªÙ‚Ø§Ø±ÙŠØ± ÙˆØ§Ù„Ø§Ø³ØªÙ„Ø§Ù…
            return redirect(url_for("index"))
        else:
            flash("âŒ Ø§Ø³Ù… Ø§Ù„Ù…Ø³ØªØ®Ø¯Ù… Ø£Ùˆ ÙƒÙ„Ù…Ø© Ø§Ù„Ù…Ø±ÙˆØ± ØºÙŠØ± ØµØ­ÙŠØ­Ø©", "danger")

    return render_template("login.html")

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

# ---------------- Ù„ÙˆØ­Ø© Ø§Ù„Ù…ÙˆØ¸Ù ----------------
VAPID_PUBLIC_KEY = "BFNeZpjEro8pwFxR1H20twlTd2pL5MZtWrDATu4ME2RcbzhN"  # Ø§Ù„Ù…ÙØªØ§Ø­ Ø§Ù„Ù„ÙŠ ÙˆÙ„Ø¯ØªÙ‡

@app.route("/employee")
def employee_dashboard():
    if session.get("role") != "employee":
        return redirect(url_for("login"))

    transactions = Transaction.query.filter_by(assigned_to=session.get("user_id")).all()
    # ğŸ§® Ø¥Ø­ØµØ§Ø¦ÙŠØ§Øª Ø­Ø³Ø¨ Ù…Ù† Ø¬Ù„Ø¨ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© (Ù‡Ø°Ø§ Ø§Ù„Ù…ÙˆØ¸Ù)
    current_user = User.query.get(session.get("user_id"))
    brought_name = current_user.username if current_user else None
    # ÙÙ„ØªØ±Ø© ÙØªØ±Ø© Ø²Ù…Ù†ÙŠØ© Ø§Ø®ØªÙŠØ§Ø±ÙŠØ© Ù…Ù† ÙˆØ§Ø¬Ù‡Ø© Ø§Ù„Ù…Ø³ØªØ®Ø¯Ù…
    start_date_str = request.args.get("start_date")
    end_date_str = request.args.get("end_date")
    start_date = None
    end_date = None
    try:
        if start_date_str:
            start_date = datetime.strptime(start_date_str, "%Y-%m-%d")
        if end_date_str:
            # Ø§Ø¬Ø¹Ù„ Ù†Ù‡Ø§ÙŠØ© Ø§Ù„ÙŠÙˆÙ… Ø´Ø§Ù…Ù„Ø©
            end_date = datetime.strptime(end_date_str, "%Y-%m-%d") + timedelta(days=1)
    except Exception:
        start_date = None
        end_date = None

    def base_brought_query(ttype: str):
        q = Transaction.query.filter(Transaction.brought_by == brought_name, Transaction.transaction_type == ttype)
        if start_date:
            q = q.filter(Transaction.date >= start_date)
        if end_date:
            q = q.filter(Transaction.date < end_date)
        return q

    real_estate_brought_count = 0
    vehicle_brought_count = 0
    if brought_name:
        real_estate_brought_count = base_brought_query("real_estate").count()
        vehicle_brought_count = base_brought_query("vehicle").count()
    banks = Bank.query.all()

    # Ù…Ø³ØªÙ†Ø¯Ø§Øª Ø§Ù„ÙØ±Ø¹ Ø§Ù„Ø®Ø§ØµØ© Ø¨Ø§Ù„Ù…ÙˆØ¸Ù
    user = User.query.get(session.get("user_id"))
    branch_docs = []
    if user and getattr(user, "branch_id", None):
        try:
            branch_docs = BranchDocument.query.filter_by(branch_id=user.branch_id)\
                .order_by(BranchDocument.expires_at.asc().nulls_last()).all()
        except Exception:
            branch_docs = BranchDocument.query.filter_by(branch_id=user.branch_id).all()

    # ØªÙ…Ø±ÙŠØ± Ø§Ù„Ø³Ø¹Ø± Ø§Ù„Ø§ÙØªØ±Ø§Ø¶ÙŠ (Ù„Ùˆ Ù…Ø§ ÙÙŠÙ‡ Ø°Ø§ÙƒØ±Ø© Ù†Ø®Ù„ÙŠÙ‡ ØµÙØ±)
    price_per_meter = 0.0  

    return render_template(
        "employee.html",
        transactions=transactions,
        banks=banks,
        vapid_public_key=VAPID_PUBLIC_KEY,
        price_per_meter=price_per_meter,
        docs=branch_docs,
        status_for=document_status,
        start_date=start_date_str,
        end_date=end_date_str,
        real_estate_brought_count=real_estate_brought_count,
        vehicle_brought_count=vehicle_brought_count
    )

@app.route("/add_transaction", methods=["POST"])
def add_transaction():
    if session.get("role") != "employee":
        return redirect(url_for("login"))
    
    user = User.query.get(session["user_id"])
    transaction_type = request.form.get("transaction_type")  # âœ… Ù†Ø­Ø¯Ø¯ Ù†ÙˆØ¹ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø©
    client_name = (request.form.get("client_name") or "").strip()
    client_phone = (request.form.get("client_phone") or "").strip()
    fee = float(request.form.get("fee") or 0)
    # ğŸ†• Ø§Ù„Ø­Ù‚ÙˆÙ„ Ø§Ù„Ø¬Ø¯ÙŠØ¯Ø©: Ù…Ù† Ø¬Ù„Ø¨ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© ÙˆÙ…Ù† Ù‚Ø§Ù… Ø¨Ø§Ù„Ø²ÙŠØ§Ø±Ø© (Ø§Ø®ØªÙŠØ§Ø±ÙŠ)
    brought_by_form = (request.form.get("brought_by") or "").strip()
    visited_by_form = (request.form.get("visited_by") or "").strip()

    # âœ… ØªØ­Ù‚Ù‚ Ø¥Ø¬Ø¨Ø§Ø±ÙŠ Ù„Ù„Ø­Ù‚Ù„ÙŠÙ†
    if not brought_by_form or not visited_by_form:
        flash("âš ï¸ ÙŠØ¬Ø¨ Ø¥Ø¯Ø®Ø§Ù„ Ù…Ù† Ø¬Ù„Ø¨ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© ÙˆÙ…Ù† Ù‚Ø§Ù… Ø¨Ø§Ù„Ø²ÙŠØ§Ø±Ø©", "danger")
        return redirect(url_for("employee_dashboard"))

    t = None  # Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø©

    # âœ… ØªØ­Ù‚Ù‚ Ù…Ù† Ø±Ù‚Ù… Ø§Ù„Ø¹Ù…ÙŠÙ„
    if not client_phone:
        flash("âš ï¸ ÙŠØ¬Ø¨ Ø¥Ø¯Ø®Ø§Ù„ Ø±Ù‚Ù… Ø§Ù„Ø¹Ù…ÙŠÙ„", "danger")
        return redirect(url_for("employee_dashboard"))

    # ğŸ§¾ Ø­ÙØ¸/ØªØ­Ø¯ÙŠØ« Ø§Ù„Ø¹Ù…ÙŠÙ„ ÙÙŠ Ø¬Ø¯ÙˆÙ„ Ø§Ù„Ø¹Ù…Ù„Ø§Ø¡
    existing_customer = Customer.query.filter_by(phone=client_phone).first()
    if existing_customer:
        # Ù†Ø­Ø¯Ù‘Ø« Ø§Ù„Ø§Ø³Ù… Ø¥Ø°Ø§ ÙƒØ§Ù† Ù…Ø®ØªÙ„ÙÙ‹Ø§
        if client_name and existing_customer.name != client_name:
            existing_customer.name = client_name
    else:
        db.session.add(Customer(name=client_name or "-", phone=client_phone))
        db.session.flush()

    # ğŸ  Ù…Ø¹Ø§Ù…Ù„Ø© Ø¹Ù‚Ø§Ø±
    if transaction_type == "real_estate":
        state = request.form.get("state")
        region = request.form.get("region")
        bank_id = request.form.get("bank_id")
        bank_branch = (request.form.get("bank_branch") or "").strip()
        bank_employee_name = (request.form.get("bank_employee_name") or "").strip()
        try:
            bank_id = int(bank_id) if bank_id else None
        except Exception:
            bank_id = None

        # âœ… Ø§Ù„Ø¨Ø­Ø« ÙÙŠ ValuationMemory
        vm = None
        if state and region and bank_id:
            vm = ValuationMemory.query.filter_by(
                state=state, region=region, bank_id=bank_id
            ).order_by(ValuationMemory.updated_at.desc()).first()

        if vm:
            price_per_meter = vm.price_per_meter
        else:
            lp = LandPrice.query.filter_by(state=state, region=region, bank_id=bank_id).first()
            price_per_meter = lp.price_per_meter if lp else 0.0

        # ğŸ“ Ø§Ù„Ø¨ÙŠØ§Ù†Ø§Øª
        area          = float(request.form.get("area") or 0)
        building_area = float(request.form.get("building_area") or 0)
        building_age  = int(request.form.get("building_age") or 0)

        # âœ… Ø­Ø³Ø§Ø¨ Ø§Ù„ØªØ«Ù…ÙŠÙ†
        land_value = area * price_per_meter if price_per_meter else 0.0
        building_value = 0
        if building_area > 0 and building_age > 0:
            building_value = building_area * (185 / 50) * building_age

        total_estimate = land_value + building_value

        # ØªØ­Ù‚Ù‚ Ø£Ø³Ø§Ø³ÙŠ: Ø§Ù„Ø¨Ù†Ùƒ ÙˆÙØ±Ø¹ Ø§Ù„Ø¨Ù†Ùƒ Ù…Ø·Ù„ÙˆØ¨Ø§Ù†
        if not bank_id or not bank_branch:
            flash("âš ï¸ ÙŠØ±Ø¬Ù‰ Ø§Ø®ØªÙŠØ§Ø± Ø§Ù„Ø¨Ù†Ùƒ ÙˆÙƒØªØ§Ø¨Ø© ÙØ±Ø¹ Ø§Ù„Ø¨Ù†Ùƒ", "danger")
            return redirect(url_for("employee_dashboard"))

        t = Transaction(
            client=client_name,
            employee=user.username,
            date=datetime.utcnow(),
            status="Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³",
            fee=fee,
            branch_id=user.branch_id,
            land_value=land_value,
            building_value=building_value,
            total_estimate=total_estimate,
            valuation_amount=total_estimate,  # ğŸ‘ˆ Ù†Ø®Ø²Ù† Ø§Ù„ØªØ«Ù…ÙŠÙ† Ù‡Ù†Ø§
            area=area,
            building_area=building_area,
            building_age=building_age,
            state=state,
            region=region,
            bank_id=bank_id,
            bank_branch=bank_branch,
            bank_employee_name=bank_employee_name,
            brought_by=(brought_by_form or user.username),
            visited_by=(visited_by_form or None),
            created_by=user.id,
            payment_status="ØºÙŠØ± Ù…Ø¯ÙÙˆØ¹Ø©",
            transaction_type="real_estate",
            assigned_to=None
        )

    # ğŸš— Ù…Ø¹Ø§Ù…Ù„Ø© Ù…Ø±ÙƒØ¨Ø©
    elif transaction_type == "vehicle":
        vehicle_type  = request.form.get("vehicle_type")
        vehicle_model = request.form.get("vehicle_model")
        vehicle_year  = request.form.get("vehicle_year")
        vehicle_value = float(request.form.get("vehicle_value") or 0)

        # ØªØ­Ù‚Ù‚ Ø£Ø³Ø§Ø³ÙŠ: Ø§Ù„Ø¨Ù†Ùƒ ÙˆÙØ±Ø¹ Ø§Ù„Ø¨Ù†Ùƒ Ù…Ø·Ù„ÙˆØ¨Ø§Ù† Ù„Ù…Ø¹Ø§Ù…Ù„Ø§Øª Ø§Ù„Ù…Ø±ÙƒØ¨Ø§Øª Ø£ÙŠØ¶Ù‹Ø§
        bank_id = request.form.get("bank_id")
        bank_branch = (request.form.get("bank_branch") or "").strip()
        bank_employee_name = (request.form.get("bank_employee_name") or "").strip()
        if not bank_id or not bank_branch:
            flash("âš ï¸ ÙŠØ±Ø¬Ù‰ Ø§Ø®ØªÙŠØ§Ø± Ø§Ù„Ø¨Ù†Ùƒ ÙˆÙƒØªØ§Ø¨Ø© ÙØ±Ø¹ Ø§Ù„Ø¨Ù†Ùƒ", "danger")
            return redirect(url_for("employee_dashboard"))

        t = Transaction(
            client=client_name,
            employee=user.username,
            date=datetime.utcnow(),
            status="Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³",  # âœ… Ø¨Ø¯ÙˆÙ† Ù‡Ù…Ø²Ø©
            fee=fee,
            branch_id=user.branch_id,
            total_estimate=vehicle_value,
            brought_by=(brought_by_form or user.username),
            visited_by=(visited_by_form or None),
            created_by=user.id,
            payment_status="ØºÙŠØ± Ù…Ø¯ÙÙˆØ¹Ø©",
            transaction_type="vehicle",
            vehicle_type=vehicle_type,
            vehicle_model=vehicle_model,
            vehicle_year=vehicle_year,
            state=None,
            region=None,
            bank_id=bank_id,
            bank_branch=bank_branch,
            bank_employee_name=bank_employee_name,
            assigned_to=None   # âœ…
        )


        # Ø¥Ø¨Ù‚Ø§Ø¡ Ù…Ø¹Ø§Ù…Ù„Ø© Ø§Ù„Ù…Ø±ÙƒØ¨Ø© ØºÙŠØ± Ù…Ø³Ù†Ø¯Ø© Ø­ØªÙ‰ ÙŠÙ‚ÙˆÙ… Ù…Ù‡Ù†Ø¯Ø³ Ø¨Ø§Ù„Ø§Ø³ØªÙ„Ø§Ù… Ù…Ù† Ù„ÙˆØ­Ø© Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³

    # Ø±ÙØ¹ Ø§Ù„Ù…Ù„ÙØ§Øª
    files = request.files.getlist("files")
    saved_files = []
    for file in files:
        if file and file.filename:
            filename = secure_filename(file.filename)
            file.save(os.path.join(app.config["UPLOAD_FOLDER"], filename))
            saved_files.append(filename)
    t.files = ",".join(saved_files)

    db.session.add(t)
    db.session.commit()
    bump_transactions_version()

    # ğŸ”” Ø¥Ø´Ø¹Ø§Ø± Ø¬Ù…ÙŠØ¹ Ù…Ù‡Ù†Ø¯Ø³ÙŠ Ù†ÙØ³ Ø§Ù„ÙØ±Ø¹ Ø¨ÙˆØ¬ÙˆØ¯ Ù…Ø¹Ø§Ù…Ù„Ø© Ø¬Ø¯ÙŠØ¯Ø©
    try:
        engineers = User.query.filter_by(role="engineer", branch_id=user.branch_id).all()
        for eng in engineers:
            send_notification(eng.id, "ğŸ“‹ Ù…Ø¹Ø§Ù…Ù„Ø© Ø¬Ø¯ÙŠØ¯Ø©", f"ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ù…Ø¹Ø§Ù…Ù„Ø© Ø±Ù‚Ù… {t.id}")
        # ğŸ”” Ø¥Ø´Ø¹Ø§Ø± Ù‚Ø³Ù… Ø§Ù„Ù…Ø§Ù„ÙŠØ© Ø¨ÙˆØ¬ÙˆØ¯ Ù…Ø¹Ø§Ù…Ù„Ø© Ø¬Ø¯ÙŠØ¯Ø©
        finances = User.query.filter_by(role="finance").all()
        for fin in finances:
            send_notification(fin.id, "ğŸ“‹ Ù…Ø¹Ø§Ù…Ù„Ø© Ø¬Ø¯ÙŠØ¯Ø©", f"ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ù…Ø¹Ø§Ù…Ù„Ø© Ø±Ù‚Ù… {t.id}")
    except Exception:
        pass
    flash("âœ… ØªÙ… Ø¥Ø¶Ø§ÙØ© Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ø¨Ù†Ø¬Ø§Ø­", "success")
    return redirect(url_for("employee_dashboard"))


# ğŸ¢ Ø¥Ø¯Ø§Ø±Ø© Ø§Ù„ÙØ±ÙˆØ¹
@app.route("/manage_branches", methods=["GET", "POST"])
def manage_branches():
    if session.get("role") != "manager":
        return redirect(url_for("login"))

    # ØªØ£ÙƒØ¯ Ù…Ù† ÙˆØ¬ÙˆØ¯ Ø¹Ù…ÙˆØ¯ department + ØªØ±Ø­ÙŠÙ„ Ø§Ù„Ø£Ù‚Ø³Ø§Ù… Ø¥Ù„Ù‰ Ø§Ù„Ø¬Ø¯ÙˆÙ„ Ø§Ù„Ø¬Ø¯ÙŠØ¯
    ensure_branch_department_column()
    ensure_branch_sections_from_department()

    if request.method == "POST":
        action = (request.form.get("_action") or "create_branch").strip()
        if action == "create_branch":
            name = (request.form.get("name") or "").strip()
            department = (request.form.get("department") or "").strip() or None
            if not name:
                flash("âš ï¸ ÙŠØ¬Ø¨ Ø¥Ø¯Ø®Ø§Ù„ Ø§Ø³Ù… Ø§Ù„ÙØ±Ø¹", "danger")
            else:
                existing = Branch.query.filter_by(name=name).first()
                if existing:
                    flash("âš ï¸ Ø§Ù„ÙØ±Ø¹ Ù…ÙˆØ¬ÙˆØ¯ Ù…Ø³Ø¨Ù‚Ø§Ù‹", "warning")
                else:
                    branch = Branch(name=name, department=department)
                    db.session.add(branch)
                    db.session.commit()
                    flash("âœ… ØªÙ… Ø¥Ø¶Ø§ÙØ© Ø§Ù„ÙØ±Ø¹", "success")
                    return redirect(url_for("manage_branches"))
        elif action == "add_section":
            branch_id = request.form.get("branch_id")
            section_name = (request.form.get("section_name") or "").strip()
            try:
                branch_id = int(branch_id)
            except Exception:
                branch_id = None
            if not branch_id or not section_name:
                flash("âš ï¸ ÙŠØ±Ø¬Ù‰ Ø§Ø®ØªÙŠØ§Ø± Ø§Ù„ÙØ±Ø¹ ÙˆØ§Ø³Ù… Ø§Ù„Ù‚Ø³Ù…", "danger")
            else:
                # ØªÙˆØ­ÙŠØ¯ ÙƒØªØ§Ø¨Ø© Ø§Ø³Ù… Ø§Ù„Ù‚Ø³Ù… Ù„Ø³Ù‡ÙˆÙ„Ø© Ø§Ù„ØªÙˆØ¬ÙŠÙ‡
                normalized = section_name.strip().lower()
                # Ø£Ø³Ù…Ø§Ø¡ Ù…Ø®ØªØµØ±Ø© Ù…Ù‚Ø¨ÙˆÙ„Ø©
                aliases = {
                    "consultation": "consultations",
                    "consulting": "consultations",
                    "Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª": "consultations",
                    "Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª Ø§Ù„Ù‡Ù†Ø¯Ø³ÙŠØ©": "consultations",
                    "valuation": "valuation",
                    "Ø§Ù„ØªØ«Ù…ÙŠÙ†": "valuation",
                    "owners": "owners_associations",
                    "Ø¬Ù…Ø¹ÙŠØ§Øª": "owners_associations",
                    "Ø¬Ù…Ø¹ÙŠØ§Øª Ø§Ù„Ù…Ù„Ø§Ùƒ": "owners_associations",
                    "Ø§Ø¯Ø§Ø±Ø© Ø¬Ù…Ø¹ÙŠØ§Øª Ø§Ù„Ù…Ù„Ø§Ùƒ": "owners_associations",
                    "Ø¥Ø¯Ø§Ø±Ø© Ø¬Ù…Ø¹ÙŠØ§Øª Ø§Ù„Ù…Ù„Ø§Ùƒ": "owners_associations",
                    "properties": "property_management",
                    "Ø§Ù„Ù…Ù…ØªÙ„ÙƒØ§Øª": "property_management",
                    "Ø§Ø¯Ø§Ø±Ø© Ø§Ù„Ù…Ù…ØªÙ„ÙƒØ§Øª": "property_management",
                    "Ø¥Ø¯Ø§Ø±Ø© Ø§Ù„Ù…Ù…ØªÙ„ÙƒØ§Øª": "property_management",
                    "Ø§Ø¯Ø§Ø±Ø© Ø§Ù„Ø¹Ù‚Ø§Ø±Ø§Øª": "property_management",
                    "Ø¥Ø¯Ø§Ø±Ø© Ø§Ù„Ø¹Ù‚Ø§Ø±Ø§Øª": "property_management",
                }
                normalized = aliases.get(normalized, normalized)
                exists = BranchSection.query.filter_by(branch_id=branch_id, name=normalized).first()
                if exists:
                    flash("âš ï¸ Ø§Ù„Ù‚Ø³Ù… Ù…ÙˆØ¬ÙˆØ¯ Ù„Ù‡Ø°Ø§ Ø§Ù„ÙØ±Ø¹", "warning")
                else:
                    db.session.add(BranchSection(branch_id=branch_id, name=normalized))
                    db.session.commit()
                    flash("âœ… ØªÙ… Ø¥Ø¶Ø§ÙØ© Ø§Ù„Ù‚Ø³Ù… Ù„Ù„ÙØ±Ø¹", "success")
                return redirect(url_for("manage_branches"))
        elif action == "delete_section":
            sid = request.form.get("section_id")
            try:
                sid = int(sid)
            except Exception:
                sid = None
            if not sid:
                flash("âš ï¸ Ù‚Ø³Ù… ØºÙŠØ± ØµØ§Ù„Ø­", "danger")
            else:
                s = BranchSection.query.get(sid)
                if not s:
                    flash("âš ï¸ Ø§Ù„Ù‚Ø³Ù… ØºÙŠØ± Ù…ÙˆØ¬ÙˆØ¯", "warning")
                else:
                    db.session.delete(s)
                    db.session.commit()
                    flash("âœ… ØªÙ… Ø­Ø°Ù Ø§Ù„Ù‚Ø³Ù…", "success")
            return redirect(url_for("manage_branches"))

    branches = Branch.query.all()
    return render_template("manage_branches.html", branches=branches)

# ğŸ—‘ Ø­Ø°Ù ÙØ±Ø¹
@app.route("/delete_branch/<int:bid>")
def delete_branch(bid):
    if session.get("role") != "manager":
        return redirect(url_for("login"))

    branch = Branch.query.get_or_404(bid)

    # ØªØ­Ù‚Ù‚ Ø¥Ø°Ø§ Ø§Ù„ÙØ±Ø¹ Ù…Ø±ØªØ¨Ø· Ø¨Ù…ÙˆØ¸ÙÙŠÙ†
    if branch.users:
        flash("ğŸš« Ù„Ø§ ÙŠÙ…ÙƒÙ† Ø­Ø°Ù ÙØ±Ø¹ Ù…Ø±ØªØ¨Ø· Ø¨Ù…ÙˆØ¸ÙÙŠÙ†", "danger")
        return redirect(url_for("manage_branches"))

    db.session.delete(branch)
    db.session.commit()
    flash("âœ… ØªÙ… Ø­Ø°Ù Ø§Ù„ÙØ±Ø¹", "success")
    return redirect(url_for("manage_branches"))

# âœ… ØµÙØ­Ø© Ø§Ù„Ø¹Ù…ÙˆÙ„Ø§Øª
@app.route("/commission", methods=["GET", "POST"])
def commissions_page():
    role = session.get("role")
    if not role:
        return redirect(url_for("login"))

    # ğŸ”¹ Ø¥Ø°Ø§ ÙƒØ§Ù† Ø§Ù„Ù…Ø¯ÙŠØ± â†’ ÙŠÙ‚Ø¯Ø± ÙŠÙÙ„ØªØ± Ø¨Ø§Ù„Ø£Ø³Ù…Ø§Ø¡ Ø§Ù„ØªÙŠ Ø¬Ù„Ø¨Øª Ù…Ø¹Ø§Ù…Ù„Ø§Øª
    selected_brought_by = None
    brought_by_names = []
    if role == "manager":
        # Ø§Ø¬Ù„Ø¨ ÙƒÙ„ Ø§Ù„Ø£Ø³Ù…Ø§Ø¡ Ø§Ù„ØªÙŠ Ø¬Ù„Ø¨Øª Ù…Ø¹Ø§Ù…Ù„Ø§Øª (Ø¨Ø¯ÙˆÙ† ØªÙƒØ±Ø§Ø±)
        raw_names = db.session.query(Transaction.brought_by)\
            .filter(Transaction.brought_by.isnot(None), Transaction.brought_by != "")\
            .distinct().all()
        # Ù†Ø¸Ù ÙˆÙ…ÙŠÙ‘Ø² Ø§Ù„Ø£Ø³Ù…Ø§Ø¡
        brought_by_names = sorted({name.strip() for (name,) in raw_names if name and name.strip()})
        if request.method == "POST":
            selected_brought_by = request.form.get("brought_by") or None
    else:
        # Ø§Ù„Ù…ÙˆØ¸Ù ÙŠØ´ÙˆÙ Ø¨ÙŠØ§Ù†Ø§ØªÙ‡ ÙÙ‚Ø·
        selected_brought_by = session.get("username")

    query = Transaction.query.filter(Transaction.payment_status == "Ù…Ø¯ÙÙˆØ¹Ø©")

    if selected_brought_by:
        # Ø§Ø­ØªØ³Ø§Ø¨ Ø¹Ù…ÙˆÙ„Ø© Ø­Ø³Ø¨ Ù…Ù† Ø¬Ù„Ø¨ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø©
        query = query.filter(Transaction.brought_by == selected_brought_by)

    transactions = query.all()

    # ğŸ”¹ Ø­Ø³Ø§Ø¨ Ø§Ù„Ø¹Ù‚Ø§Ø±Ø§Øª
    real_estate_txns = [t for t in transactions if t.transaction_type == "real_estate"]
    real_estate_income = sum(t.fee for t in real_estate_txns)
    # ÙƒÙ„ 50 Ø±ÙŠØ§Ù„ = 1 Ù…Ø¹Ø§Ù…Ù„Ø©
    real_estate_count = sum(max(1, int(t.fee // 50)) for t in real_estate_txns)

    # ğŸ”¹ Ø­Ø³Ø§Ø¨ Ø§Ù„Ø³ÙŠØ§Ø±Ø§Øª
    vehicle_txns = [t for t in transactions if t.transaction_type == "vehicle"]
    vehicle_income = sum(t.fee for t in vehicle_txns)
    # ÙƒÙ„ 3 Ø³ÙŠØ§Ø±Ø§Øª = 1 Ù…Ø¹Ø§Ù…Ù„Ø©
    vehicle_count = len(vehicle_txns) // 3

    # ğŸ”¹ Ø§Ù„Ø¥Ø¬Ù…Ø§Ù„ÙŠ
    total_income = real_estate_income + vehicle_income
    total_count = real_estate_count + vehicle_count

    # ğŸ”¹ Ø§Ù„Ø¹Ù…ÙˆÙ„Ø© (Ø¨Ø¹Ø¯ 30 Ù…Ø¹Ø§Ù…Ù„Ø©)
    commission_count = max(0, total_count - 30)
    commission = commission_count * 15

    return render_template(
        "commission.html",
        brought_by_names=brought_by_names,
        role=role,
        selected_brought_by=selected_brought_by,
        real_estate_count=real_estate_count,
        real_estate_income=real_estate_income,
        vehicle_count=vehicle_count,
        vehicle_income=vehicle_income,
        total_income=total_income,
        total_count=total_count,
        commission=commission
    )


# ---------------- Ù„ÙˆØ­Ø© Ø§Ù„Ù…Ø¯ÙŠØ± ----------------
VAPID_PUBLIC_KEY = "BFNeZpjEro8pwFxR1H20twlTd2pL5MZtWrDATu4ME2RcbzhN"  # Ø§Ù„Ù…ÙØªØ§Ø­ Ø§Ù„Ù„ÙŠ ÙˆÙ„Ø¯ØªÙ‡
# ğŸ“Œ Ù„ÙˆØ­Ø© Ø§Ù„Ù…Ø¯ÙŠØ±
@app.route("/manager")
def manager_dashboard():
    if session.get("role") != "manager":
        return redirect(url_for("login"))

    now = datetime.utcnow()
    hidden_statuses =     "in_progress"   ,  "Ø¨Ø¥Ù†ØªØ¸Ø§Ø± Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³" , "Ù‚ÙŠØ¯ Ø§Ù„Ù…Ø¹Ø§ÙŠÙ†Ø©", "ğŸ“‘ ØªÙ‚Ø±ÙŠØ± Ù…Ø±ÙÙˆØ¹" ,  "Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³",
    # Current month boundaries (UTC) for Postgres-compatible filtering
    month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    if now.month == 12:
        next_month_start = datetime(now.year + 1, 1, 1)
    else:
        next_month_start = datetime(now.year, now.month + 1, 1)
    VAPID_PUBLIC_KEY = "BFNeZpjEro8pwFxR1H20twlTd2pL5MZtWrDATu4ME2RcbzhN"  # Ø§Ù„Ù…ÙØªØ§Ø­ Ø§Ù„Ù„ÙŠ ÙˆÙ„Ø¯ØªÙ‡

    # âœ… ÙÙ‚Ø· Ù…Ø¹Ø§Ù…Ù„Ø§Øª Ø§Ù„Ø¹Ù‚Ø§Ø±Ø§Øª ØªØ¸Ù‡Ø± Ø¹Ù†Ø¯ Ø§Ù„Ù…Ø¯ÙŠØ± + Ø§Ø³ØªØ¨Ø¹Ø§Ø¯ Ø§Ù„Ø­Ø§Ù„Ø§Øª Ø§Ù„Ù…Ø®ÙÙŠØ©
    transactions = Transaction.query.filter(
        Transaction.transaction_type == "real_estate",
        ~Transaction.status.in_(hidden_statuses),
        Transaction.status.notin_(["Ù…Ø±ÙÙˆØ¶Ø©",  "Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ù…Ø§Ù„ÙŠØ©"  , "Ù…ÙƒØªÙ…Ù„Ø©", "Ù…Ù†Ø¬Ø²Ø©"])
    ).order_by(Transaction.id.desc()).all()
    
    users = User.query.all()

    branches_data = []
    branches = Branch.query.all()
    for b in branches:
        income = db.session.query(func.coalesce(func.sum(Payment.amount), 0.0))\
            .join(Transaction, Payment.transaction_id == Transaction.id)\
            .filter(Transaction.branch_id == b.id)\
            .scalar() or 0.0
        expenses = db.session.query(func.coalesce(func.sum(Expense.amount), 0.0))\
            .filter(Expense.branch_id == b.id)\
            .scalar() or 0.0
        profit = income - expenses

        # âœ… Ø¥Ø­ØµØ§Ø¦ÙŠØ© Ø§Ù„Ø¨Ù†ÙˆÙƒ (Ø§Ù„Ø´Ù‡Ø± Ø§Ù„Ø­Ø§Ù„ÙŠ) Ù„ÙƒÙ† ÙÙ‚Ø· Ù„Ù„Ø¹Ù‚Ø§Ø±Ø§Øª
        banks_stats = (
            db.session.query(Bank.name, func.count(Transaction.id))
            .join(Transaction, Transaction.bank_id == Bank.id)
            .filter(Transaction.branch_id == b.id)
            .filter(Transaction.transaction_type == "real_estate")   # ğŸš« Ø§Ø³ØªØ¨Ø¹Ø§Ø¯ Ø§Ù„Ø³ÙŠØ§Ø±Ø§Øª
            .filter(Transaction.date >= month_start)
            .filter(Transaction.date < next_month_start)
            .group_by(Bank.name)
            .all()
        )
        banks_list = [{"name": x[0], "count": x[1]} for x in banks_stats]

        branches_data.append({
            "name": b.name,
            "income": income,
            "expenses": expenses,
            "profit": profit,
            "banks": banks_list
        })

    # ğŸ”” Ù…Ø³ØªÙ†Ø¯Ø§Øª Ø¹Ù„Ù‰ ÙˆØ´Ùƒ Ø§Ù„Ø§Ù†ØªÙ‡Ø§Ø¡ Ø®Ù„Ø§Ù„ 30 ÙŠÙˆÙ…Ù‹Ø§ (ÙƒÙ„ Ø§Ù„ÙØ±ÙˆØ¹)
    expiring_docs = BranchDocument.query.filter(
        BranchDocument.expires_at != None,
        BranchDocument.expires_at <= (now + timedelta(days=30))
    ).order_by(BranchDocument.expires_at.asc()).all()

    # âš ï¸ Ù…Ø¹Ø§Ù…Ù„Ø§Øª Ù…ØªØ£Ø®Ø±Ø© (5 Ø³Ø§Ø¹Ø§Øª)
    five_hours_ago = now - timedelta(hours=5)
    # 1) ØµØ§Ø± Ù„Ù‡Ø§ 5 Ø³Ø§Ø¹Ø§Øª ÙˆÙ„Ù… ÙŠØªÙ… Ø§Ø³ØªÙ„Ø§Ù…Ù‡Ø§ Ù…Ù† Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³
    delayed_not_received = Transaction.query.filter(
        Transaction.transaction_type == "real_estate",
        Transaction.status == "Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³",
        or_(Transaction.assigned_to == None, Transaction.assigned_to.is_(None)),
        Transaction.date <= five_hours_ago,
    ).order_by(Transaction.date.asc()).all()

    # 2) ØªÙ… Ø§Ø³ØªÙ„Ø§Ù…Ù‡Ø§ (Ù‚ÙŠØ¯ Ø§Ù„Ù…Ø¹Ø§ÙŠÙ†Ø©/Ø§Ù„ØªÙ†ÙÙŠØ°) Ù…Ù†Ø° 5 Ø³Ø§Ø¹Ø§Øª Ù„ÙƒÙ† Ù„Ù… ÙŠÙØ±ÙØ¹ Ø§Ù„ØªÙ‚Ø±ÙŠØ± Ø¨Ø¹Ø¯
    delayed_received_no_report = Transaction.query.filter(
        Transaction.transaction_type == "real_estate",
        Transaction.status.in_(["Ù‚ÙŠØ¯ Ø§Ù„Ù…Ø¹Ø§ÙŠÙ†Ø©", "Ù‚ÙŠØ¯ Ø§Ù„ØªÙ†ÙÙŠØ°"]),
        Transaction.date <= five_hours_ago,
        and_(
            or_(Transaction.report_file == None, func.length(func.trim(Transaction.report_file)) == 0),
            or_(Transaction.engineer_report == None, func.length(func.trim(Transaction.engineer_report)) == 0),
            or_(Transaction.report_b2_file_name == None, func.length(func.trim(Transaction.report_b2_file_name)) == 0),
        ),
    ).order_by(Transaction.date.asc()).all()

    # Ø¨ÙŠØ§Ù†Ø§Øª Ù…Ø³Ø§Ø¹Ø¯Ø© Ù„Ø±ÙØ¹ Ù‚ÙˆØ§Ù„Ø¨ Ø§Ù„ÙÙˆØ§ØªÙŠØ± Ù…Ø¨Ø§Ø´Ø±Ø© Ù…Ù† Ù„ÙˆØ­Ø© Ø§Ù„Ù…Ø¯ÙŠØ±
    current_user = User.query.get(session.get("user_id"))
    current_branch_id = getattr(current_user, "branch_id", None)
    template_branches = Branch.query.order_by(Branch.name.asc()).all()

    return render_template(
        "manager_dashboard.html",
        transactions=transactions,
        users=users,
        branches=branches_data,
        vapid_public_key=VAPID_PUBLIC_KEY,
        net_profit=sum(b["profit"] for b in branches_data),
        expiring_docs=expiring_docs,
        delayed_not_received=delayed_not_received,
        delayed_received_no_report=delayed_received_no_report,
        template_branches=template_branches,
        current_branch_id=current_branch_id
    )


# âœ… ØªØ­Ø¯ÙŠØ« Ø­Ø§Ù„Ø© Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø©
@app.route("/update_status/<int:tid>/<status>")
def update_status(tid, status):
    role = session.get("role")
    if not role:
        return redirect(url_for("login"))

    t = Transaction.query.get_or_404(tid)

    # âœ… Ø§Ù„Ù…Ø¯ÙŠØ± Ù…Ø§ ÙŠÙ‚Ø¯Ø± ÙŠØ±Ø³Ù„ Ù„Ù„Ù…Ø§Ù„ÙŠØ©
    if role == "manager" and status == "Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ø¯ÙØ¹":
        flash("âš ï¸ Ù„Ø§ ÙŠÙ…ÙƒÙ† Ù„Ù„Ù…Ø¯ÙŠØ± Ø¥Ø±Ø³Ø§Ù„ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ù„Ù„Ù…Ø§Ù„ÙŠØ©. ÙÙ‚Ø· Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³.", "danger")
        return redirect(url_for("manager_dashboard"))

    # âœ… Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³ ÙÙ‚Ø· ÙŠØ±Ø³Ù„ Ù„Ù„Ù…Ø§Ù„ÙŠØ©
    if role == "engineer" and status == "Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ø¯ÙØ¹":
        if not t.engineer_report:  # ØªØªØ£ÙƒØ¯ Ø¥Ù†Ù‡ ÙƒØªØ¨ Ø§Ù„ØªÙ‚Ø±ÙŠØ±
            flash("âš ï¸ Ù„Ø§ ÙŠÙ…ÙƒÙ†Ùƒ Ø¥Ø±Ø³Ø§Ù„ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ù„Ù„Ù…Ø§Ù„ÙŠØ© Ø¨Ø¯ÙˆÙ† ØªÙ‚Ø±ÙŠØ±.", "danger")
            return redirect(url_for("engineer_dashboard"))

    t.status = status
    db.session.commit()
    bump_transactions_version()
      # Ø¨Ø¹Ø¯ db.session.commit() ÙÙŠ send_to_visit Ø£Ùˆ update_status
    engineer = User.query.filter_by(role="engineer").first()
    if engineer:
        send_notification(engineer.id, "ğŸ“© Ù…Ø¹Ø§Ù…Ù„Ø© Ø¬Ø¯ÙŠØ¯Ø©", f"ØªÙ… Ø¥Ø±Ø³Ø§Ù„ Ù…Ø¹Ø§Ù…Ù„Ø© Ø±Ù‚Ù… {t.id} Ø¥Ù„ÙŠÙƒ")

    # Ø¥Ø´Ø¹Ø§Ø± Ø§Ù„Ù…Ø§Ù„ÙŠØ© Ø¹Ù†Ø¯Ù…Ø§ ØªØµØ¨Ø­ Ø§Ù„Ø­Ø§Ù„Ø© "Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ø¯ÙØ¹"
    if status == "Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ø¯ÙØ¹":
        try:
            finances = User.query.filter_by(role="finance").all()
            for fin in finances:
                send_notification(fin.id, "ğŸ’³ Ù…Ø¹Ø§Ù…Ù„Ø© Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ø¯ÙØ¹", f"Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ø±Ù‚Ù… {t.id} Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ø¯ÙØ¹")
        except Exception:
            pass


    if role == "manager":
        return redirect(url_for("manager_dashboard"))
    elif role == "engineer":
        return redirect(url_for("engineer_dashboard"))
    elif role == "employee":
        return redirect(url_for("employee_dashboard"))
    return redirect(url_for("login"))


# Ø±Ø§ÙˆØª Ø§Ø¹ØªÙ…Ø§Ø¯ Ù…Ù† Ø§Ù„Ù…Ø¯ÙŠØ±
@app.route("/approve_transaction/<int:tid>")
def approve_transaction(tid):
    if session.get("role") != "manager":
        return redirect(url_for("login"))

    transaction = Transaction.query.get_or_404(tid)
    transaction.status = "Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³"   # ğŸ‘ˆ ÙƒÙ„ Ù…Ù‡Ù†Ø¯Ø³ Ø¨Ø§Ù„ÙØ±Ø¹ Ø¨ÙŠØ´ÙˆÙÙ‡Ø§
    db.session.commit()
    bump_transactions_version()

    flash("âœ… ØªÙ… Ø§Ø¹ØªÙ…Ø§Ø¯ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© ÙˆØ¥Ø±Ø³Ø§Ù„Ù‡Ø§ Ù„Ø¬Ù…ÙŠØ¹ Ù…Ù‡Ù†Ø¯Ø³ÙŠ Ø§Ù„ÙØ±Ø¹", "success")
    return redirect(url_for("manager_dashboard"))





# ğŸ¢ Ø¥Ø¶Ø§ÙØ© ÙØ±Ø¹ Ø¬Ø¯ÙŠØ¯
@app.route("/add_branch", methods=["POST"])
def add_branch():
    if session.get("role") != "manager":
        return redirect(url_for("login"))
    # ØªØ£ÙƒØ¯ Ù…Ù† ÙˆØ¬ÙˆØ¯ Ø¹Ù…ÙˆØ¯ department
    ensure_branch_department_column()

    name = (request.form.get("name") or "").strip()
    department = (request.form.get("department") or "").strip() or None
    if name:
        db.session.add(Branch(name=name, department=department))
        db.session.commit()
        flash("âœ… ØªÙ… Ø¥Ø¶Ø§ÙØ© Ø§Ù„ÙØ±Ø¹ Ø¨Ù†Ø¬Ø§Ø­", "success")
    else:
        flash("âš ï¸ ÙŠØ¬Ø¨ Ø¥Ø¯Ø®Ø§Ù„ Ø§Ø³Ù… Ø§Ù„ÙØ±Ø¹", "danger")
    return redirect(url_for("manager_dashboard"))


# ğŸ”— ÙØªØ­ ÙˆØ§Ø¬Ù‡Ø© Ø§Ù„Ù‚Ø³Ù… Ø§Ù„Ø®Ø§ØµØ© Ø¨ÙØ±Ø¹ Ù…Ø¹ÙŠÙ‘Ù†
@app.route("/branch/<int:bid>/interface", endpoint="open_branch_interface")
def open_branch_interface(bid: int):
    if not session.get("user_id"):
        return redirect(url_for("login"))

    # Ø§Ù„Ø³Ù…Ø§Ø­ Ø¨ØªÙ…Ø±ÙŠØ± Ø§Ù„Ù‚Ø³Ù… Ø§Ù„Ù…Ø·Ù„ÙˆØ¨ Ø¹Ø¨Ø± Ø¨Ø§Ø±Ø§Ù…ÙŠØªØ±
    section_param = (request.args.get("section") or "").strip().lower()
    if section_param:
        return _redirect_to_section(section_param)

    ensure_branch_department_column()
    ensure_branch_sections_from_department()
    b = Branch.query.get_or_404(bid)
    dept = (b.department or "").lower()

    # Ù„Ùˆ ØªÙ… ØªØ¹Ø±ÙŠÙ Ø£Ù‚Ø³Ø§Ù… Ù…ØªØ¹Ø¯Ø¯Ø© Ù„Ù‡Ø°Ø§ Ø§Ù„ÙØ±Ø¹ØŒ ÙˆÙØ¬Ø¯ Ù‚Ø³Ù… ÙˆØ§Ø­Ø¯ ÙÙ‚Ø·ØŒ Ø§ÙØªØ­Ù‡ Ù…Ø¨Ø§Ø´Ø±Ø©Ù‹
    try:
        branch_sections = BranchSection.query.filter_by(branch_id=bid).all()
    except Exception:
        branch_sections = []
    if branch_sections and len(branch_sections) == 1:
        return _redirect_to_section((branch_sections[0].name or "").lower())

    # Ø®Ø±Ø§Ø¦Ø· Ø¨Ø³ÙŠØ·Ø© Ù„Ù„Ø£Ù‚Ø³Ø§Ù… Ø¥Ù„Ù‰ Ø§Ù„ÙˆØ§Ø¬Ù‡Ø§Øª Ø§Ù„Ø­Ø§Ù„ÙŠØ©
    if dept in ("consultations", "consultation", "consulting", "Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª"):
        return redirect(url_for("consultations_list"))
    if dept in ("finance", "financial", "Ø§Ù„Ù…Ø§Ù„ÙŠØ©"):
        return redirect(url_for("finance_dashboard"))

    # Ø§Ù„Ø§ÙØªØ±Ø§Ø¶ÙŠ: ÙˆØ§Ø¬Ù‡Ø§Øª Ø§Ù„ØªØ«Ù…ÙŠÙ† Ø§Ù„Ù…Ø¹ØªØ§Ø¯Ø© Ø¨Ø­Ø³Ø¨ Ø¯ÙˆØ± Ø§Ù„Ù…Ø³ØªØ®Ø¯Ù…
    role = session.get("role")
    if role == "manager":
        return redirect(url_for("manager_dashboard"))
    if role == "employee":
        return redirect(url_for("employee_dashboard"))
    if role == "engineer":
        return redirect(url_for("engineer_dashboard"))
    if role == "finance":
        return redirect(url_for("finance_dashboard"))
    return redirect(url_for("index"))

@app.route("/branch/<int:bid>/interface/<string:section>", endpoint="open_branch_section")
def open_branch_section(bid: int, section: str):
    if not session.get("user_id"):
        return redirect(url_for("login"))
    # Ù„Ø§ Ø­Ø§Ø¬Ø© Ù„Ù„ØªØ­Ù‚Ù‚ Ù…Ù† Ù…Ù„ÙƒÙŠØ© Ø§Ù„ÙØ±Ø¹ Ù‡Ù†Ø§ØŒ Ù…Ø¬Ø±Ø¯ ØªÙˆØ¬ÙŠÙ‡ Ø­Ø³Ø¨ Ø§Ù„Ù‚Ø³Ù…
    return _redirect_to_section(section)

def _redirect_to_section(section: str):
    s = (section or "").strip().lower()
    if s in ("consultations", "consultation", "consulting", "Ø§Ù„Ø§Ø³ØªØ´Ø§Ø±Ø§Øª"):
        return redirect(url_for("consultations_list"))
    if s in ("finance", "financial", "Ø§Ù„Ù…Ø§Ù„ÙŠØ©"):
        return redirect(url_for("finance_dashboard"))
    # ğŸ†• Ø£Ù‚Ø³Ø§Ù… Ø¥Ø¶Ø§ÙÙŠØ©: Ø¥Ø¯Ø§Ø±Ø© Ø¬Ù…Ø¹ÙŠØ§Øª Ø§Ù„Ù…Ù„Ø§Ùƒ / Ø¥Ø¯Ø§Ø±Ø© Ø§Ù„Ù…Ù…ØªÙ„ÙƒØ§Øª
    if s in ("owners_associations", "Ø¬Ù…Ø¹ÙŠØ§Øª Ø§Ù„Ù…Ù„Ø§Ùƒ", "owners", "associations"):
        return redirect(url_for("owners_associations_dashboard"))
    if s in ("property_management", "Ø§Ø¯Ø§Ø±Ø© Ø§Ù„Ù…Ù…ØªÙ„ÙƒØ§Øª", "properties"):
        return redirect(url_for("property_management_dashboard"))
    # valuation Ø£Ùˆ ØºÙŠØ± Ù…Ø¹Ø±ÙˆÙ => Ø§Ù„Ø§ÙØªØ±Ø§Ø¶ÙŠ Ø­Ø³Ø¨ Ø§Ù„Ø¯ÙˆØ±
    role = session.get("role")
    if role == "manager":
        return redirect(url_for("manager_dashboard"))
    if role == "employee":
        return redirect(url_for("employee_dashboard"))
    if role == "engineer":
        return redirect(url_for("engineer_dashboard"))
    if role == "finance":
        return redirect(url_for("finance_dashboard"))
    return redirect(url_for("index"))


# ================= Department Dashboards (Owners Associations / Property Management) =================
@app.route("/owners_associations")
def owners_associations_dashboard():
    if not session.get("user_id"):
        return redirect(url_for("login"))
    # Ø³ÙŠØ§Ù‚ Ø¨Ø³ÙŠØ·: ÙØ±Ø¹ Ø§Ù„Ù…Ø³ØªØ®Ø¯Ù… (Ø¥Ù† ÙˆÙØ¬Ø¯)
    user = User.query.get(session.get("user_id"))
    branch = Branch.query.get(user.branch_id) if user and getattr(user, "branch_id", None) else None
    return render_template("owners_associations.html", user=user, branch=branch)


@app.route("/property_management")
def property_management_dashboard():
    if not session.get("user_id"):
        return redirect(url_for("login"))
    user = User.query.get(session.get("user_id"))
    branch = Branch.query.get(user.branch_id) if user and getattr(user, "branch_id", None) else None
    return render_template("property_management.html", user=user, branch=branch)


# ğŸ¦ Ø¥Ø¶Ø§ÙØ© Ø¨Ù†Ùƒ Ø¬Ø¯ÙŠØ¯
@app.route("/add_bank", methods=["GET", "POST"])
def add_bank():
    if session.get("role") != "manager":
        return redirect(url_for("login"))
    
    if request.method == "POST":
        name = request.form.get("name")
        if name:
            db.session.add(Bank(name=name))
            db.session.commit()
            flash("âœ… ØªÙ… Ø¥Ø¶Ø§ÙØ© Ø§Ù„Ø¨Ù†Ùƒ Ø¨Ù†Ø¬Ø§Ø­", "success")
            return redirect(url_for("manager_dashboard"))
        else:
            flash("âš ï¸ ÙŠØ¬Ø¨ Ø¥Ø¯Ø®Ø§Ù„ Ø§Ø³Ù… Ø§Ù„Ø¨Ù†Ùƒ", "danger")
    
    return render_template("add_bank.html")

@app.route("/transaction/<int:tid>")
def transaction_detail(tid):
    if session.get("role") != "manager":
        return redirect(url_for("login"))
    t = Transaction.query.get_or_404(tid)
    return render_template("transaction_detail.html", t=t)

# âœ… ØªÙˆÙ„ÙŠØ¯ Ø±Ù‚Ù… Ø§Ù„ØªÙ‚Ø±ÙŠØ±
def generate_report_number():
    last_number = db.session.query(Transaction.report_number) \
        .filter(Transaction.report_number.isnot(None)) \
        .order_by(Transaction.report_number.desc()) \
        .first()
    if last_number and last_number[0]:
        match = re.search(r"(\d+)", last_number[0])
        if match:
            return f"ref{int(match.group(1)) + 1}"
    return "ref1001"

# âœ… Ù†Ù‚Ù„ Ø§Ù„ØªØ«Ù…ÙŠÙ† Ø¥Ù„Ù‰ Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³
@app.route("/engineer/valuate/<int:tid>", methods=["POST"])
def engineer_valuate_transaction(tid):
    if session.get("role") != "engineer":
        return redirect(url_for("login"))

    t = Transaction.query.get_or_404(tid)

    if t.transaction_type == "real_estate":
        land_value     = parse_float_input(request.form.get("land_value", 0))
        building_value = parse_float_input(request.form.get("building_value", 0))
        total_estimate = land_value + building_value

        t.land_value      = land_value
        t.building_value  = building_value
        t.total_estimate  = total_estimate
        t.valuation_amount = total_estimate
        t.status          = "Ù‚ÙŠØ¯ Ø§Ù„Ù…Ø¹Ø§ÙŠÙ†Ø©"

        # âœ… ØªØ­Ø¯ÙŠØ« Ø°Ø§ÙƒØ±Ø© Ø§Ù„ØªØ«Ù…ÙŠÙ†
        memory = ValuationMemory.query.filter_by(
            state=t.state, region=t.region, bank_id=t.bank_id
        ).first()
        if memory:
            memory.price_per_meter = land_value / t.area if t.area > 0 else 0
            memory.updated_at = datetime.utcnow()
        else:
            memory = ValuationMemory(
                state=t.state,
                region=t.region,
                bank_id=t.bank_id,
                price_per_meter=land_value / t.area if t.area > 0 else 0
            )
            db.session.add(memory)

    elif t.transaction_type == "vehicle":
        vehicle_value = parse_float_input(request.form.get("vehicle_value", 0))
        t.total_estimate = vehicle_value
        t.valuation_amount = vehicle_value
        t.status = "Ù‚ÙŠØ¯ Ø§Ù„Ù…Ø¹Ø§ÙŠÙ†Ø©"

    # âœ… Ø¥Ø¶Ø§ÙØ© Ø±Ù‚Ù… Ù…Ø±Ø¬Ø¹ÙŠ Ø¥Ø°Ø§ Ù…Ø§ ÙƒØ§Ù† Ù…ÙˆØ¬ÙˆØ¯
    if not t.report_number:
        last_txn = Transaction.query.filter(
            Transaction.report_number != None
        ).order_by(Transaction.id.desc()).first()

        if last_txn and last_txn.report_number.startswith("ref"):
            last_num = int(last_txn.report_number.replace("ref", ""))
            t.report_number = f"ref{last_num + 1}"
        else:
            t.report_number = "ref1001"

    db.session.commit()
    bump_transactions_version()
    flash("âœ… ØªÙ… Ø­ÙØ¸ Ø§Ù„ØªØ«Ù…ÙŠÙ† Ø¨ÙˆØ§Ø³Ø·Ø© Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³", "success")
    return redirect(url_for("engineer_transaction_details", tid=tid))




@app.route("/save-subscription", methods=["POST"])
def save_subscription():
    if not session.get("user_id"):
        return {"error": "Unauthorized"}, 401

    data = request.get_json()
    if not data:
        return {"error": "Invalid subscription"}, 400

    # Ù†Ø­Ø°Ù Ø£ÙŠ Ø§Ø´ØªØ±Ø§Ùƒ Ù‚Ø¯ÙŠÙ… Ù„Ù†ÙØ³ Ø§Ù„Ù…Ø³ØªØ®Ø¯Ù…
    NotificationSubscription.query.filter_by(user_id=session["user_id"]).delete()

    sub = NotificationSubscription(user_id=session["user_id"],
                                   subscription_json=json.dumps(data))
    db.session.add(sub)
    db.session.commit()
    return {"success": True}


@app.route("/assign_to_engineer/<int:tid>/<int:engineer_id>")
def assign_to_engineer(tid, engineer_id):
    if session.get("role") != "manager":
        return redirect(url_for("login"))
    
    transaction = Transaction.query.get_or_404(tid)
    transaction.assigned_to = engineer_id
    transaction.status = "Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³"   # âœ… Ø§Ù„Ø¢Ù† ÙŠÙ‚Ø¯Ø± ÙŠØ´ÙˆÙÙ‡Ø§ Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³
    db.session.commit()

    flash("âœ… ØªÙ… ØªØ­ÙˆÙŠÙ„ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ø¥Ù„Ù‰ Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³")
    return redirect(url_for("manager_dashboard"))



@app.route("/new_transaction", methods=["GET", "POST"])
def new_transaction():
    if request.method == "POST":
        state = request.form.get("state")
        region = request.form.get("region")
        bank_id = request.form.get("bank_id")
        price = float(request.form.get("price") or 0)
        try:
            bank_id = int(bank_id) if bank_id else None
        except Exception:
            bank_id = None
        if bank_id is None:
            flash("âš ï¸ Ø§Ù„Ø±Ø¬Ø§Ø¡ Ø§Ø®ØªÙŠØ§Ø± Ø¨Ù†Ùƒ ØµØ­ÙŠØ­", "warning")
        else:
            save_price(state, region, bank_id, price)
            flash("ØªÙ… Ø­ÙØ¸ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© ÙˆØ§Ù„Ø³Ø¹Ø± ÙÙŠ Ø§Ù„Ø°Ø§ÙƒØ±Ø© âœ…", "success")
        return redirect(url_for("new_transaction"))
    state = request.args.get("state")
    region = request.args.get("region")
    bank_id = request.args.get("bank_id")
    suggested_price = None
    try:
        bank_id_int = int(bank_id) if bank_id else None
    except Exception:
        bank_id_int = None
    if state and region and bank_id_int:
        suggested_price = get_last_price(state, region, bank_id_int)
    banks = Bank.query.order_by(Bank.name.asc()).all()
    return render_template("new_transaction.html", suggested_price=suggested_price, banks=banks)


@app.route("/get_price", methods=["POST"])
def get_price():
    state = request.form.get("state")
    region = request.form.get("region")
    bank_id = request.form.get("bank_id")

    price_per_meter = 0.0
    if state and region and bank_id:
        try:
            bank_id_int = int(bank_id)
        except Exception:
            bank_id_int = None

        vm = None
        if bank_id_int is not None:
            vm = ValuationMemory.query.filter_by(
                state=state, region=region, bank_id=bank_id_int
            ).order_by(ValuationMemory.updated_at.desc()).first()

        if vm:
            price_per_meter = vm.price_per_meter
        else:
            lp = LandPrice.query.filter_by(
                state=state, region=region, bank_id=bank_id_int
            ).first()
            if lp:
                price_per_meter = lp.price_per_meter

    return {"price_per_meter": price_per_meter}





# ---------------- Ù„ÙˆØ­Ø© Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³ ----------------
# ğŸ‘¨â€ğŸ”§ Ù„ÙˆØ­Ø© Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³
@app.route("/engineer")
def engineer_dashboard():
    if session.get("role") != "engineer":
        return redirect(url_for("login"))

    engineer_id = session.get("user_id")
    engineer = User.query.get_or_404(engineer_id)

    transactions = Transaction.query.filter(
        Transaction.branch_id == engineer.branch_id,
        or_(
            Transaction.status == "Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³",
            and_(
                Transaction.assigned_to == engineer_id,
                Transaction.status.in_(["Ù‚ÙŠØ¯ Ø§Ù„Ù…Ø¹Ø§ÙŠÙ†Ø©", "Ù‚ÙŠØ¯ Ø§Ù„ØªÙ†ÙÙŠØ°"])
            )
        )
    ).order_by(Transaction.id.desc()).all()
                
    return render_template("engineer.html", transactions=transactions, engineer=engineer, vapid_public_key=VAPID_PUBLIC_KEY)


# âœ… Ø¹Ù†Ø¯ Ø§Ø³ØªÙ„Ø§Ù… Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø©
@app.route("/engineer_take/<int:tid>")
def engineer_take(tid):
    if session.get("role") != "engineer":
        return redirect(url_for("login"))

    t = Transaction.query.get_or_404(tid)
    engineer_id = session.get("user_id")

    # ğŸ†• ØªØ­Ø¯ÙŠØ¯ Ø£Ù† Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³ Ø§Ø³ØªÙ„Ù… Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø©
    t.assigned_to = engineer_id
    t.status = "Ù‚ÙŠØ¯ Ø§Ù„Ù…Ø¹Ø§ÙŠÙ†Ø©"

    # ğŸ†• ØªØ®ØµÙŠØµ Ø±Ø³Ø§Ù„Ø© Ø­Ø³Ø¨ Ù†ÙˆØ¹ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø©
    if t.transaction_type == "vehicle":
        flash("ğŸš— ØªÙ… Ø§Ø³ØªÙ„Ø§Ù… Ù…Ø¹Ø§Ù…Ù„Ø© Ø§Ù„Ø³ÙŠØ§Ø±Ø©", "success")
    else:
        flash("ğŸ  ØªÙ… Ø§Ø³ØªÙ„Ø§Ù… Ù…Ø¹Ø§Ù…Ù„Ø© Ø§Ù„Ø¹Ù‚Ø§Ø±", "success")

    db.session.commit()
    return redirect(url_for("engineer_dashboard"))


# ---------------- Ù‚ÙˆØ§Ù„Ø¨ Ø§Ù„ØªÙ‚Ø§Ø±ÙŠØ± ÙˆØ¥Ø¹Ø¯Ø§Ø¯ Ù…Ø­Ø±Ø± Ø§Ù„ØªÙ‚Ø±ÙŠØ± Ù„Ù„Ù…Ù‡Ù†Ø¯Ø³ ----------------
def get_template_by_type(template_type: str) -> ReportTemplate | None:
    return ReportTemplate.query.filter_by(template_type=template_type).first()


@app.route("/manager/report_templates", methods=["GET", "POST"])
def manage_report_templates():
    if session.get("role") != "manager":
        return redirect(url_for("login"))

    real_estate_tpl = get_template_by_type("real_estate")
    vehicle_tpl = get_template_by_type("vehicle")

    if request.method == "POST":
        # Ø­ÙØ¸ Ø§Ù„Ù†ØµÙˆØµ ÙÙ‚Ø· (ØªÙ… ØªØ¹Ø·ÙŠÙ„ Ø±ÙØ¹ Ø§Ù„Ù‚ÙˆØ§Ù„Ø¨)
        re_content = (request.form.get("real_estate_content") or "").strip()
        ve_content = (request.form.get("vehicle_content") or "").strip()

        if re_content:
            if not real_estate_tpl:
                real_estate_tpl = ReportTemplate(template_type="real_estate", content=re_content)
                db.session.add(real_estate_tpl)
            else:
                real_estate_tpl.content = re_content

        if ve_content:
            if not vehicle_tpl:
                vehicle_tpl = ReportTemplate(template_type="vehicle", content=ve_content)
                db.session.add(vehicle_tpl)
            else:
                vehicle_tpl.content = ve_content

        db.session.commit()
        bump_transactions_version()
        flash("âœ… ØªÙ… Ø­ÙØ¸ Ø§Ù„Ù‚ÙˆØ§Ù„Ø¨ Ø§Ù„Ù†ØµÙŠØ© (ØªÙ… ØªØ¹Ø·ÙŠÙ„ Ø§Ù„Ø±ÙØ¹)", "success")
        return redirect(url_for("manage_report_templates"))

    return render_template(
        "manager_report_templates.html",
        real_estate_content=real_estate_tpl.content if real_estate_tpl else "",
        vehicle_content=vehicle_tpl.content if vehicle_tpl else "",
        real_estate_files=[],
        vehicle_files=[]
    )


def extract_placeholders(template_text: str) -> list[str]:
    if not template_text:
        return []
    return sorted(set(re.findall(r"\{([a-zA-Z0-9_]+)\}", template_text)))


def default_values_for_placeholders(t: Transaction, placeholders: list[str]) -> dict[str, str]:
    mapping = {
        "client_name": t.client or "",
        "sketch_number": "",
        "bank_name": t.bank.name if t.bank else "",
        "bank_branch": t.bank_branch or "",
        "property_state": t.state or "",
        "property_region": t.region or "",
        "area": str(t.area or 0),
        "building_area": str(t.building_area or 0),
        "building_age": str(t.building_age or 0),
        "land_value": str(t.land_value or 0),
        "building_value": str(t.building_value or 0),
        "total_estimate": str(t.total_estimate or 0),
        "vehicle_type": t.vehicle_type or "",
        "vehicle_model": t.vehicle_model or "",
        "vehicle_year": t.vehicle_year or "",
        "vehicle_value": str(t.total_estimate or 0),
        "today": datetime.utcnow().strftime("%Y-%m-%d"),
        "transaction_id": str(t.id),
    }
    return {ph: mapping.get(ph, "") for ph in placeholders}


def fill_template(template_text: str, values: dict[str, str]) -> str:
    def repl(match):
        key = match.group(1)
        return str(values.get(key, match.group(0)))
    return re.sub(r"\{([a-zA-Z0-9_]+)\}", repl, template_text)


@app.route("/engineer/report_editor/<int:tid>", methods=["GET", "POST"])
def engineer_report_editor(tid):
    if session.get("role") != "engineer":
        return redirect(url_for("login"))

    t = Transaction.query.get_or_404(tid)
    template_type = t.transaction_type or "real_estate"
    tpl = get_template_by_type(template_type)
    template_text = tpl.content if tpl else ""

    placeholders = extract_placeholders(template_text)

    if request.method == "POST":
        # Ø¬Ù…Ø¹ Ø§Ù„Ù‚ÙŠÙ…
        values = {ph: (request.form.get(ph) or "").strip() for ph in placeholders}
        # ØªÙˆÙ„ÙŠØ¯ Ø§Ù„Ù†Øµ Ø§Ù„Ù†Ù‡Ø§Ø¦ÙŠ
        final_text = fill_template(template_text, values)

        t.engineer_report = final_text
        t.status = "ğŸ“‘ ØªÙ‚Ø±ÙŠØ± Ù…Ø¨Ø¯Ø¦ÙŠ"  # Ø­Ø§Ù„Ø© ÙˆØ³Ø·ÙŠØ© Ø­ØªÙ‰ Ø§Ù„Ø±ÙØ¹ Ø§Ù„Ù†Ù‡Ø§Ø¦ÙŠ PDF
        db.session.commit()
        bump_transactions_version()
        flash("âœ… ØªÙ… Ø­ÙØ¸ Ù†Øµ Ø§Ù„ØªÙ‚Ø±ÙŠØ± Ù…Ù† Ø§Ù„Ù‚Ø§Ù„Ø¨", "success")
        return redirect(url_for("engineer_transaction_details", tid=tid))

    # Ù‚ÙŠÙ‘Ù… Ø§ÙØªØ±Ø§Ø¶ÙŠØ©
    defaults = default_values_for_placeholders(t, placeholders)

    return render_template(
        "engineer_report_editor.html",
        t=t,
        template_text=template_text,
        placeholders=placeholders,
        defaults=defaults,
    )



@app.route("/add_transaction_engineer", methods=["GET", "POST"])
def add_transaction_engineer():
    if session.get("role") != "engineer":
        return redirect(url_for("login"))

    user = User.query.get(session["user_id"])
    banks = Bank.query.all()

    if request.method == "POST":
        transaction_type = request.form.get("transaction_type")
        client_name = (request.form.get("client_name") or "").strip()
        client_phone = (request.form.get("client_phone") or "").strip()
        fee = float(request.form.get("fee") or 0)
        # ğŸ†• Ø§Ù„Ø­Ù‚ÙˆÙ„ Ø§Ù„Ø¬Ø¯ÙŠØ¯Ø©
        brought_by = (request.form.get("brought_by") or "").strip()
        visited_by = (request.form.get("visited_by") or "").strip()

        # âœ… ØªØ­Ù‚Ù‚ Ù…Ù† Ø±Ù‚Ù… Ø§Ù„Ø¹Ù…ÙŠÙ„ ÙˆØ§Ù„Ø­Ù‚ÙˆÙ„ Ø§Ù„Ø¥Ø¬Ø¨Ø§Ø±ÙŠØ© Ø§Ù„Ø¬Ø¯ÙŠØ¯Ø©
        if not client_phone:
            flash("âš ï¸ ÙŠØ¬Ø¨ Ø¥Ø¯Ø®Ø§Ù„ Ø±Ù‚Ù… Ø§Ù„Ø¹Ù…ÙŠÙ„", "danger")
            return redirect(url_for("add_transaction_engineer"))
        if not brought_by or not visited_by:
            flash("âš ï¸ ÙŠØ¬Ø¨ Ø¥Ø¯Ø®Ø§Ù„ Ù…Ù† Ø¬Ù„Ø¨ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© ÙˆÙ…Ù† Ù‚Ø§Ù… Ø¨Ø§Ù„Ø²ÙŠØ§Ø±Ø©", "danger")
            return redirect(url_for("add_transaction_engineer"))

        # ğŸ§¾ Ø­ÙØ¸/ØªØ­Ø¯ÙŠØ« Ø§Ù„Ø¹Ù…ÙŠÙ„ ÙÙŠ Ø¬Ø¯ÙˆÙ„ Ø§Ù„Ø¹Ù…Ù„Ø§Ø¡
        existing_customer = Customer.query.filter_by(phone=client_phone).first()
        if existing_customer:
            if client_name and existing_customer.name != client_name:
                existing_customer.name = client_name
        else:
            db.session.add(Customer(name=client_name or "-", phone=client_phone))
            db.session.flush()

        t = None

        if transaction_type == "real_estate":
            state = request.form.get("state")
            region = request.form.get("region")
            bank_id = request.form.get("bank_id")
            bank_branch = (request.form.get("bank_branch") or "").strip()
            bank_employee_name = (request.form.get("bank_employee_name") or "").strip()

            area = float(request.form.get("area") or 0)
            building_area = float(request.form.get("building_area") or 0)
            building_age = int(request.form.get("building_age") or 0)

            # ØªØ­Ù‚Ù‚ Ø£Ø³Ø§Ø³ÙŠ: Ø§Ù„Ø¨Ù†Ùƒ ÙˆÙØ±Ø¹ Ø§Ù„Ø¨Ù†Ùƒ Ù…Ø·Ù„ÙˆØ¨Ø§Ù†
            if not bank_id or not bank_branch:
                flash("âš ï¸ ÙŠØ±Ø¬Ù‰ Ø§Ø®ØªÙŠØ§Ø± Ø§Ù„Ø¨Ù†Ùƒ ÙˆÙƒØªØ§Ø¨Ø© ÙØ±Ø¹ Ø§Ù„Ø¨Ù†Ùƒ", "danger")
                return redirect(url_for("add_transaction_engineer"))

            # Ø§Ø­Ø³Ø¨ Ø³Ø¹Ø± Ø§Ù„Ù…ØªØ± Ù…Ù† Ø§Ù„Ø°Ø§ÙƒØ±Ø© Ø£Ùˆ Ù…Ù† Ø¬Ø¯ÙˆÙ„ Ø§Ù„Ø£Ø³Ø¹Ø§Ø± Ø¥Ù† ÙˆÙØ¬Ø¯
            price_per_meter = 0.0
            try:
                bank_id_int = int(bank_id)
            except Exception:
                bank_id_int = None

            if state and region and bank_id_int is not None:
                vm = ValuationMemory.query.filter_by(
                    state=state, region=region, bank_id=bank_id_int
                ).order_by(ValuationMemory.updated_at.desc()).first()
                if vm:
                    price_per_meter = vm.price_per_meter or 0.0
                else:
                    lp = LandPrice.query.filter_by(state=state, region=region, bank_id=bank_id_int).first()
                    if lp:
                        price_per_meter = lp.price_per_meter or 0.0

            # Ø­Ø³Ø§Ø¨ Ø§Ù„ØªØ«Ù…ÙŠÙ† Ø§Ù„Ø§Ø¨ØªØ¯Ø§Ø¦ÙŠ
            land_value = (area * price_per_meter) if price_per_meter else 0.0
            building_value = 0.0
            if building_area > 0 and building_age > 0:
                building_value = building_area * (185 / 50) * building_age
            total_estimate = land_value + building_value

            t = Transaction(
                client=client_name,
                employee=user.username,
                date=datetime.utcnow(),
                status="Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³",
                fee=fee,
                branch_id=user.branch_id,
                land_value=land_value,
                building_value=building_value,
                total_estimate=total_estimate,
                valuation_amount=total_estimate,
                area=area,
                building_area=building_area,
                building_age=building_age,
                state=state,
                region=region,
                bank_id=bank_id,
                bank_branch=bank_branch,
                bank_employee_name=bank_employee_name,
                brought_by=brought_by,
                visited_by=visited_by,
                created_by=user.id,
                transaction_type="real_estate",
                payment_status="ØºÙŠØ± Ù…Ø¯ÙÙˆØ¹Ø©",
                assigned_to=None   # âœ… ØºÙŠØ± Ù…Ø³Ù†Ø¯

            )

        elif transaction_type == "vehicle":
            vehicle_type  = request.form.get("vehicle_type")
            vehicle_model = request.form.get("vehicle_model")
            vehicle_year  = request.form.get("vehicle_year")
            vehicle_value = float(request.form.get("vehicle_value") or 0)

            # ØªØ­Ù‚Ù‚ Ø£Ø³Ø§Ø³ÙŠ: Ø§Ù„Ø¨Ù†Ùƒ ÙˆÙØ±Ø¹ Ø§Ù„Ø¨Ù†Ùƒ Ù…Ø·Ù„ÙˆØ¨Ø§Ù† Ù„Ù…Ø¹Ø§Ù…Ù„Ø§Øª Ø§Ù„Ù…Ø±ÙƒØ¨Ø§Øª Ø£ÙŠØ¶Ù‹Ø§
            bank_id = request.form.get("bank_id")
            bank_branch = (request.form.get("bank_branch") or "").strip()
            bank_employee_name = (request.form.get("bank_employee_name") or "").strip()
            if not bank_id or not bank_branch:
                flash("âš ï¸ ÙŠØ±Ø¬Ù‰ Ø§Ø®ØªÙŠØ§Ø± Ø§Ù„Ø¨Ù†Ùƒ ÙˆÙƒØªØ§Ø¨Ø© ÙØ±Ø¹ Ø§Ù„Ø¨Ù†Ùƒ", "danger")
                return redirect(url_for("add_transaction_engineer"))

            t = Transaction(
                client=client_name,
                employee=user.username,
                date=datetime.utcnow(),
                status="Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³",
                fee=fee,
                branch_id=user.branch_id,
                total_estimate=vehicle_value,
                created_by=user.id,
                transaction_type="vehicle",
                payment_status="ØºÙŠØ± Ù…Ø¯ÙÙˆØ¹Ø©",
                vehicle_type=vehicle_type,
                vehicle_model=vehicle_model,
                vehicle_year=vehicle_year,
                valuation_amount = vehicle_value,
                bank_id=bank_id,
                bank_branch=bank_branch,
                bank_employee_name=bank_employee_name,
                brought_by=brought_by,
                visited_by=visited_by,

                assigned_to=None   # âœ…
            )

        if t:
            # Ø±ÙØ¹ Ø§Ù„Ù…Ù„ÙØ§Øª Ø§Ù„Ù…Ø±ÙÙ‚Ø© Ù…Ù† Ù†Ù…ÙˆØ°Ø¬ Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³ ÙˆØ­ÙØ¸Ù‡Ø§ Ø¯Ø§Ø®Ù„ uploads
            try:
                files = request.files.getlist("files")
                saved_files = []
                for file in files:
                    if file and file.filename:
                        filename = secure_filename(file.filename)
                        file.save(os.path.join(app.config["UPLOAD_FOLDER"], filename))
                        saved_files.append(filename)
                if saved_files:
                    t.files = ",".join(saved_files)
            except Exception:
                # Ø¥Ù† Ø­Ø¯Ø« Ø®Ø·Ø£ Ø£Ø«Ù†Ø§Ø¡ Ø±ÙØ¹ Ø§Ù„Ù…Ù„ÙØ§ØªØŒ Ù†ÙÙƒÙ…Ù„ Ø¥Ù†Ø´Ø§Ø¡ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ø¨Ø¯ÙˆÙ† Ù…Ù„ÙØ§Øª
                pass

            db.session.add(t)
            db.session.commit()
            flash("âœ… ØªÙ… Ø¥Ø¶Ø§ÙØ© Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ø¨Ù†Ø¬Ø§Ø­", "success")
            return redirect(url_for("engineer_dashboard"))

    return render_template("add_transaction_engineer.html", banks=banks)


@app.route("/engineer/transaction/<int:tid>")
def engineer_transaction_details(tid):
    if session.get("role") != "engineer":
        return redirect(url_for("login"))

    engineer_id = session.get("user_id")
    engineer = User.query.get_or_404(engineer_id)

    t = Transaction.query.filter_by(id=tid, branch_id=engineer.branch_id).first_or_404()

    return render_template("engineer_transaction_details.html", t=t, engineer=engineer)


# âœ… Ø¹Ù†Ø¯ Ø±ÙØ¹ Ø§Ù„ØªÙ‚Ø±ÙŠØ± (Ø§Ù„Ù…Ø¹ØªÙ…Ø¯ ÙÙ‚Ø·)
@app.route("/engineer/upload_report/<int:tid>", methods=["POST"])
def engineer_upload_report(tid):
    if session.get("role") not in ["engineer", "manager"]:
        return redirect(url_for("login"))

    t = Transaction.query.get_or_404(tid)

    if "report_file" not in request.files or request.files["report_file"].filename == "":
        flash("âš ï¸ Ù„Ù… ÙŠØªÙ… Ø§Ø®ØªÙŠØ§Ø± Ù…Ù„Ù", "danger")
        return redirect(url_for("engineer_dashboard"))

    file = request.files["report_file"]
    original_name = file.filename
    # Ø§Ù„Ø³Ù…Ø§Ø­ ÙÙ‚Ø· Ø¨Ø±ÙØ¹ Ù…Ù„ÙØ§Øª PDF
    if not original_name.lower().endswith(".pdf"):
        flash("âš ï¸ ÙŠØ¬Ø¨ Ø±ÙØ¹ Ù…Ù„Ù Ø¨ØµÙŠØºØ© PDF.", "danger")
        return redirect(url_for("engineer_transaction_details", tid=tid))

    # ØªÙˆÙ„ÙŠØ¯ Ø±Ù‚Ù… Ø§Ù„ØªÙ‚Ø±ÙŠØ± Ø¥Ù† Ù„Ù… ÙŠÙˆØ¬Ø¯ Ù…Ø³Ø¨Ù‚Ù‹Ø§ (Ù‚Ø¨Ù„ Ø­ÙØ¸ Ø§Ù„Ù…Ù„Ù Ù„ØªØ³Ù…ÙŠØ© Ø§Ù„Ù…Ù„Ù)
    if not t.report_number:
        last_txn = Transaction.query.filter(
            Transaction.report_number != None
        ).order_by(Transaction.id.desc()).first()

        if last_txn and last_txn.report_number.startswith("ref"):
            last_num = int(last_txn.report_number.replace("ref", ""))
            t.report_number = f"ref{last_num + 1}"
        else:
            t.report_number = "ref1001"

    # Ø§Ø­ÙØ¸ Ø§Ù„Ù…Ù„Ù Ø¨Ø§Ø³Ù… Ø±Ù‚Ù… Ø§Ù„ØªÙ‚Ø±ÙŠØ± Ù…Ø«Ù„ ref1010.pdf
    filename = secure_filename(f"{t.report_number}.pdf")
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(filepath)

    # 1) Ø­Ø³Ø§Ø¨ Ø¨ØµÙ…Ø© Ø§Ù„Ù…Ù„Ù Ø§Ù„Ø£ØµÙ„ÙŠ ÙƒÙ…Ø§ ÙÙŠ index.html
    try:
        original_hash = compute_file_sha256(filepath)
    except Exception:
        original_hash = None

    # 2) Ø®ØªÙ… Ø§Ù„Ù…Ù„Ù Ø¨Ø§Ù„Ù€ QR Ø§Ù„Ø°ÙŠ ÙŠØ´ÙŠØ± Ø¥Ù„Ù‰ /file?hash=<hash>
    try:
        if original_hash:
            stamp_pdf_with_qr(filepath, original_hash)
        # Ø¥Ø¶Ø§ÙØ© Ø®ØªÙ… Ù†ØµÙŠ Ø¨Ø³ÙŠØ· Ø¥Ø¶Ø§ÙÙŠ (Ù…Ø¹Ù„ÙˆÙ…Ø§Øª Ø£Ø³Ø§Ø³ÙŠØ©)
        stamp_lines = [
            f"Ø±Ù‚Ù… Ø§Ù„ØªÙ‚Ø±ÙŠØ±: {t.report_number or '-'}",
            f"Ù…Ø¹Ø§Ù…Ù„Ø©: {t.id}",
            f"Ø§Ù„ØªØ§Ø±ÙŠØ®: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            "Ø®ØªÙ… Ø§Ù„Ù†Ø¸Ø§Ù… - ØºÙŠØ± Ù‚Ø§Ø¨Ù„ Ù„Ù„ØªØ¹Ø¯ÙŠÙ„",
        ]
        stamp_pdf_with_seal(filepath, "Ø®ØªÙ… Ø§Ù„ØªÙ‚Ø±ÙŠØ±", stamp_lines)
    except Exception:
        pass

    # 3) Ø­ÙØ¸ Ø§Ø³Ù… Ø§Ù„Ù…Ù„Ù ÙˆØªØ­Ø¯ÙŠØ« Ø§Ù„Ø­Ø§Ù„Ø©
    t.report_file = filename
    t.status = "ğŸ“‘ ØªÙ‚Ø±ÙŠØ± Ù…Ø±ÙÙˆØ¹"
    # ØªÙˆÙ„ÙŠØ¯ Ø±Ù…Ø² Ù…Ø´Ø§Ø±ÙƒØ© Ø¹Ø§Ù… Ø¥Ù† Ù„Ù… ÙŠÙƒÙ† Ù…ÙˆØ¬ÙˆØ¯Ù‹Ø§
    if not getattr(t, "public_share_token", None):
        try:
            t.public_share_token = secrets.token_urlsafe(24)
        except Exception:
            t.public_share_token = None

    # Ø®ØªÙ… Ø§Ù„Ù…Ù„Ù Ù…Ø¨Ø§Ø´Ø±Ø© Ø¨Ø¹Ø¯ Ø§Ù„Ø±ÙØ¹
    try:
        stamp_lines = [
            f"Ø±Ù‚Ù… Ø§Ù„ØªÙ‚Ø±ÙŠØ±: {t.report_number}",
            f"Ù…Ø¹Ø§Ù…Ù„Ø©: {t.id}",
            f"Ø§Ù„ØªØ§Ø±ÙŠØ®: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            "Ø®ØªÙ… Ø§Ù„Ù†Ø¸Ø§Ù… - ØºÙŠØ± Ù‚Ø§Ø¨Ù„ Ù„Ù„ØªØ¹Ø¯ÙŠÙ„"
        ]
        stamp_pdf_with_seal(filepath, "Ø®ØªÙ… Ø§Ù„ØªÙ‚Ø±ÙŠØ±", stamp_lines)
    except Exception:
        pass

    # 4) Ø­Ø³Ø§Ø¨ Ø¨ØµÙ…Ø© SHA-256 Ø§Ù„Ù†Ù‡Ø§Ø¦ÙŠØ© (Ø¨Ø¹Ø¯ Ø§Ù„Ø®ØªÙ…) Ù„Ù„Ø§Ø³ØªØ®Ø¯Ø§Ù… ÙÙŠ /verify Ùˆ/file
    try:
        final_hash = compute_file_sha256(filepath)
    except Exception:
        final_hash = None
    t.report_sha256 = final_hash or original_hash

    # 5) Ø±ÙØ¹ Ø§Ù„ØªÙ‚Ø±ÙŠØ± Ø§Ù„Ù†Ù‡Ø§Ø¦ÙŠ Ø¥Ù„Ù‰ Backblaze B2 ÙˆØªØ®Ø²ÙŠÙ† Ø¨ÙŠØ§Ù†Ø§ØªÙ‡ (Ø¥Ù† Ø£Ù…ÙƒÙ†)
    try:
        bucket = get_b2_bucket()
        with open(filepath, "rb") as fh:
            data = fh.read()
        safe_ref = (t.report_number or "ref").replace("/", "-")
        b2_name = f"reports/{t.id}_{safe_ref}_{int(time.time())}.pdf"
        uploaded = bucket.upload_bytes(data, file_name=b2_name)
        t.report_b2_file_name = b2_name
        t.report_b2_file_id = getattr(uploaded, "id_", None) or getattr(uploaded, "file_id", None)
    except Exception as e:
        # Ø¥Ø°Ø§ Ù„Ù… ØªÙØ¶Ø¨Ø· Ù…ÙØ§ØªÙŠØ­ B2 Ø£Ùˆ Ø­Ø¯Ø« Ø®Ø·Ø£ØŒ Ù†ØªØ¬Ø§Ù‡Ù„ Ø¨Ø¯ÙˆÙ† Ø¥ÙŠÙ‚Ø§Ù Ø§Ù„Ø¹Ù…Ù„ÙŠØ©
        try:
            print(
                f"âš ï¸ B2 upload failed for transaction {t.id}: {e} | "
                f"has_key_id={bool(app.config.get('B2_KEY_ID'))}, "
                f"has_application_key={bool(app.config.get('B2_APPLICATION_KEY'))}, "
                f"bucket_id={app.config.get('B2_BUCKET_ID')}"
            )
        except Exception:
            pass
        pass

    db.session.commit()
    bump_transactions_version()

    # Ø¨Ø¹Ø¯ db.session.commit() ÙÙŠ upload_report
    finance = User.query.filter_by(role="finance").first()
    employee = User.query.filter_by(username=t.employee).first()

    if finance:
        send_notification(finance.id, "ğŸ“„ ØªÙ‚Ø±ÙŠØ± Ø¬Ø¯ÙŠØ¯", f"ØªÙ… Ø±ÙØ¹ ØªÙ‚Ø±ÙŠØ± Ù„Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ø±Ù‚Ù… {t.id}")
    if employee:
        send_notification(employee.id, "ğŸ“„ ØªÙ‚Ø±ÙŠØ± Ø¬Ø§Ù‡Ø²", f"ØªÙ… Ø±ÙØ¹ Ø§Ù„ØªÙ‚Ø±ÙŠØ± Ù„Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ø±Ù‚Ù… {t.id}")

    flash(f"âœ… ØªÙ… Ø±ÙØ¹ Ø§Ù„ØªÙ‚Ø±ÙŠØ± (Ø§Ù„Ø±Ù‚Ù… Ø§Ù„Ù…Ø±Ø¬Ø¹ÙŠ: {t.report_number})", "success")

    # Ø§Ù„ØªÙˆØ¬ÙŠÙ‡ Ù…Ø¨Ø§Ø´Ø±Ø©Ù‹ Ø¥Ù„Ù‰ Ù‚Ø³Ù… Ø§Ù„ØªÙ‚Ø§Ø±ÙŠØ± Ø¨Ø¹Ø¯ Ø§Ù„Ø±ÙØ¹ ÙˆØ§Ù„Ø®ØªÙ…
    return redirect(url_for("reports_page"))

@app.route("/reports", endpoint="reports_page")
def reports():

    if session.get("role") not in ["manager", "admin", "finance" ,"employee" ,"engineer"]:
        return redirect(url_for("login"))

    q = request.args.get("q", "").strip()
    query = Transaction.query

    if q:
        query = query.filter(Transaction.report_number.contains(q))

    reports = query.filter(Transaction.report_file != None).order_by(Transaction.date.desc()).all()

    return render_template("reports.html", reports=reports, q=q)



    
# ---------------- ØµÙØ­Ø© Ø§Ù„Ù…Ø§Ù„ÙŠØ© ----------------
@app.route("/finance", methods=["GET", "POST"])
def finance_dashboard():
    if session.get("role") != "finance":
        return redirect(url_for("login"))

    user = User.query.get(session["user_id"])

    # âœ… Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø§Øª Ø§Ù„Ø®Ø§ØµØ© Ø¨Ø§Ù„ÙØ±Ø¹
    transactions = Transaction.query.filter_by(branch_id=user.branch_id).all()

    # âœ… Ø¥Ø¶Ø§ÙØ© Ù…ØµØ±ÙˆÙ Ø®Ø§Øµ Ø¨Ø§Ù„ÙØ±Ø¹
    if request.method == "POST" and "expense_name" in request.form:
        expense_name = request.form["expense_name"]
        amount = float(request.form.get("amount") or 0)
        file = request.files.get("file")
        filename = None
        if file and file.filename:
            filename = secure_filename(file.filename)
            file.save(os.path.join(app.config["UPLOAD_FOLDER"], filename))
        e = Expense(
            description=expense_name,
            amount=amount,
            file=filename,
            branch_id=user.branch_id
        )
        db.session.add(e)
        db.session.commit()
        flash("âœ… ØªÙ… ØªØ³Ø¬ÙŠÙ„ Ø§Ù„Ù…ØµØ±ÙˆÙ", "success")
        return redirect(url_for("finance_dashboard"))

    # âœ… ÙÙ‚Ø· Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø§Øª ØºÙŠØ± Ø§Ù„Ù…Ø¯ÙÙˆØ¹Ø© Ù„Ù‡Ø°Ø§ Ø§Ù„ÙØ±Ø¹
    unpaid_transactions = Transaction.query.filter_by(
        payment_status="ØºÙŠØ± Ù…Ø¯ÙÙˆØ¹Ø©",
        branch_id=user.branch_id
    ).all()

    # âœ… Ø§Ù„Ù…Ø¯ÙÙˆØ¹Ø§Øª Ø§Ù„Ø®Ø§ØµØ© Ø¨Ù‡Ø°Ø§ Ø§Ù„ÙØ±Ø¹ (Ø³ÙˆØ§Ø¡ Ù…Ø±ØªØ¨Ø·Ø© Ø¨Ù…Ø¹Ø§Ù…Ù„Ø© Ø§Ù„ÙØ±Ø¹ Ø£Ùˆ ØºÙŠØ± Ù…Ø±ØªØ¨Ø·Ø© ÙˆÙ„ÙƒÙ† Ù…Ù†Ø³ÙˆØ¨Ø© Ù„Ù„ÙØ±Ø¹)
    paid_transactions = Payment.query.outerjoin(Transaction, Payment.transaction_id == Transaction.id) \
        .filter(or_(Transaction.branch_id == user.branch_id, Payment.branch_id == user.branch_id)) \
        .order_by(Payment.id.desc()).all()

    # âœ… Ù…Ø¬Ù…ÙˆØ¹ Ø§Ù„Ø¯Ø®Ù„ Ù„Ù„ÙØ±Ø¹ ÙÙ‚Ø· (ÙŠØ´Ù…Ù„ Ø§Ù„Ù…Ø¯ÙÙˆØ¹Ø§Øª ØºÙŠØ± Ø§Ù„Ù…Ø±ØªØ¨Ø·Ø© Ø¨Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø§Øª ÙˆØ§Ù„Ù…Ù†Ø³ÙˆØ¨Ø© Ù„Ù„ÙØ±Ø¹)
    total_income = db.session.query(func.coalesce(func.sum(Payment.amount), 0.0)) \
        .select_from(Payment) \
        .outerjoin(Transaction, Payment.transaction_id == Transaction.id) \
        .filter(or_(Transaction.branch_id == user.branch_id, Payment.branch_id == user.branch_id)).scalar() or 0.0

    # âœ… Ù…ØµØ§Ø±ÙŠÙ Ø§Ù„ÙØ±Ø¹ ÙÙ‚Ø·
    expenses = Expense.query.filter_by(branch_id=user.branch_id).order_by(Expense.id.desc()).all()

    total_expenses = db.session.query(func.coalesce(func.sum(Expense.amount), 0.0))\
        .filter(Expense.branch_id == user.branch_id).scalar() or 0.0

    net_profit = total_income - total_expenses

    # âœ… Ø¨ÙŠØ§Ù†Ø§Øª Ø§Ù„Ø¹Ø±ÙˆØ¶ ÙˆØ§Ù„ÙÙˆØ§ØªÙŠØ± Ù„Ù„Ø¨Ù†ÙˆÙƒ Ù„Ø¹Ø±Ø¶Ù‡Ø§ Ø¨Ø³Ø±Ø¹Ø© ÙÙŠ Ù„ÙˆØ­Ø© Ø§Ù„Ù…Ø§Ù„ÙŠØ©
    banks = Bank.query.order_by(Bank.name.asc()).all()
    # Ø§Ø³ØªØ¨Ø¹Ø§Ø¯ Ø§Ù„Ø¹Ø±ÙˆØ¶ Ø§Ù„Ù…Ù†ØªÙ‡ÙŠØ© Ù…Ù† Ø§Ù„Ø¹Ø±Ø¶ Ø§Ù„Ø³Ø±ÙŠØ¹
    now_utc = datetime.utcnow()
    recent_quotes = Quote.query \
        .filter(or_(Quote.valid_until == None, Quote.valid_until >= now_utc)) \
        .order_by(Quote.id.desc()).limit(20).all()
    recent_bank_invoices = BankInvoice.query \
        .filter(BankInvoice.received_at == None) \
        .order_by(BankInvoice.id.desc()).limit(20).all()
    recent_customer_quotes = CustomerQuote.query \
        .filter(or_(CustomerQuote.valid_until == None, CustomerQuote.valid_until >= now_utc)) \
        .order_by(CustomerQuote.id.desc()).limit(20).all()
    recent_customer_invoices = CustomerInvoice.query.order_by(CustomerInvoice.id.desc()).limit(20).all()

    return render_template(
        "finance.html",
        transactions=unpaid_transactions,
        expenses=expenses,
        total_income=total_income,
        total_expenses=total_expenses,
        net_profit=net_profit,
        banks=banks,
        recent_quotes=recent_quotes,
        recent_bank_invoices=recent_bank_invoices,
        recent_customer_quotes=recent_customer_quotes,
        recent_customer_invoices=recent_customer_invoices,
        vat_default_percent=int(_get_vat_rate() * 100)
    )

# ---------------- ØµÙØ­Ø© Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø§Øª Ø§Ù„Ù…Ø¯ÙÙˆØ¹Ø© (Ù…Ø§Ù„ÙŠØ©) ----------------
@app.route("/finance/paid")
def finance_paid():
    if session.get("role") != "finance":
        return redirect(url_for("login"))

    user = User.query.get(session["user_id"])

    # Ù…Ø¹Ø§Ù…Ù„Ø§Øª Ù‡Ø°Ø§ Ø§Ù„ÙØ±Ø¹ Ø§Ù„ØªÙŠ Ù„Ø¯ÙŠÙ‡Ø§ Ù…Ø¯ÙÙˆØ¹Ø§Øª + Ø§Ù„Ø¯ÙØ¹Ø§Øª Ø§Ù„Ù…Ù†Ø³ÙˆØ¨Ø© Ù„Ù„ÙØ±Ø¹ Ø¨Ø¯ÙˆÙ† Ù…Ø¹Ø§Ù…Ù„Ø©
    payments = Payment.query.outerjoin(Transaction, Payment.transaction_id == Transaction.id) \
        .filter(or_(Transaction.branch_id == user.branch_id, Payment.branch_id == user.branch_id)) \
        .order_by(Payment.id.desc()).all()

    total_income = db.session.query(func.coalesce(func.sum(Payment.amount), 0.0)) \
        .select_from(Payment) \
        .outerjoin(Transaction, Payment.transaction_id == Transaction.id) \
        .filter(or_(Transaction.branch_id == user.branch_id, Payment.branch_id == user.branch_id)).scalar() or 0.0

    # ÙÙˆØ§ØªÙŠØ± Ø§Ù„Ø¨Ù†Ùƒ Ø§Ù„ØªÙŠ ØªÙ… Ø§Ø³ØªÙ„Ø§Ù… Ù…Ø¨Ù„ØºÙ‡Ø§
    received_bank_invoices = BankInvoice.query \
        .filter(BankInvoice.received_at != None) \
        .order_by(BankInvoice.received_at.desc()).all()

    return render_template("finance_paid.html", payments=payments, total_income=total_income, received_bank_invoices=received_bank_invoices)

# ---------------- Ø¥Ø¯Ø§Ø±Ø© Ù‚ÙˆØ§Ù„Ø¨ ÙˆÙˆØ±Ø¯ (Ù…Ø§Ù„ÙŠØ©) ----------------
@app.route("/finance/templates")
def finance_templates():
    if session.get("role") not in ["finance", "manager"]:
        return redirect(url_for("login"))

    # ğŸ†• Ø¹Ø±Ø¶ Ø§Ù„Ù‚Ø§Ù„Ø¨ Ø§Ù„Ø­Ø§Ù„ÙŠ Ù„Ù„ÙØ±Ø¹ Ø§Ù„Ø­Ø§Ù„ÙŠ Ù„Ù…ÙˆØ¸Ù Ø§Ù„Ù…Ø§Ù„ÙŠØ© + Ø§Ù„Ø¹Ø§Ù…
    user = User.query.get(session.get("user_id"))
    current_branch_id = getattr(user, "branch_id", None)
    templates = {
        "invoice": get_template_filename("invoice", current_branch_id) or get_template_filename("invoice", None),
        "quote": get_template_filename("quote", current_branch_id) or get_template_filename("quote", None),
    }
    branches = Branch.query.order_by(Branch.name.asc()).all()
    return render_template("finance_templates.html", templates=templates, branches=branches, current_branch_id=current_branch_id)

def _replace_placeholders_in_xml_bytes(xml_bytes: bytes, mapping: dict) -> bytes:
    try:
        text = xml_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = xml_bytes.decode("latin-1")
    for key, value in mapping.items():
        placeholder = "{" + str(key) + "}"
        replacement = str(value)
        if placeholder in text:
            text = text.replace(placeholder, replacement)
    return text.encode("utf-8")


def _fill_docx_from_template_xml(template_path: str, out_path: str, mapping: dict) -> None:
    import zipfile
    with zipfile.ZipFile(template_path, "r") as zin:
        with zipfile.ZipFile(out_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename.startswith("word/") and info.filename.lower().endswith(".xml"):
                    data = _replace_placeholders_in_xml_bytes(data, mapping)
                zout.writestr(info, data)


def _set_paragraph_rtl(paragraph, rtl: bool = True) -> None:
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1' if rtl else '0')
        pPr.append(bidi)
    except Exception:
        pass


def _generate_default_docx(doc_type: str, placeholders: dict, out_path: str) -> None:
    # ÙŠÙ†Ø´Ø¦ Ù…Ù„Ù DOCX Ø§ÙØªØ±Ø§Ø¶ÙŠ Ø¹Ø±Ø¨ÙŠ Ù…Ù†Ø³Ù‚ ÙƒØ¬Ø¯ÙˆÙ„ Ù„ÙØ§ØªÙˆØ±Ø©/Ø¹Ø±Ø¶ Ø³Ø¹Ø±
    from docx import Document as DocxDocument
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    document = DocxDocument()

    # ØªØ±ÙˆÙŠØ³Ø©
    header_p = document.add_paragraph()
    header_text = "invoice" if doc_type == "invoice" else "Ø¹Ø±Ø¶ Ø³Ø¹Ø±"
    run = header_p.add_run(header_text)
    run.bold = True
    try:
        run.font.size = Pt(16)
        run.font.name = "Arial"
    except Exception:
        pass
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_rtl(header_p, True)

    document.add_paragraph().add_run(" ")

    # Ù…Ù„Ø®Øµ Ø£Ø³Ø§Ø³ÙŠ
    meta_pairs = [
        ("Ø±Ù‚Ù… Ø§Ù„Ø¹Ù…Ù„ÙŠØ©", placeholders.get("TRANSACTION_ID", "")),
        ("Ø§Ù„ØªØ§Ø±ÙŠØ®", placeholders.get("DATE", "")),
        ("Ø§Ù„Ø¹Ù…ÙŠÙ„", placeholders.get("CLIENT_NAME", placeholders.get("NAME", ""))),
        ("Ø§Ù„Ù…ÙˆØ¸Ù", placeholders.get("EMPLOYEE", "")),
        ("Ø§Ù„Ø¨Ù†Ùƒ", placeholders.get("BANK_NAME", "")),
        ("ÙØ±Ø¹ Ø§Ù„Ø¨Ù†Ùƒ", placeholders.get("BANK_BRANCH", "")),
    ]

    table = document.add_table(rows=0, cols=2)
    try:
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    except Exception:
        pass

    for label, value in meta_pairs:
        row_cells = table.add_row().cells
        lc = row_cells[0].paragraphs[0]
        lr = lc.add_run(str(label))
        lr.bold = True
        try:
            lr.font.name = "Arial"; lr.font.size = Pt(11)
        except Exception:
            pass
        _set_paragraph_rtl(lc, True)
        lc.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        rc = row_cells[1].paragraphs[0]
        rr = rc.add_run(str(value))
        try:
            rr.font.name = "Arial"; rr.font.size = Pt(11)
        except Exception:
            pass
        _set_paragraph_rtl(rc, True)
        rc.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_paragraph().add_run(" ")

    # Ø¬Ø¯ÙˆÙ„ Ø§Ù„Ù…Ø¨Ù„Øº ÙˆØ§Ù„Ø¶Ø±ÙŠØ¨Ø© ÙˆØ§Ù„Ø¥Ø¬Ù…Ø§Ù„ÙŠ
    amounts = [
        ("Ø§Ù„Ø³Ø¹Ø± Ù‚Ø¨Ù„ Ø§Ù„Ø¶Ø±ÙŠØ¨Ø©", placeholders.get("PRICE", placeholders.get("AMOUNT", "0.00"))),
        ("Ø§Ù„Ø¶Ø±ÙŠØ¨Ø©", placeholders.get("TAX", "0.00")),
        ("Ø§Ù„Ø¥Ø¬Ù…Ø§Ù„ÙŠ Ø¨Ø¹Ø¯ Ø§Ù„Ø¶Ø±ÙŠØ¨Ø©", placeholders.get("TOTAL_PRICE", placeholders.get("TOTAL", "0.00"))),
    ]

    amt_table = document.add_table(rows=1, cols=3)
    try:
        amt_table.style = 'Table Grid'
        amt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    except Exception:
        pass

    hdr_cells = amt_table.rows[0].cells
    headers = ["Ø§Ù„Ø¨Ù†Ø¯", "Ø§Ù„Ù‚ÙŠÙ…Ø©", "Ø§Ù„Ø¹Ù…Ù„Ø©"]
    for idx, text in enumerate(headers):
        p = hdr_cells[idx].paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        try:
            r.font.name = "Arial"; r.font.size = Pt(11)
        except Exception:
            pass
        _set_paragraph_rtl(p, True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for label, value in amounts:
        row = amt_table.add_row().cells
        p0 = row[0].paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        _set_paragraph_rtl(p0, True)
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p1 = row[1].paragraphs[0]
        p1.add_run(str(value))
        _set_paragraph_rtl(p1, True)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p2 = row[2].paragraphs[0]
        p2.add_run("Ø±ÙŠØ§Ù„")
        _set_paragraph_rtl(p2, True)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph().add_run(" ")

    # ØªÙØ§ØµÙŠÙ„ Ø¥Ø¶Ø§ÙÙŠØ©
    details_title = document.add_paragraph()
    dr = details_title.add_run("Ø§Ù„ØªÙØ§ØµÙŠÙ„")
    dr.bold = True
    _set_paragraph_rtl(details_title, True)
    details_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    details_p = document.add_paragraph()
    details_p.add_run(placeholders.get("DETAILS", ""))
    _set_paragraph_rtl(details_p, True)
    details_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Ø£Ø±Ù‚Ø§Ù… Ø§Ù„Ù…Ø³ØªÙ†Ø¯
    ref_p = document.add_paragraph()
    ref_text = placeholders.get("INVOICE_NO") or placeholders.get("QUOTE_NO") or placeholders.get("QUTE_NO")
    if ref_text:
        ref_run = ref_p.add_run(f"Ø§Ù„Ù…Ø±Ø¬Ø¹: {ref_text}")
        ref_run.bold = True
    _set_paragraph_rtl(ref_p, True)
    ref_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Ø­ÙØ¸
    document.save(out_path)
def _render_docx_from_template(doc_type: str, placeholders: dict, out_name: str, branch_id: int | None = None):
    template_filename = get_template_filename(doc_type, branch_id)
    output_path = os.path.join(app.config["UPLOAD_FOLDER"], out_name)
    if not template_filename:
        # Ù„Ø§ ÙŠÙˆØ¬Ø¯ Ù‚Ø§Ù„Ø¨ Ù…Ø±ÙÙˆØ¹: Ø£Ù†Ø´Ø¦ Ù…Ù„Ù DOCX Ø§ÙØªØ±Ø§Ø¶ÙŠ Ø¹Ø±Ø¨ÙŠ Ù…Ù†Ø³Ù‚
        try:
            _generate_default_docx(doc_type, placeholders, output_path)
            return send_file(output_path, as_attachment=True, download_name=out_name)
        except Exception:
            flash("âš ï¸ ØªØ¹Ø°Ø± Ø¥Ù†Ø´Ø§Ø¡ Ø§Ù„Ù‚Ø§Ù„Ø¨ Ø§Ù„Ø§ÙØªØ±Ø§Ø¶ÙŠ", "warning")
            return redirect(url_for("finance_templates"))
    else:
        path = os.path.join(app.config["UPLOAD_FOLDER"], template_filename)
        _fill_docx_from_template_xml(path, output_path, placeholders)
        return send_file(output_path, as_attachment=True, download_name=out_name)


def _get_vat_rate() -> float:
    # ÙŠÙ…ÙƒÙ† Ø¶Ø¨Ø·Ù‡Ø§ Ø¹Ø¨Ø± Ù…ØªØºÙŠØ± Ø§Ù„Ø¨ÙŠØ¦Ø© VAT_RATE ÙƒÙ‚ÙŠÙ…Ø© Ø¹Ø´Ø±ÙŠØ© (Ù…Ø«Ø§Ù„ 0.05)
    try:
        return float(os.environ.get("VAT_RATE", "0.05"))
    except Exception:
        return 0.05


def _compute_tax_and_total(base_amount: float) -> tuple[float, float]:
    vat = _get_vat_rate()
    tax = round((base_amount or 0.0) * vat, 2)
    total = round((base_amount or 0.0) + tax, 2)
    return tax, total


def _sanitize_description(raw_text: str, transaction: Transaction | None = None) -> str:
    """Ø¥Ø²Ø§Ù„Ø© Ø£ÙŠ Ø°ÙƒØ± Ù„Ø­Ø§Ù„Ø© Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ù…Ù† Ø§Ù„ÙˆØµÙ/Ø§Ù„Ù…Ù„Ø§Ø­Ø¸Ø§Øª Ù‚Ø¨Ù„ Ø§Ù„Ø·Ø¨Ø§Ø¹Ø©.

    - ÙŠØ­Ø°Ù ØµØ±Ø§Ø­Ø© Ù‚ÙŠÙ…Ø© Ø­Ø§Ù„Ø© Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ø¥Ù† ØªÙˆÙØ±Øª
    - ÙŠØ­Ø°Ù Ø§Ù„Ù…Ù‚Ø§Ø·Ø¹ Ø§Ù„ØªÙŠ ØªØ¨Ø¯Ø£ Ø¨Ù€ "Ø§Ù„Ø­Ø§Ù„Ø©:" Ø£Ùˆ "Ø­Ø§Ù„Ø© Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø©:" Ø­ØªÙ‰ Ù†Ù‡Ø§ÙŠØ© Ø§Ù„Ø³Ø·Ø±
    - ÙŠØ­Ø°Ù Ø£Ø´Ù‡Ø± Ø§Ù„Ø¹Ø¨Ø§Ø±Ø§Øª Ø§Ù„Ù…Ø³ØªØ®Ø¯Ù…Ø© ÙƒØ­Ø§Ù„Ø©
    - ÙŠÙ†Ø¸Ù Ø§Ù„ÙÙˆØ§ØµÙ„ Ø§Ù„Ø²Ø§Ø¦Ø¯Ø© ÙÙŠ Ø§Ù„Ù†Ù‡Ø§ÙŠØ©
    """
    if not raw_text:
        return ""
    try:
        text = str(raw_text)
        # Ø¥Ø²Ø§Ù„Ø© Ù‚ÙŠÙ…Ø© Ø­Ø§Ù„Ø© Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ø¥Ù† ÙˆÙØ¬Ø¯Øª
        if transaction is not None and getattr(transaction, "status", None):
            status_value = str(transaction.status or "").strip()
            if status_value:
                text = text.replace(status_value, "")

        # Ø¥Ø²Ø§Ù„Ø© Ø§Ù„Ø£Ù†Ù…Ø§Ø· Ø§Ù„Ù†ØµÙŠØ© Ø§Ù„Ø´Ø§Ø¦Ø¹Ø© Ù„Ù„Ø­Ø§Ù„Ø©
        patterns = [
            r"\bØ§Ù„Ø­Ø§Ù„Ø©\s*(?:Ø§Ù„Ø­Ø§Ù„ÙŠØ©)?\s*[:ï¼š]\s*.*$",
            r"\bØ­Ø§Ù„Ø©\s*(?:Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø©|Ø§Ù„ÙØ§ØªÙˆØ±Ø©|Ø§Ù„Ø·Ù„Ø¨)?\s*[:ï¼š]\s*.*$",
        ]
        for pat in patterns:
            text = re.sub(pat, "", text).strip()

        # Ø¥Ø²Ø§Ù„Ø© Ø£Ø´Ù‡Ø± Ù‚ÙŠÙ… Ø§Ù„Ø­Ø§Ù„Ø© Ø§Ù„Ù…Ø¹Ø±ÙˆÙØ© ÙÙŠ Ø§Ù„Ù†Ø¸Ø§Ù…
        known_statuses = (
            "Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ù…Ù‡Ù†Ø¯Ø³",
            "Ù‚ÙŠØ¯ Ø§Ù„Ù…Ø¹Ø§ÙŠÙ†Ø©",
            "Ù‚ÙŠØ¯ Ø§Ù„ØªÙ†ÙÙŠØ°",
            "Ù…Ø±ÙÙˆØ¶Ø©",
            "Ø¨Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ø¯ÙØ¹",
            "Ù…ÙƒØªÙ…Ù„Ø©",
            "Ù…Ù†Ø¬Ø²Ø©",
            "ğŸ“‘ ØªÙ‚Ø±ÙŠØ± Ù…Ø±ÙÙˆØ¹",
            "ğŸ“‘ ØªÙ‚Ø±ÙŠØ± Ù…Ø¨Ø¯Ø¦ÙŠ",
        )
        for s in known_statuses:
            text = text.replace(s, "")

        # ØªÙ†Ø¸ÙŠÙ ÙÙˆØ§ØµÙ„/Ø±Ù…ÙˆØ² Ø²Ø§Ø¦Ø¯Ø© ÙÙŠ Ù†Ù‡Ø§ÙŠØ© Ø§Ù„Ù†Øµ
        text = re.sub(r"[\-\|\(\)\[\]Â·â€¢ØŒØŒ:,\s]+$", "", text).strip()
        return text
    except Exception:
        return str(raw_text)

@app.route("/finance/templates/quote/<int:transaction_id>")
def download_quote_doc(transaction_id: int):
    if session.get("role") != "finance":
        return redirect(url_for("login"))
    t = Transaction.query.get_or_404(transaction_id)
    bank_name = None
    if t.bank_id:
        bank = Bank.query.get(t.bank_id)
        bank_name = bank.name if bank else None
    amount = float(t.fee or 0)
    # ØªØ®ØµÙŠØµ Ø§Ù„ÙˆØµÙ ÙˆØ§Ù„Ø¶Ø±ÙŠØ¨Ø©
    details_override = (request.args.get("details") or "").strip()
    apply_vat = (request.args.get("apply_vat") or "1") == "1"
    vat_percent = request.args.get("vat")
    if vat_percent is not None:
        try:
            os.environ["VAT_RATE"] = str(float(vat_percent) / 100.0)
        except Exception:
            pass
    tax, total_with_tax = _compute_tax_and_total(amount) if apply_vat else (0.0, amount)
    placeholders = {
        "NAME": t.client or "",
        "CLIENT_NAME": t.client or "",
        "AMOUNT": f"{amount:.2f}",
        "PRICE": f"{amount:.2f}",
        "TAX": f"{tax:.2f}",
        "TOTAL_PRICE": f"{total_with_tax:.2f}",
        # Ù„Ù„Ø­ÙØ§Ø¸ Ø¹Ù„Ù‰ Ø§Ù„ØªÙˆØ§ÙÙ‚ Ù…Ø¹ Ø§Ù„Ù‚ÙˆØ§Ù„Ø¨ Ø§Ù„Ù‚Ø¯ÙŠÙ…Ø©
        "TOTAL": f"{amount:.2f}",
        "DATE": datetime.utcnow().strftime("%Y-%m-%d"),
        "DETAILS": details_override or "Ø±Ø³ÙˆÙ… Ø§Ù„ØªØ«Ù…ÙŠÙ†",
        "QUOTE_NO": f"QUOTE-{t.id}",
        "QUTE_NO": f"QUOTE-{t.id}",
        "TRANSACTION_ID": str(t.id),
        "EMPLOYEE": t.employee or "",
        "BANK_NAME": bank_name or "",
        "BANK_BRANCH": t.bank_branch or "",
        "STATE": t.state or "",
        "REGION": t.region or "",
        "AREA": str(t.area or 0),
        "BUILDING_AREA": str(t.building_area or 0),
        "BUILDING_AGE": str(t.building_age or 0),
        "LAND_VALUE": f"{float(t.land_value or 0):.2f}",
        "BUILDING_VALUE": f"{float(t.building_value or 0):.2f}",
        "TOTAL_ESTIMATE": f"{float(t.total_estimate or 0):.2f}",
    }
    out_name = f"quote_{t.id}.docx"
    return _render_docx_from_template(
        "quote",
        placeholders,
        out_name,
        branch_id=t.branch_id,
    )

@app.route("/finance/templates/invoice/<int:transaction_id>")
def download_invoice_doc(transaction_id: int):
    if session.get("role") != "finance":
        return redirect(url_for("login"))
    t = Transaction.query.get_or_404(transaction_id)
    bank_name = None
    if t.bank_id:
        bank = Bank.query.get(t.bank_id)
        bank_name = bank.name if bank else None
    amount = float(t.fee or 0)
    # ØªØ®ØµÙŠØµ Ø§Ù„ÙˆØµÙ ÙˆØ§Ù„Ø¶Ø±ÙŠØ¨Ø©
    details_override = (request.args.get("details") or "").strip()
    apply_vat = (request.args.get("apply_vat") or "1") == "1"
    vat_percent = request.args.get("vat")
    if vat_percent is not None:
        try:
            os.environ["VAT_RATE"] = str(float(vat_percent) / 100.0)
        except Exception:
            pass
    tax, total_with_tax = _compute_tax_and_total(amount) if apply_vat else (0.0, amount)
    placeholders = {
        "NAME": t.client or "",
        "CLIENT_NAME": t.client or "",
        "AMOUNT": f"{amount:.2f}",
        "PRICE": f"{amount:.2f}",
        "TAX": f"{tax:.2f}",
        "TOTAL_PRICE": f"{total_with_tax:.2f}",
        # Ù„Ù„ØªÙˆØ§ÙÙ‚ Ù…Ø¹ Ø§Ù„Ù‚ÙˆØ§Ù„Ø¨ Ø§Ù„Ù‚Ø¯ÙŠÙ…Ø©
        "TOTAL": f"{amount:.2f}",
        "DATE": datetime.utcnow().strftime("%Y-%m-%d"),
        "DETAILS": _sanitize_description(details_override or "Ø±Ø³ÙˆÙ… Ø§Ù„ØªØ«Ù…ÙŠÙ†", t),
        "INVOICE_NO": f"INV-{t.id}",
        "TRANSACTION_ID": str(t.id),
        "EMPLOYEE": t.employee or "",
        "BANK_NAME": bank_name or "",
        "BANK_BRANCH": t.bank_branch or "",
        "STATE": t.state or "",
        "REGION": t.region or "",
        "AREA": str(t.area or 0),
        "BUILDING_AREA": str(t.building_area or 0),
        "BUILDING_AGE": str(t.building_age or 0),
        "LAND_VALUE": f"{float(t.land_value or 0):.2f}",
        "BUILDING_VALUE": f"{float(t.building_value or 0):.2f}",
        "TOTAL_ESTIMATE": f"{float(t.total_estimate or 0):.2f}",
    }
    out_name = f"invoice_{t.id}.docx"
    return _render_docx_from_template(
        "invoice",
        placeholders,
        out_name,
        branch_id=t.branch_id,
    )

# âœ… Ø·Ø¨Ø§Ø¹Ø© ÙØ§ØªÙˆØ±Ø© HTML Ø§Ø­ØªØ±Ø§ÙÙŠØ© Ù„Ù„Ù…Ø¹Ø§Ù…Ù„Ø©
@app.route("/finance/print/invoice/<int:transaction_id>")
def print_invoice_html(transaction_id: int):
    if session.get("role") != "finance":
        return redirect(url_for("login"))
    t = Transaction.query.get_or_404(transaction_id)
    bank_name = None
    if t.bank_id:
        bank = Bank.query.get(t.bank_id)
        bank_name = bank.name if bank else None

    amount = float(t.fee or 0)
    tax, total_with_tax = _compute_tax_and_total(amount)
    details_override = (request.args.get("details") or "").strip()

    # Ù…Ø¹Ù„ÙˆÙ…Ø§Øª Ø§Ù„Ù…Ø¤Ø³Ø³Ø© Ø§Ù„Ø§ÙØªØ±Ø§Ø¶ÙŠØ© (ÙŠÙ…ÙƒÙ† Ù„Ø§Ø­Ù‚Ù‹Ø§ Ø±Ø¨Ø·Ù‡Ø§ Ù…Ù† Ø§Ù„Ø¥Ø¹Ø¯Ø§Ø¯Ø§Øª/Ø§Ù„ÙØ±Ø¹)
    org_name = "Ø´Ø±ÙƒØ© Ø§Ù„ØªØ«Ù…ÙŠÙ†"
    org_meta = "Ø§Ù„Ø¹Ù†ÙˆØ§Ù† Â· Ø§Ù„Ù‡Ø§ØªÙ Â· Ø§Ù„Ø¨Ø±ÙŠØ¯ Ø§Ù„Ø¥Ù„ÙƒØªØ±ÙˆÙ†ÙŠ"

    return render_template(
        "print_invoice.html",
        transaction=t,
        bank_name=bank_name,
        amount=amount,
        tax=tax,
        total_with_tax=total_with_tax,
        vat_rate=_get_vat_rate(),
        date_str=(datetime.utcnow().strftime("%Y-%m-%d")),
        org_name=org_name,
        org_meta=org_meta,
        notes=_sanitize_description(details_override, t),
    )

# âœ… Ø·Ø¨Ø§Ø¹Ø© ÙØ§ØªÙˆØ±Ø© Ø¨Ù†Ùƒ HTML Ø¨Ù†ÙØ³ ØªØµÙ…ÙŠÙ… Ø§Ù„Ø·Ø¨Ø§Ø¹Ø©
@app.route("/finance/print/bank_invoice/<int:invoice_id>")
def print_bank_invoice_html(invoice_id: int):
    if session.get("role") not in ["finance", "manager"]:
        return redirect(url_for("login"))

    inv = BankInvoice.query.get_or_404(invoice_id)
    bank = Bank.query.get(inv.bank_id) if inv.bank_id else None
    transaction = Transaction.query.get(inv.transaction_id) if inv.transaction_id else None

    bank_name = bank.name if bank else None
    amount = float(inv.amount or 0)
    details_override = (request.args.get("details") or "").strip()
    # Ø¶Ø±ÙŠØ¨Ø© Ø§Ø®ØªÙŠØ§Ø±ÙŠØ©
    apply_vat = (request.args.get("apply_vat") or "1") == "1"
    vat_percent = request.args.get("vat")
    if vat_percent is not None:
        try:
            os.environ["VAT_RATE"] = str(float(vat_percent) / 100.0)
        except Exception:
            pass
    tax, total_with_tax = _compute_tax_and_total(amount) if apply_vat else (0.0, amount)

    org_name = "Ø´Ø±ÙƒØ© Ø§Ù„ØªØ«Ù…ÙŠÙ†"
    org_meta = "Ø§Ù„Ø¹Ù†ÙˆØ§Ù† Â· Ø§Ù„Ù‡Ø§ØªÙ Â· Ø§Ù„Ø¨Ø±ÙŠØ¯ Ø§Ù„Ø¥Ù„ÙƒØªØ±ÙˆÙ†ÙŠ"

    return render_template(
        "print_invoice.html",
        # unified template
        transaction=transaction,
        bank_name=bank_name,
        amount=amount,
        tax=tax,
        total_with_tax=total_with_tax,
        vat_rate=_get_vat_rate(),
        date_str=(inv.issued_at or datetime.utcnow()).strftime("%Y-%m-%d"),
        org_name=org_name,
        org_meta=org_meta,
        notes=_sanitize_description(details_override or (inv.note or ""), transaction),
        # metadata for header
        badge_label="ÙØ§ØªÙˆØ±Ø©",
        invoice_code=(inv.invoice_number or f"INV-BANK-{inv.id}"),
        reference_code="INV-BANK",
        client_name=(transaction.client if transaction else (bank_name or "")),
        employee_name=(transaction.employee if transaction else "-"),
    )

# âœ… Ø·Ø¨Ø§Ø¹Ø© ÙØ§ØªÙˆØ±Ø© Ø¹Ù…ÙŠÙ„ HTML Ø¨Ù†ÙØ³ ØªØµÙ…ÙŠÙ… Ø§Ù„Ø·Ø¨Ø§Ø¹Ø©
@app.route("/finance/print/customer_invoice/<int:invoice_id>")
def print_customer_invoice_html(invoice_id: int):
    if session.get("role") != "finance":
        return redirect(url_for("login"))

    inv = CustomerInvoice.query.get_or_404(invoice_id)
    transaction = Transaction.query.get(inv.transaction_id) if inv.transaction_id else None

    bank_name = None
    if transaction and transaction.bank_id:
        bank = Bank.query.get(transaction.bank_id)
        bank_name = bank.name if bank else None

    amount = float(inv.amount or 0)
    details_override = (request.args.get("details") or "").strip()
    # Ø¶Ø±ÙŠØ¨Ø© Ø§Ø®ØªÙŠØ§Ø±ÙŠØ©
    apply_vat = (request.args.get("apply_vat") or "1") == "1"
    vat_percent = request.args.get("vat")
    if vat_percent is not None:
        try:
            os.environ["VAT_RATE"] = str(float(vat_percent) / 100.0)
        except Exception:
            pass
    tax, total_with_tax = _compute_tax_and_total(amount) if apply_vat else (0.0, amount)

    org_name = "Ø´Ø±ÙƒØ© Ø§Ù„ØªØ«Ù…ÙŠÙ†"
    org_meta = "Ø§Ù„Ø¹Ù†ÙˆØ§Ù† Â· Ø§Ù„Ù‡Ø§ØªÙ Â· Ø§Ù„Ø¨Ø±ÙŠØ¯ Ø§Ù„Ø¥Ù„ÙƒØªØ±ÙˆÙ†ÙŠ"

    return render_template(
        "print_invoice.html",
        # unified template
        transaction=transaction,
        bank_name=bank_name,
        amount=amount,
        tax=tax,
        total_with_tax=total_with_tax,
        vat_rate=_get_vat_rate(),
        date_str=(inv.issued_at or datetime.utcnow()).strftime("%Y-%m-%d"),
        org_name=org_name,
        org_meta=org_meta,
        notes=_sanitize_description(details_override or (inv.note or ""), transaction),
        # metadata for header
        badge_label="ÙØ§ØªÙˆØ±Ø©",
        invoice_code=(inv.invoice_number or f"INV-CUST-{inv.id}"),
        reference_code="INV-CUST",
        client_name=(inv.customer_name or (transaction.client if transaction else "")),
        employee_name=(transaction.employee if transaction else "-"),
    )

# âœ… Ø·Ø¨Ø§Ø¹Ø© Ø¹Ø±Ø¶ Ø³Ø¹Ø± Ø¹Ù…ÙŠÙ„ HTML Ø¨Ù†ÙØ³ ØªØµÙ…ÙŠÙ… Ø§Ù„Ø·Ø¨Ø§Ø¹Ø©
@app.route("/finance/print/customer_quote/<int:quote_id>")
def print_customer_quote_html(quote_id: int):
    if session.get("role") != "finance":
        return redirect(url_for("login"))

    q = CustomerQuote.query.get_or_404(quote_id)
    transaction = Transaction.query.get(q.transaction_id) if q.transaction_id else None

    bank_name = None
    if transaction and transaction.bank_id:
        bank = Bank.query.get(transaction.bank_id)
        bank_name = bank.name if bank else None

    amount = float(q.amount or 0)
    tax, total_with_tax = _compute_tax_and_total(amount)

    org_name = "Ø´Ø±ÙƒØ© Ø§Ù„ØªØ«Ù…ÙŠÙ†"
    org_meta = "Ø§Ù„Ø¹Ù†ÙˆØ§Ù† Â· Ø§Ù„Ù‡Ø§ØªÙ Â· Ø§Ù„Ø¨Ø±ÙŠØ¯ Ø§Ù„Ø¥Ù„ÙƒØªØ±ÙˆÙ†ÙŠ"

    return render_template(
        "print_invoice.html",
        # unified template
        transaction=transaction,
        bank_name=bank_name,
        amount=amount,
        tax=tax,
        total_with_tax=total_with_tax,
        vat_rate=_get_vat_rate(),
        date_str=(q.valid_until or datetime.utcnow()).strftime("%Y-%m-%d"),
        org_name=org_name,
        org_meta=org_meta,
        notes=q.note or "",
        # metadata for header
        badge_label="Ø¹Ø±Ø¶ Ø³Ø¹Ø±",
        invoice_code=f"QUOTE-CUST-{q.id}",
        reference_code="QUOTE-CUST",
        client_name=(q.customer_name or (transaction.client if transaction else "")),
        employee_name=(transaction.employee if transaction else "-"),
    )

# âœ… ØªÙ†Ø²ÙŠÙ„ ÙØ§ØªÙˆØ±Ø© Ø¨Ù†Ùƒ (Ù…Ù† Ø¬Ø¯ÙˆÙ„ BankInvoice)
@app.route("/finance/download/bank_invoice/<int:invoice_id>")
def download_bank_invoice_doc(invoice_id: int):
    if session.get("role") not in ["finance", "manager"]:
        return redirect(url_for("login"))
    inv = BankInvoice.query.get_or_404(invoice_id)
    bank = Bank.query.get(inv.bank_id)
    # Ø§Ø®ØªÙŠØ§Ø± Ø§Ù„ÙØ±Ø¹: Ø¥Ù† ÙˆØ¬Ø¯Øª Ù…Ø¹Ø§Ù…Ù„Ø© Ù…Ø±ØªØ¨Ø·Ø© Ù†Ø³ØªØ®Ø¯Ù… ÙØ±Ø¹Ù‡Ø§ØŒ ÙˆØ¥Ù„Ø§ ÙØ±Ø¹ Ù…ÙˆØ¸Ù Ø§Ù„Ù…Ø§Ù„ÙŠØ©
    preferred_branch_id = None
    transaction = None
    if inv.transaction_id:
        transaction = Transaction.query.get(inv.transaction_id)
        preferred_branch_id = transaction.branch_id if transaction else None
    if preferred_branch_id is None:
        user = User.query.get(session.get("user_id"))
        preferred_branch_id = getattr(user, "branch_id", None)

    amount = float(inv.amount or 0)
    # Ø¶Ø±ÙŠØ¨Ø© Ø§Ø®ØªÙŠØ§Ø±ÙŠØ© Ø¹Ø¨Ø± Ø§Ù„Ø§Ø³ØªØ¹Ù„Ø§Ù…
    apply_vat = (request.args.get("apply_vat") or "1") == "1"
    vat_percent = request.args.get("vat")
    if vat_percent is not None:
        try:
            os.environ["VAT_RATE"] = str(float(vat_percent) / 100.0)
        except Exception:
            pass
    tax, total_with_tax = _compute_tax_and_total(amount) if apply_vat else (0.0, amount)
    placeholders = {
        "NAME": (bank.name if bank else f"Bank #{inv.bank_id}"),
        "CLIENT_NAME": (bank.name if bank else f"Bank #{inv.bank_id}"),
        "AMOUNT": f"{amount:.2f}",
        "PRICE": f"{amount:.2f}",
        "TAX": f"{tax:.2f}",
        "TOTAL_PRICE": f"{total_with_tax:.2f}",
        # ØªÙˆØ§ÙÙ‚ Ù‚Ø¯ÙŠÙ…
        "TOTAL": f"{amount:.2f}",
        "DATE": (inv.issued_at or datetime.utcnow()).strftime("%Y-%m-%d"),
        "DETAILS": _sanitize_description(inv.note or "", transaction),
        "INVOICE_NO": (inv.invoice_number or f"INV-BANK-{inv.id}"),
        "TRANSACTION_ID": str(inv.transaction_id or ""),
        "BANK_NAME": (bank.name if bank else ""),
    }
    # Ù„Ùˆ Ø¹Ù†Ø¯Ù†Ø§ Ù…Ø¹Ø§Ù…Ù„Ø©ØŒ Ù†Ø¶ÙŠÙ ØªÙØ§ØµÙŠÙ„ Ø¥Ø¶Ø§ÙÙŠØ©
    if transaction:
        placeholders.update({
            "CLIENT_NAME": transaction.client or placeholders.get("CLIENT_NAME", ""),
            "BANK_BRANCH": transaction.bank_branch or "",
            "EMPLOYEE": transaction.employee or "",
            "STATE": transaction.state or "",
            "REGION": transaction.region or "",
            "AREA": str(transaction.area or 0),
            "BUILDING_AREA": str(transaction.building_area or 0),
            "BUILDING_AGE": str(transaction.building_age or 0),
            "LAND_VALUE": f"{float(transaction.land_value or 0):.2f}",
            "BUILDING_VALUE": f"{float(transaction.building_value or 0):.2f}",
            "TOTAL_ESTIMATE": f"{float(transaction.total_estimate or 0):.2f}",
        })

    out_name = f"bank_invoice_{inv.id}.docx"
    return _render_docx_from_template(
        "invoice",
        placeholders,
        out_name,
        branch_id=preferred_branch_id,
    )

# âœ… ØªÙ†Ø²ÙŠÙ„ ÙØ§ØªÙˆØ±Ø© Ø¹Ù…ÙŠÙ„ (Ù…Ù† Ø¬Ø¯ÙˆÙ„ CustomerInvoice)
@app.route("/finance/download/customer_invoice/<int:invoice_id>")
def download_customer_invoice_doc(invoice_id: int):
    if session.get("role") != "finance":
        return redirect(url_for("login"))
    inv = CustomerInvoice.query.get_or_404(invoice_id)
    preferred_branch_id = None
    transaction = None
    if inv.transaction_id:
        transaction = Transaction.query.get(inv.transaction_id)
        preferred_branch_id = transaction.branch_id if transaction else None
    if preferred_branch_id is None:
        user = User.query.get(session.get("user_id"))
        preferred_branch_id = getattr(user, "branch_id", None)

    amount = float(inv.amount or 0)
    apply_vat = (request.args.get("apply_vat") or "1") == "1"
    vat_percent = request.args.get("vat")
    if vat_percent is not None:
        try:
            os.environ["VAT_RATE"] = str(float(vat_percent) / 100.0)
        except Exception:
            pass
    tax, total_with_tax = _compute_tax_and_total(amount) if apply_vat else (0.0, amount)
    placeholders = {
        "NAME": inv.customer_name or "",
        "CLIENT_NAME": inv.customer_name or "",
        "AMOUNT": f"{amount:.2f}",
        "PRICE": f"{amount:.2f}",
        "TAX": f"{tax:.2f}",
        "TOTAL_PRICE": f"{total_with_tax:.2f}",
        # ØªÙˆØ§ÙÙ‚ Ù‚Ø¯ÙŠÙ…
        "TOTAL": f"{amount:.2f}",
        "DATE": (inv.issued_at or datetime.utcnow()).strftime("%Y-%m-%d"),
        "DETAILS": _sanitize_description(inv.note or "", transaction),
        "INVOICE_NO": (inv.invoice_number or f"INV-CUST-{inv.id}"),
        "TRANSACTION_ID": str(inv.transaction_id or ""),
    }
    if transaction:
        bank_name = None
        if transaction.bank_id:
            bank = Bank.query.get(transaction.bank_id)
            bank_name = bank.name if bank else None
        placeholders.update({
            "BANK_NAME": bank_name or "",
            "BANK_BRANCH": transaction.bank_branch or "",
            "STATE": transaction.state or "",
            "REGION": transaction.region or "",
            "AREA": str(transaction.area or 0),
            "BUILDING_AREA": str(transaction.building_area or 0),
            "BUILDING_AGE": str(transaction.building_age or 0),
            "LAND_VALUE": f"{float(transaction.land_value or 0):.2f}",
            "BUILDING_VALUE": f"{float(transaction.building_value or 0):.2f}",
            "TOTAL_ESTIMATE": f"{float(transaction.total_estimate or 0):.2f}",
        })

    out_name = f"customer_invoice_{inv.id}.docx"
    return _render_docx_from_template(
        "invoice",
        placeholders,
        out_name,
        branch_id=preferred_branch_id,
    )

# âœ… ØªÙ†Ø²ÙŠÙ„ Ø¹Ø±Ø¶ Ø³Ø¹Ø± Ø¹Ù…ÙŠÙ„ (Ù…Ù† Ø¬Ø¯ÙˆÙ„ CustomerQuote)
@app.route("/finance/download/customer_quote/<int:quote_id>")
def download_customer_quote_doc(quote_id: int):
    if session.get("role") != "finance":
        return redirect(url_for("login"))
    q = CustomerQuote.query.get_or_404(quote_id)
    preferred_branch_id = None
    transaction = None
    if q.transaction_id:
        transaction = Transaction.query.get(q.transaction_id)
        preferred_branch_id = transaction.branch_id if transaction else None
    if preferred_branch_id is None:
        user = User.query.get(session.get("user_id"))
        preferred_branch_id = getattr(user, "branch_id", None)

    placeholders = {
        "NAME": q.customer_name or "",
        "CLIENT_NAME": q.customer_name or "",
        "AMOUNT": f"{float(q.amount or 0):.2f}",
        "PRICE": f"{float(q.amount or 0):.2f}",
        "TOTAL": f"{float(q.amount or 0):.2f}",
        "DATE": datetime.utcnow().strftime("%Y-%m-%d"),
        "DETAILS": q.note or "",
        "QUOTE_NO": f"QUOTE-CUST-{q.id}",
        "QUTE_NO": f"QUOTE-CUST-{q.id}",
        "TRANSACTION_ID": str(q.transaction_id or ""),
        "VALID_UNTIL": q.valid_until.strftime("%Y-%m-%d") if q.valid_until else "",
    }
    if transaction:
        bank_name = None
        if transaction.bank_id:
            bank = Bank.query.get(transaction.bank_id)
            bank_name = bank.name if bank else None
        placeholders.update({
            "BANK_NAME": bank_name or "",
            "BANK_BRANCH": transaction.bank_branch or "",
            "STATE": transaction.state or "",
            "REGION": transaction.region or "",
            "AREA": str(transaction.area or 0),
            "BUILDING_AREA": str(transaction.building_area or 0),
            "BUILDING_AGE": str(transaction.building_age or 0),
            "LAND_VALUE": f"{float(transaction.land_value or 0):.2f}",
            "BUILDING_VALUE": f"{float(transaction.building_value or 0):.2f}",
            "TOTAL_ESTIMATE": f"{float(transaction.total_estimate or 0):.2f}",
        })

    out_name = f"customer_quote_{q.id}.docx"
    return _render_docx_from_template(
        "quote",
        placeholders,
        out_name,
        branch_id=preferred_branch_id,
    )

# âœ… Ø¥Ø¶Ø§ÙØ© Ø¯ÙØ¹Ø© Ø¬Ø¯ÙŠØ¯Ø©
@app.route("/add_payment/<int:tid>", methods=["POST"])
def add_payment(tid):
    if session.get("role") != "finance":
        return redirect(url_for("login"))

    user = User.query.get(session["user_id"])
    transaction = Transaction.query.get_or_404(tid)

    # ğŸš¨ Ù…Ù†Ø¹ Ø§Ù„ØªÙ„Ø§Ø¹Ø¨: Ù„Ø§Ø²Ù… ØªÙƒÙˆÙ† Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ù„Ù†ÙØ³ ÙØ±Ø¹ Ù…ÙˆØ¸Ù Ø§Ù„Ù…Ø§Ù„ÙŠØ©
    if transaction.branch_id != user.branch_id:
        flash("â›” ØºÙŠØ± Ù…Ø³Ù…ÙˆØ­ ØªØ¹Ø¯ÙŠÙ„ Ù…Ø¹Ø§Ù…Ù„Ø§Øª Ù…Ù† ÙØ±Ø¹ Ø¢Ø®Ø±", "danger")
        return redirect(url_for("finance_dashboard"))

    amount = float(request.form.get("amount") or 0)
    if amount > 0:
        receipt = request.files.get("receipt_file")
        filename = None
        if receipt and receipt.filename:
            filename = secure_filename(receipt.filename)
            receipt.save(os.path.join(app.config["UPLOAD_FOLDER"], filename))
        payment = Payment(
            transaction_id=transaction.id,
            amount=amount,
            method=request.form.get("method"),
            receipt_file=filename,
            date_received=datetime.utcnow(),
            received_by=session.get("username")
        )
        db.session.add(payment)
        db.session.commit()

        total_paid = db.session.query(func.coalesce(func.sum(Payment.amount), 0.0))\
                               .filter_by(transaction_id=transaction.id).scalar() or 0.0
        transaction.payment_status = "Ù…Ø¯ÙÙˆØ¹Ø©" if total_paid >= transaction.fee else "ØºÙŠØ± Ù…Ø¯ÙÙˆØ¹Ø©"
        db.session.commit()
        flash("âœ… ØªÙ… ØªØ³Ø¬ÙŠÙ„ Ø§Ù„Ø¯ÙØ¹Ø© Ø¨Ù†Ø¬Ø§Ø­", "success")
    return redirect(url_for("finance_dashboard"))

# âœ… Ø¥Ù†Ø´Ø§Ø¡ Ø¹Ø±Ø¶ Ø³Ø¹Ø± Ù„Ù„Ø¨Ù†Ùƒ (Ù…Ù† Ø§Ù„Ù…Ø§Ù„ÙŠØ©)
@app.route("/finance/quotes", methods=["POST"])
def finance_create_quote():
    if session.get("role") != "finance":
        return redirect(url_for("login"))

    bank_id = int(request.form.get("bank_id"))
    amount = float(request.form.get("amount") or 0)
    valid_until_str = request.form.get("valid_until")
    note = request.form.get("note")
    transaction_id = request.form.get("transaction_id")

    valid_until_dt = None
    if valid_until_str:
        try:
            valid_until_dt = datetime.fromisoformat(valid_until_str)
        except Exception:
            valid_until_dt = None

    q = Quote(
        bank_id=bank_id,
        amount=amount,
        valid_until=valid_until_dt,
        note=note,
        transaction_id=int(transaction_id) if transaction_id else None,
        created_by=session.get("user_id"),
    )
    db.session.add(q)
    db.session.commit()
    flash("âœ… ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ Ø¹Ø±Ø¶ Ø§Ù„Ø³Ø¹Ø±", "success")
    return redirect(url_for("download_customer_quote_doc", quote_id=q.id))

# âœ… Ø¥Ù†Ø´Ø§Ø¡ ÙØ§ØªÙˆØ±Ø© Ø¨Ù†Ùƒ Ø¨Ù…Ø¨Ù„Øº Ù…Ø­Ø¯Ø¯ (Ù…Ø±Ø­Ù„Ø© Ø§Ù„Ø¥ØµØ¯Ø§Ø±)
@app.route("/finance/bank_invoices", methods=["POST"])
def finance_create_bank_invoice():
    if session.get("role") != "finance":
        return redirect(url_for("login"))

    bank_id = int(request.form.get("bank_id"))
    amount = float(request.form.get("amount") or 0)
    transaction_id = request.form.get("transaction_id")
    note = request.form.get("note")

    inv = BankInvoice(
        bank_id=bank_id,
        amount=amount,
        transaction_id=int(transaction_id) if transaction_id else None,
        note=note,
        issued_at=datetime.utcnow(),
    )
    db.session.add(inv)
    db.session.commit()

    # ØªÙˆÙ„ÙŠØ¯ Ø±Ù‚Ù… ÙØ§ØªÙˆØ±Ø© ÙØ±ÙŠØ¯ Ø¨Ø¹Ø¯ Ø¥Ù†Ø´Ø§Ø¡ Ø§Ù„Ø³Ø¬Ù„ ÙˆØ§Ù„Ø­ØµÙˆÙ„ Ø¹Ù„Ù‰ id
    try:
        inv.invoice_number = generate_unique_invoice_number(prefix="INV", kind="BANK")
        db.session.commit()
    except Exception:
        db.session.rollback()
        try:
            inv.invoice_number = generate_unique_invoice_number(prefix="INV")
            db.session.commit()
        except Exception:
            db.session.rollback()
    flash("âœ… ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ ÙØ§ØªÙˆØ±Ø© Ø§Ù„Ø¨Ù†Ùƒ", "success")
    return redirect(url_for("print_bank_invoice_html", invoice_id=inv.id) + "?auto=1")

# âœ… ØªØ­Ø¯ÙŠØ« Ø­Ø§Ù„Ø© ÙØ§ØªÙˆØ±Ø© Ø§Ù„Ø¨Ù†Ùƒ (ØªØ³Ù„ÙŠÙ… / Ø§Ø³ØªÙ„Ø§Ù…)
@app.route("/finance/bank_invoices/<int:invoice_id>/status", methods=["POST"])
def finance_update_bank_invoice_status(invoice_id: int):
    if session.get("role") != "finance":
        return redirect(url_for("login"))

    action = (request.form.get("action") or "").strip().lower()
    invoice = BankInvoice.query.get_or_404(invoice_id)

    if action == "deliver":
        invoice.delivered_at = datetime.utcnow()
        db.session.commit()
        flash("âœ… ØªÙ… ØªØ­Ø¯ÙŠØ« Ø­Ø§Ù„Ø© Ø§Ù„ÙØ§ØªÙˆØ±Ø©: ØªÙ… Ø§Ù„ØªØ³Ù„ÙŠÙ…", "success")
    elif action == "receive":
        # âœ… ÙŠØ¬Ø¨ Ø±ÙØ¹ Ø¥ÙŠØµØ§Ù„ Ø§Ù„Ø§Ø³ØªÙ„Ø§Ù… Ù…Ù† Ø§Ù„Ø¨Ù†Ùƒ
        receipt = request.files.get("receipt_file")
        if not receipt or not receipt.filename:
            flash("â›” Ø¥ÙŠØµØ§Ù„ Ø§Ù„Ø§Ø³ØªÙ„Ø§Ù… Ù…Ù† Ø§Ù„Ø¨Ù†Ùƒ Ù…Ø·Ù„ÙˆØ¨", "danger")
            return redirect(url_for("finance_dashboard"))

        # Ø­ÙØ¸ Ø§Ù„Ø¥ÙŠØµØ§Ù„
        filename = secure_filename(receipt.filename)
        receipt.save(os.path.join(app.config["UPLOAD_FOLDER"], filename))

        invoice.received_at = datetime.utcnow()
        db.session.commit()

        # Ø¹Ù†Ø¯ ØªØ£ÙƒÙŠØ¯ Ø§Ù„Ø§Ø³ØªÙ„Ø§Ù…ØŒ Ù†Ø³Ø¬Ù„ Ø§Ù„Ø¯Ø®Ù„ ÙƒÙ€ Payment Ù„ÙØ±Ø¹ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© (Ø¥Ù† ÙˆÙØ¬Ø¯Øª)
        created_income = False
        if invoice.transaction_id:
            t = Transaction.query.get(invoice.transaction_id)
            if t:
                # Ù†ØªØ¬Ù†Ø¨ Ø§Ù„ØªÙƒØ±Ø§Ø±: Ù†ØªØ­Ù‚Ù‚ Ù…Ù† ÙˆØ¬ÙˆØ¯ Ø¯ÙØ¹Ø© Ø¨Ù†ÙØ³ Ø§Ù„Ù…Ø¨Ù„Øº ÙˆØ§Ù„Ø·Ø±ÙŠÙ‚Ø© "Ø¨Ù†Ùƒ"
                existing_payment = Payment.query.filter_by(
                    transaction_id=t.id,
                    amount=invoice.amount,
                    method="Ø¨Ù†Ùƒ",
                ).first()
                if not existing_payment:
                    p = Payment(
                        transaction_id=t.id,
                        amount=invoice.amount,
                        method="Ø¨Ù†Ùƒ",
                        date_received=datetime.utcnow(),
                        received_by=session.get("username"),
                        branch_id=t.branch_id,
                        receipt_file=filename,
                    )
                    db.session.add(p)
                    db.session.commit()
                    created_income = True

                # âœ… Ø¨Ø¹Ø¯ ØªØ³Ø¬ÙŠÙ„/ØªØ£ÙƒÙŠØ¯ Ø§Ù„Ø¯ÙØ¹Ø©ØŒ Ù†Ø¹ÙŠØ¯ Ø§Ø­ØªØ³Ø§Ø¨ Ø­Ø§Ù„Ø© Ø§Ù„Ø¯ÙØ¹ Ù„Ù„Ù…Ø¹Ø§Ù…Ù„Ø©
                total_paid = db.session.query(func.coalesce(func.sum(Payment.amount), 0.0))\
                    .filter_by(transaction_id=t.id).scalar() or 0.0
                t.payment_status = "Ù…Ø¯ÙÙˆØ¹Ø©" if total_paid >= t.fee else "ØºÙŠØ± Ù…Ø¯ÙÙˆØ¹Ø©"
                db.session.commit()

        # Ù„Ø§ ØªÙˆØ¬Ø¯ Ù…Ø¹Ø§Ù…Ù„Ø© Ù…Ø±ØªØ¨Ø·Ø©: Ù†Ù†Ø´Ø¦ Ø¯Ø®Ù„Ù‹Ø§ ØºÙŠØ± Ù…Ø±ØªØ¨Ø· Ø¨Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø©ØŒ Ù…Ù†Ø³ÙˆØ¨Ù‹Ø§ Ù„ÙØ±Ø¹ Ù…ÙˆØ¸Ù Ø§Ù„Ù…Ø§Ù„ÙŠØ©
        if not created_income and not invoice.transaction_id:
            user = User.query.get(session.get("user_id"))
            finance_branch_id = getattr(user, "branch_id", None)
            # ØªØ¬Ù†Ø¨ Ø§Ù„ØªÙƒØ±Ø§Ø±: ØªØ­Ù‚Ù‚ Ù…Ù† ÙˆØ¬ÙˆØ¯ Ø¯ÙØ¹Ø© Ù„Ù†ÙØ³ Ø§Ù„ÙØ§ØªÙˆØ±Ø© Ø¹Ø¨Ø± Ù†ÙØ³ Ø§Ù„Ù…Ø¨Ù„Øº ÙˆØ§Ù„ØªØ§Ø±ÙŠØ® Ø§Ù„ØªÙ‚Ø±ÙŠØ¨ÙŠ
            existing_unlinked = Payment.query.filter_by(
                transaction_id=None,
                amount=invoice.amount,
                method="Ø¨Ù†Ùƒ",
            ).first()
            if not existing_unlinked:
                p = Payment(
                    transaction_id=None,
                    amount=invoice.amount,
                    method="Ø¨Ù†Ùƒ",
                    date_received=datetime.utcnow(),
                    received_by=session.get("username"),
                    branch_id=finance_branch_id,
                    receipt_file=filename,
                )
                db.session.add(p)
                db.session.commit()
                created_income = True

        if created_income:
            flash("âœ… ØªÙ… ØªØ­Ø¯ÙŠØ« Ø§Ù„Ø­Ø§Ù„Ø© ÙˆØ¥Ø¶Ø§ÙØ© Ø§Ù„Ø¯Ø®Ù„", "success")
        else:
            flash("âœ… ØªÙ… ØªØ­Ø¯ÙŠØ« Ø§Ù„Ø­Ø§Ù„Ø©", "success")
    else:
        flash("âš ï¸ Ø¥Ø¬Ø±Ø§Ø¡ ØºÙŠØ± Ù…Ø¹Ø±ÙˆÙ", "warning")

    return redirect(url_for("finance_dashboard"))

# âœ… Ø¥Ù†Ø´Ø§Ø¡ Ø¹Ø±Ø¶ Ø³Ø¹Ø± Ù„Ù„Ø¹Ù…ÙŠÙ„ (Ù…Ù† Ø§Ù„Ù…Ø§Ù„ÙŠØ©)
@app.route("/finance/customer_quotes", methods=["POST"])
def finance_create_customer_quote():
    if session.get("role") != "finance":
        return redirect(url_for("login"))

    customer_name = (request.form.get("customer_name") or "").strip()
    amount = float(request.form.get("amount") or 0)
    valid_until_str = request.form.get("valid_until")
    note = request.form.get("note")
    transaction_id = request.form.get("transaction_id")

    if not customer_name:
        flash("â›” Ø§Ø³Ù… Ø§Ù„Ø¹Ù…ÙŠÙ„ Ù…Ø·Ù„ÙˆØ¨", "danger")
        return redirect(url_for("finance_dashboard"))

    valid_until_dt = None
    if valid_until_str:
        try:
            valid_until_dt = datetime.fromisoformat(valid_until_str)
        except Exception:
            valid_until_dt = None

    q = CustomerQuote(
        customer_name=customer_name,
        amount=amount,
        valid_until=valid_until_dt,
        note=note,
        transaction_id=int(transaction_id) if transaction_id else None,
        created_by=session.get("user_id"),
    )
    db.session.add(q)
    db.session.commit()
    flash("âœ… ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ Ø¹Ø±Ø¶ Ø§Ù„Ø³Ø¹Ø± Ù„Ù„Ø¹Ù…ÙŠÙ„", "success")
    return redirect(url_for("print_customer_quote_html", quote_id=q.id) + "?auto=1")

# âœ… Ø¥Ù†Ø´Ø§Ø¡ ÙØ§ØªÙˆØ±Ø© Ù„Ù„Ø¹Ù…ÙŠÙ„ (Ù…Ù† Ø§Ù„Ù…Ø§Ù„ÙŠØ©)
@app.route("/finance/customer_invoices", methods=["POST"])
def finance_create_customer_invoice():
    if session.get("role") != "finance":
        return redirect(url_for("login"))

    customer_name = (request.form.get("customer_name") or "").strip()
    amount = float(request.form.get("amount") or 0)
    note = request.form.get("note")
    transaction_id = request.form.get("transaction_id")
    user = User.query.get(session["user_id"]) if session.get("user_id") else None

    if not customer_name:
        flash("â›” Ø§Ø³Ù… Ø§Ù„Ø¹Ù…ÙŠÙ„ Ù…Ø·Ù„ÙˆØ¨", "danger")
        return redirect(url_for("finance_dashboard"))

    # Ø±Ø¨Ø· ØªÙ„Ù‚Ø§Ø¦ÙŠ Ø¨Ø£Ø­Ø¯Ø« Ù…Ø¹Ø§Ù…Ù„Ø© ØªØ­Ù…Ù„ Ù†ÙØ³ Ø§Ø³Ù… Ø§Ù„Ø¹Ù…ÙŠÙ„ Ø¥Ù† Ù„Ù… ÙŠÙØ­Ø¯Ø¯ transaction_id
    resolved_transaction_id = None
    try:
        resolved_transaction_id = int(transaction_id) if transaction_id else None
    except Exception:
        resolved_transaction_id = None
    if not resolved_transaction_id and customer_name:
        try:
            # Ø£Ø­Ø¯Ø« Ù…Ø¹Ø§Ù…Ù„Ø© ØºÙŠØ± Ù…Ø¯ÙÙˆØ¹Ø© Ø¨Ù†ÙØ³ Ø§Ø³Ù… Ø§Ù„Ø¹Ù…ÙŠÙ„ ÙˆØ¶Ù…Ù† Ù†ÙØ³ ÙØ±Ø¹ Ø§Ù„Ù…Ø³ØªØ®Ø¯Ù… Ø§Ù„Ù…Ø§Ù„ÙŠ
            tx_query = Transaction.query.filter_by(client=customer_name, payment_status="ØºÙŠØ± Ù…Ø¯ÙÙˆØ¹Ø©")
            if user and getattr(user, "branch_id", None) is not None:
                tx_query = tx_query.filter(Transaction.branch_id == user.branch_id)
            candidate_tx = tx_query.order_by(Transaction.id.desc()).first()
            if candidate_tx:
                resolved_transaction_id = candidate_tx.id
        except Exception:
            resolved_transaction_id = None

    inv = CustomerInvoice(
        customer_name=customer_name,
        amount=amount,
        note=note,
        transaction_id=resolved_transaction_id,
        created_by=session.get("user_id"),
    )
    db.session.add(inv)
    db.session.commit()

    # ØªÙˆÙ„ÙŠØ¯ Ø±Ù‚Ù… ÙØ§ØªÙˆØ±Ø© ÙØ±ÙŠØ¯
    try:
        inv.invoice_number = generate_unique_invoice_number(prefix="INV", kind="CUST")
        db.session.commit()
    except Exception:
        db.session.rollback()
        try:
            inv.invoice_number = generate_unique_invoice_number(prefix="INV")
            db.session.commit()
        except Exception:
            db.session.rollback()
    # ØªÙ…Ø±ÙŠØ± Ù…Ø¯Ø®Ù„Ø§Øª Ø§Ù„Ø¶Ø±ÙŠØ¨Ø© Ù„Ù„Ø·Ø¨Ø§Ø¹Ø© Ø¥Ù† ÙˆÙØ¬Ø¯Øª
    apply_vat = (request.form.get("apply_vat") or "1")
    vat_percent = request.form.get("vat")
    qp = {
        "auto": "1",
        "apply_vat": apply_vat,
    }
    if vat_percent:
        qp["vat"] = str(vat_percent)
    from urllib.parse import urlencode
    flash("âœ… ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ ÙØ§ØªÙˆØ±Ø© Ø§Ù„Ø¹Ù…ÙŠÙ„", "success")
    return redirect(url_for("print_customer_invoice_html", invoice_id=inv.id) + "?" + urlencode(qp))

# ---------------- ØµÙØ­Ø© Ø§Ù„Ø¨Ù†ÙˆÙƒ: Ù†Ø¸Ø±Ø© Ø¹Ø§Ù…Ø© ----------------
@app.route("/banks")
def banks_overview():
    if session.get("role") not in ["manager", "finance"]:
        return redirect(url_for("login"))

    # ÙÙ„ØªØ±Ø© Ø¨Ø§Ù„ØªØ§Ø±ÙŠØ® (Ø§Ø®ØªÙŠØ§Ø±ÙŠ)
    start_date_str = request.args.get("start")
    end_date_str = request.args.get("end")
    # Parse dates safely; treat end as inclusive by moving to next day start
    start_date = None
    end_date = None
    try:
        if start_date_str:
            start_date = datetime.fromisoformat(start_date_str)
        if end_date_str:
            parsed_end = datetime.fromisoformat(end_date_str)
            end_date = parsed_end + timedelta(days=1)
    except Exception:
        # Ignore invalid date inputs gracefully
        start_date = start_date or None
        end_date = end_date or None

    tx_query = db.session.query(Bank.id, Bank.name, func.count(Transaction.id))\
        .outerjoin(Transaction, Transaction.bank_id == Bank.id)
    if start_date:
        tx_query = tx_query.filter(Transaction.date >= start_date)
    if end_date:
        # end_date is exclusive (next day), so use < rather than <=
        tx_query = tx_query.filter(Transaction.date < end_date)

    banks_stats = tx_query.group_by(Bank.id, Bank.name)\
        .order_by(Bank.name.asc()).all()

    banks_list = [
        {"id": b_id, "name": b_name, "count": tx_count}
        for (b_id, b_name, tx_count) in banks_stats
    ]

    # Keep original strings in template for inputs
    return render_template("banks.html", banks=banks_list, start=start_date_str, end=end_date_str)


# ---------------- ØµÙØ­Ø© Ø¨Ù†Ùƒ Ù…Ø­Ø¯Ø¯: ØªÙØ§ØµÙŠÙ„ ÙˆØ¥Ø­ØµØ§Ø¦ÙŠØ§Øª ----------------
@app.route("/banks/<int:bank_id>")
def bank_detail(bank_id):
    if session.get("role") not in ["manager", "finance"]:
        return redirect(url_for("login"))

    bank = Bank.query.get_or_404(bank_id)

    # ÙÙ„ØªØ±Ø© Ø¨Ø§Ù„ØªØ§Ø±ÙŠØ® (Ø§Ø®ØªÙŠØ§Ø±ÙŠ)
    start_date_str = request.args.get("start")
    end_date_str = request.args.get("end")
    # Parse dates safely; treat end as inclusive by moving to next day start
    start_date = None
    end_date = None
    try:
        if start_date_str:
            start_date = datetime.fromisoformat(start_date_str)
        if end_date_str:
            parsed_end = datetime.fromisoformat(end_date_str)
            end_date = parsed_end + timedelta(days=1)
    except Exception:
        start_date = start_date or None
        end_date = end_date or None

    # Ø¥Ø­ØµØ§Ø¦ÙŠØ© Ø¹Ø¯Ø¯ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø§Øª Ù„ÙƒÙ„ ÙØ±Ø¹ Ø¨Ù†Ùƒ Ù„Ù‡Ø°Ø§ Ø§Ù„Ø¨Ù†Ùƒ
    br_query = db.session.query(
        Transaction.bank_branch.label("bank_branch"),
        func.count(Transaction.id)
    ).filter(
        Transaction.bank_id == bank_id,
        Transaction.bank_branch.isnot(None),
        func.length(func.trim(Transaction.bank_branch)) > 0
    )
    if start_date:
        br_query = br_query.filter(Transaction.date >= start_date)
    if end_date:
        br_query = br_query.filter(Transaction.date < end_date)
    branch_rows = br_query.group_by(text("bank_branch")).order_by(text("bank_branch ASC")).all()
    branch_stats = [
        {"name": (bname or "ØºÙŠØ± Ù…Ø­Ø¯Ø¯"), "count": bcount}
        for (bname, bcount) in branch_rows
    ]

    total_tx = sum(b["count"] for b in branch_stats)

    # ğŸ‘¤ Ø¥Ø­ØµØ§Ø¦ÙŠØ© Ø¹Ø¯Ø¯ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø§Øª Ù„ÙƒÙ„ Ø§Ø³Ù… Ù…ÙˆØ¸Ù Ø¨Ù†Ùƒ
    emp_query = db.session.query(
        func.coalesce(func.trim(Transaction.bank_employee_name), "ØºÙŠØ± Ù…Ø¹Ø±ÙˆÙ").label("emp"),
        func.count(Transaction.id)
    ).filter(Transaction.bank_id == bank_id)
    if start_date:
        emp_query = emp_query.filter(Transaction.date >= start_date)
    if end_date:
        emp_query = emp_query.filter(Transaction.date < end_date)
    employee_rows = emp_query.group_by(text("emp")).order_by(text("emp ASC")).all()
    employee_stats = [
        {"name": ename or "ØºÙŠØ± Ù…Ø¹Ø±ÙˆÙ", "count": ecount}
        for (ename, ecount) in employee_rows
    ]

    # ğŸ‘¥ ØªØ¬Ù…ÙŠØ¹ Ø§Ù„Ù…ÙˆØ¸ÙÙŠÙ† Ø¯Ø§Ø®Ù„ ÙƒÙ„ ÙØ±Ø¹ (ÙØ±Ø¹ â† [Ù…ÙˆØ¸ÙÙˆÙ† + Ø¹Ø¯Ø¯ Ù…Ø¹Ø§Ù…Ù„Ø§ØªÙ‡Ù…])
    branch_emp_query = db.session.query(
        func.coalesce(func.trim(Transaction.bank_branch), "ØºÙŠØ± Ù…Ø­Ø¯Ø¯").label("branch"),
        func.coalesce(func.trim(Transaction.bank_employee_name), "ØºÙŠØ± Ù…Ø¹Ø±ÙˆÙ").label("emp"),
        func.count(Transaction.id)
    ).filter(Transaction.bank_id == bank_id)
    if start_date:
        branch_emp_query = branch_emp_query.filter(Transaction.date >= start_date)
    if end_date:
        branch_emp_query = branch_emp_query.filter(Transaction.date < end_date)
    branch_emp_rows = (
        branch_emp_query
        .group_by(text("branch, emp"))
        .order_by(text("branch ASC, emp ASC"))
        .all()
    )

    # Ø¨Ù†Ø§Ø¡ Ù‡ÙŠÙƒÙ„ Ù…ØªØ¯Ø§Ø®Ù„: [{ name, count, employees: [{ name, count }] }]
    from collections import defaultdict
    branch_to_employees = defaultdict(list)
    branch_totals = defaultdict(int)
    for bname, ename, ecount in branch_emp_rows:
        branch_to_employees[bname].append({"name": ename or "ØºÙŠØ± Ù…Ø¹Ø±ÙˆÙ", "count": ecount})
        branch_totals[bname] += ecount

    # ØªØ±ØªÙŠØ¨ Ø§Ù„Ù…ÙˆØ¸ÙÙŠÙ† Ø¯Ø§Ø®Ù„ ÙƒÙ„ ÙØ±Ø¹ Ø­Ø³Ø¨ Ø§Ù„Ø¹Ø¯Ø¯ ØªÙ†Ø§Ø²Ù„ÙŠÙ‹Ø§ØŒ Ø«Ù… Ø¨Ø§Ù„Ø§Ø³Ù…
    for bname, items in branch_to_employees.items():
        items.sort(key=lambda x: (-int(x.get("count") or 0), str(x.get("name") or "")))

    branches_nested = [
        {"name": bname or "ØºÙŠØ± Ù…Ø­Ø¯Ø¯", "count": branch_totals.get(bname, 0), "employees": branch_to_employees.get(bname, [])}
        for bname in sorted(branch_to_employees.keys(), key=lambda s: str(s or ""))
    ]

    # Ø§Ù„ÙÙˆØ§ØªÙŠØ± Ø§Ù„Ù…Ø±ØªØ¨Ø·Ø© Ø¨Ù…Ø¹Ø§Ù…Ù„Ø§Øª Ù‡Ø°Ø§ Ø§Ù„Ø¨Ù†Ùƒ (Ø§Ø¹ØªÙ…Ø§Ø¯Ø§Ù‹ Ø¹Ù„Ù‰ Ø¬Ø¯ÙˆÙ„ Payments)
    pay_query = Payment.query.join(Transaction, Payment.transaction_id == Transaction.id)\
        .filter(Transaction.bank_id == bank_id)
    if start_date:
        pay_query = pay_query.filter(Payment.date_received >= start_date)
    if end_date:
        pay_query = pay_query.filter(Payment.date_received < end_date)
    payments = pay_query.order_by(Payment.date_received.desc()).all()

    # Ø§Ù„Ù…Ø³ØªÙ†Ø¯Ø§Øª Ø§Ù„Ù…Ø±ØªØ¨Ø·Ø© Ø¨Ù…Ø¹Ø§Ù…Ù„Ø§Øª Ù‡Ø°Ø§ Ø§Ù„Ø¨Ù†Ùƒ
    # Ù…Ù„Ø§Ø­Ø¸Ø©: Ù†Ø¹Ø±Ø¶Ù‡Ø§ Ø¯Ø§Ø¦Ù…Ù‹Ø§ Ø¨Ø¯ÙˆÙ† Ø£ÙŠ ÙÙ„ØªØ±Ø© Ø­Ø³Ø¨ Ø§Ù„ØªØ§Ø±ÙŠØ® Ø¨Ù†Ø§Ø¡Ù‹ Ø¹Ù„Ù‰ Ø·Ù„Ø¨ Ø§Ù„Ù…Ø³ØªØ®Ø¯Ù…
    txs = (
        Transaction.query
        .filter(Transaction.bank_id == bank_id)
        .order_by(Transaction.id.desc())
        .all()
    )
    documents = []
    for t in txs:
        # Ù…Ù„ÙØ§Øª Ù…ØªØ¹Ø¯Ø¯Ø© Ù…Ø­ÙÙˆØ¸Ø© ÙƒØ³Ù„Ø³Ù„Ø© Ù…ÙØµÙˆÙ„Ø© Ø¨ÙÙˆØ§ØµÙ„
        if t.files:
            for fname in (t.files or "").split(","):
                fname = (fname or "").strip()
                if fname:
                    documents.append({"transaction_id": t.id, "filename": fname})
        # Ù„Ø§ Ù†Ø¹Ø±Ø¶ Ù…Ù„ÙØ§Øª Ø§Ù„ØªÙ‚Ø§Ø±ÙŠØ± Ù‡Ù†Ø§ Ø­Ø³Ø¨ Ø§Ù„Ø·Ù„Ø¨
        # Ù…Ù„ÙØ§Øª Ø§Ù„Ø¨Ù†Ùƒ Ø§Ù„ØªÙŠ Ø±ÙØ¹Ù‡Ø§ Ø§Ù„Ù…ÙˆØ¸Ù
        if getattr(t, "bank_sent_files", None):
            for fname in (t.bank_sent_files or "").split(","):
                fname = (fname or "").strip()
                if fname:
                    documents.append({"transaction_id": t.id, "filename": fname})

    # ğŸ“¨ Ù…Ø³ØªÙ†Ø¯Ø§Øª Ø¹Ø§Ù…Ø© Ù…Ø±Ø³Ù„Ø© Ù„Ù„Ø¨Ù†Ùƒ (ØºÙŠØ± Ù…Ø±ØªØ¨Ø·Ø© Ø¨Ù…Ø¹Ø§Ù…Ù„Ø©)
    try:
        general_docs = BankDocument.query.filter_by(bank_id=bank_id).order_by(BankDocument.id.desc()).all()
    except Exception:
        general_docs = []

    # ÙÙˆØ§ØªÙŠØ± Ø§Ù„Ø¨Ù†Ùƒ Ø¨Ù…Ø±Ø§Ø­Ù„Ù‡Ø§ (Ø¥Ù† ÙˆÙØ¬Ø¯Øª)
    inv_query = BankInvoice.query.filter_by(bank_id=bank_id)
    if start_date:
        inv_query = inv_query.filter(
            or_(
                BankInvoice.issued_at >= start_date,
                BankInvoice.delivered_at >= start_date,
                BankInvoice.received_at >= start_date,
            )
        )
    if end_date:
        inv_query = inv_query.filter(
            or_(
                BankInvoice.issued_at < end_date,
                BankInvoice.delivered_at < end_date,
                BankInvoice.received_at < end_date,
            )
        )
    invoices = inv_query.order_by(BankInvoice.id.desc()).all()

    # âœ… Ù…Ù„Ø®Øµ Ù„Ù„ÙÙˆØ§ØªÙŠØ± ÙˆØ§Ù„Ù…Ø±Ø§Ø­Ù„ (Ù„Ù„Ø¥Ø¸Ù‡Ø§Ø± ÙƒÙ…Ù„Ø®Øµ Ø¹Ù†Ø¯ Ø§Ù„Ù…Ø¯ÙŠØ±)
    total_invoices = len(invoices)
    total_amount = sum((inv.amount or 0) for inv in invoices)
    issued_count = sum(1 for inv in invoices if inv.issued_at)
    delivered_count = sum(1 for inv in invoices if inv.delivered_at)
    received_count = sum(1 for inv in invoices if inv.received_at)
    pending_count = total_invoices - received_count

    invoice_summary = {
        "total_invoices": total_invoices,
        "total_amount": total_amount,
        "issued_count": issued_count,
        "delivered_count": delivered_count,
        "received_count": received_count,
        "pending_count": pending_count,
    }

    return render_template(
        "bank_detail.html",
        bank=bank,
        branches=branch_stats,
        employees=employee_stats,
        branches_nested=branches_nested,
        total_tx=total_tx,
        payments=payments,
        documents=documents,
        general_docs=general_docs,
        invoices=invoices,
        invoice_summary=invoice_summary,
        start=start_date_str,
        end=end_date_str,
    )


# ---------------- ØµÙØ­Ø© Ù…ÙˆØ¸ÙÙŠ Ø§Ù„Ø¨Ù†ÙˆÙƒ ØºÙŠØ± Ø§Ù„Ù†Ø´Ø·ÙŠÙ† (> 15 ÙŠÙˆÙ…) ----------------
@app.route("/banks/inactive_employees")
def inactive_bank_employees():
    if session.get("role") not in ["manager", "finance"]:
        return redirect(url_for("login"))

    try:
        days_param = int(request.args.get("days") or 15)
        if days_param < 1:
            days_param = 15
    except Exception:
        days_param = 15

    now_utc = datetime.utcnow()
    threshold = now_utc - timedelta(days=days_param)

    emp = func.coalesce(func.trim(Transaction.bank_employee_name), "ØºÙŠØ± Ù…Ø¹Ø±ÙˆÙ").label("emp")

    query = (
        db.session.query(
            emp,
            func.max(Transaction.date).label("last_date"),
            func.count(Transaction.id).label("tx_count"),
        )
        .filter(
            Transaction.bank_employee_name != None,
            func.length(func.trim(Transaction.bank_employee_name)) > 0,
        )
        .group_by(text("emp"))
        .having(func.max(Transaction.date) < threshold)
        .order_by(func.max(Transaction.date).asc())
    )

    rows = query.all()

    items = []
    for emp_name, last_date, tx_count in rows:
        days_since = (now_utc - (last_date or now_utc)).days
        items.append({
            "name": emp_name or "ØºÙŠØ± Ù…Ø¹Ø±ÙˆÙ",
            "last_date": last_date,
            "days_since": days_since,
            "tx_count": tx_count or 0,
        })

    return render_template(
        "inactive_bank_employees.html",
        employees=items,
        days=days_param,
        now=now_utc,
    )


# ---------------- Ù…Ø³ØªÙ†Ø¯Ø§Øª ÙˆÙÙˆØ§ØªÙŠØ± Ø§Ù„ÙØ±ÙˆØ¹ (Ø¹Ø±Ø¶ ÙÙ‚Ø· Ù„Ù„Ù…Ø¯ÙŠØ±) ----------------
@app.route("/branch_documents", methods=["GET"])
def branch_documents():
    if session.get("role") != "manager":
        return redirect(url_for("login"))

    branches = Branch.query.order_by(Branch.name.asc()).all()

    # ÙÙ„ØªØ±Ø© Ø§Ø®ØªÙŠØ§Ø±ÙŠØ© Ø¨Ø§Ù„ÙØ±Ø¹
    selected_branch_id = request.args.get("branch_id")
    q = BranchDocument.query
    if selected_branch_id:
        q = q.filter_by(branch_id=int(selected_branch_id))
    docs = q.order_by(BranchDocument.expires_at.asc().nulls_last()).all()

    # ØªØµÙ†ÙŠÙ Ø§Ù„Ù…Ø³ØªÙ†Ø¯Ø§Øª
    now = datetime.utcnow()
    def status_for(doc):
        if not doc.expires_at:
            return "Ø¨Ø¯ÙˆÙ† Ø§Ù†ØªÙ‡Ø§Ø¡"
        delta = (doc.expires_at - now).days
        if delta < 0:
            return "Ù…Ù†ØªÙ‡ÙŠ"
        if delta <= 30:
            return "Ù‚Ø±ÙŠØ¨ Ø§Ù„Ø§Ù†ØªÙ‡Ø§Ø¡"
        return "Ø³Ø§Ø±ÙŠ"

    return render_template(
        "branch_documents.html",
        branches=branches,
        docs=docs,
        selected_branch_id=selected_branch_id,
        status_for=status_for,
    )


# ØªØ­Ø¯ÙŠØ« Ù…Ø±Ø§Ø­Ù„ ÙØ§ØªÙˆØ±Ø© Ø¨Ù†Ùƒ
@app.route("/banks/<int:bank_id>/invoice_stage", methods=["POST"]) 
def update_bank_invoice_stage(bank_id):
    # âœ… Ø­ØµØ± Ø¥Ø¯Ø®Ø§Ù„ ÙˆØªØ­Ø¯ÙŠØ« Ù…Ø±Ø§Ø­Ù„ ÙÙˆØ§ØªÙŠØ± Ø§Ù„Ø¨Ù†Ùƒ Ø¹Ù„Ù‰ Ù‚Ø³Ù… Ø§Ù„Ù…Ø§Ù„ÙŠØ© ÙÙ‚Ø·
    if session.get("role") != "finance":
        return redirect(url_for("login"))

    action = request.form.get("action")  # issue/deliver/receive
    amount = float(request.form.get("amount") or 0)
    invoice_id = request.form.get("invoice_id")
    transaction_id = request.form.get("transaction_id")
    note = request.form.get("note")

    # Ø£Ù†Ø´Ø¦/Ø­Ø¯Ø« Ø³Ø¬Ù„ Ø§Ù„ÙØ§ØªÙˆØ±Ø©
    invoice = None
    if invoice_id:
        invoice = BankInvoice.query.get(invoice_id)
    if not invoice:
        invoice = BankInvoice(bank_id=bank_id)
        if transaction_id:
            invoice.transaction_id = int(transaction_id)
        if amount:
            invoice.amount = amount
        db.session.add(invoice)
        db.session.commit()

    now_ts = datetime.utcnow()
    if action == "issue":
        invoice.issued_at = now_ts
        if amount:
            invoice.amount = amount
    elif action == "deliver":
        invoice.delivered_at = now_ts
    elif action == "receive":
        invoice.received_at = now_ts
    if note:
        invoice.note = note
    db.session.commit()

    flash("âœ… ØªÙ… ØªØ­Ø¯ÙŠØ« Ù…Ø±Ø­Ù„Ø© Ø§Ù„ÙØ§ØªÙˆØ±Ø©", "success")
    return redirect(url_for("bank_detail", bank_id=bank_id, start=request.args.get('start'), end=request.args.get('end')))

# ---------------- Ø¥Ø¯Ø§Ø±Ø© Ø§Ù„Ù…ÙˆØ¸ÙÙŠÙ† ----------------
@app.route("/manage_employees", methods=["GET", "POST"])
def manage_employees():
    if session.get("role") != "manager":
        return redirect(url_for("login"))

    if request.method == "POST":
        username = request.form.get("username")
        password = request.form["password"]
        role = request.form["role"]
        branch_id = request.form.get("branch_id")
        section_id = request.form.get("section_id") or None
        try:
            section_id = int(section_id) if section_id else None
        except Exception:
            section_id = None
        try:
            branch_id = int(branch_id) if branch_id else None
        except Exception:
            branch_id = None
        hashed_pw = generate_password_hash(password)
        user = User(username=username, password=hashed_pw, role=role, branch_id=branch_id, section_id=section_id)
        db.session.add(user)
        db.session.commit()
        flash("âœ… ØªÙ… Ø¥Ø¶Ø§ÙØ© Ø§Ù„Ù…ÙˆØ¸Ù Ø¨Ù†Ø¬Ø§Ø­", "success")
        return redirect(url_for("manage_employees"))

    users = User.query.all()
    branches = Branch.query.all()
    return render_template("manage_employees.html", users=users, branches=branches)

@app.route("/manager/employees/delete/<int:uid>")
def delete_employee(uid):
    if session.get("role") != "manager":
        return redirect(url_for("login"))
    u = User.query.get_or_404(uid)
    db.session.delete(u)
    db.session.commit()
    flash("âœ… ØªÙ… Ø­Ø°Ù Ø§Ù„Ù…ÙˆØ¸Ù", "success")
    return redirect(url_for("manage_employees"))

@app.route("/assign_branch/<int:uid>", methods=["POST"])
def assign_branch(uid):
    if session.get("role") != "manager":
        return redirect(url_for("login"))
    branch_id = request.form.get("branch_id")
    user = User.query.get_or_404(uid)
    if branch_id:
        user.branch_id = branch_id
        db.session.commit()
        flash("âœ… ØªÙ… ØªØ­Ø¯ÙŠØ¯ ÙØ±Ø¹ Ø§Ù„Ù…ÙˆØ¸Ù Ø¨Ù†Ø¬Ø§Ø­", "success")
    return redirect(url_for("manager_dashboard"))

@app.route("/r/<string:token>")
def public_report(token):
    t = Transaction.query.filter_by(public_share_token=token).first_or_404()
    if not t.report_file:
        abort(404)
    # Ø¥Ù† ÙƒØ§Ù† Ù„Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ù…Ù„Ù Ù…Ø±ÙÙˆØ¹Ù‹Ø§ Ø¹Ù„Ù‰ B2 ÙˆÙ†Ø³ØªØ·ÙŠØ¹ ØªÙˆÙ„ÙŠØ¯ Ø±Ø§Ø¨Ø· Ø¹Ø§Ù…ØŒ Ø­ÙˆÙ‘Ù„ Ø§Ù„Ù…Ø³ØªØ®Ø¯Ù… Ù„Ù‡
    if getattr(t, "report_b2_file_name", None):
        b2_url = build_b2_public_url(t.report_b2_file_name)
        if b2_url:
            return redirect(b2_url)
    return send_from_directory(app.config["UPLOAD_FOLDER"], t.report_file)

# ---------------- Ø¹Ø±Ø¶ Ø§Ù„Ù…Ù„ÙØ§Øª ----------------
@app.route("/uploads/<path:filename>")
def uploaded_file(filename):
    resp = send_from_directory(app.config["UPLOAD_FOLDER"], filename)
    try:
        # Ensure inline display in browser tabs instead of forced download
        basename = os.path.basename(filename)
        resp.headers["Content-Disposition"] = f'inline; filename="{basename}"'
        # Make sure Content-Type is set appropriately for better inline rendering
        if not resp.headers.get("Content-Type") or resp.headers.get("Content-Type") == "application/octet-stream":
            import mimetypes
            guessed, _ = mimetypes.guess_type(basename)
            if guessed:
                resp.headers["Content-Type"] = guessed
    except Exception:
        pass
    return resp

# ØªÙ†Ø²ÙŠÙ„ Ù…Ù„Ù Ù…Ø­Ù„ÙŠ Ù…Ù† Ù…Ø¬Ù„Ø¯ Ø§Ù„Ø±ÙØ¹ Ù…Ø¹ Ø¥Ø¬Ø¨Ø§Ø± Ø§Ù„ØªÙ†Ø²ÙŠÙ„
@app.route("/download/local/<path:filename>")
def download_local_file(filename):
    # ÙŠØ³Ù…Ø­ Ù„Ù„Ù…Ø¯ÙŠØ± ÙˆØ§Ù„Ù…Ø§Ù„ÙŠØ© ÙÙ‚Ø·
    if session.get("role") not in ["manager", "finance"]:
        return redirect(url_for("login"))
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename, as_attachment=True)

# ØªÙ†Ø²ÙŠÙ„ Ù…Ù„Ù Ù…Ù† B2 Ø¹Ø¨Ø± Ù‚Ø±Ø§Ø¡Ø© Ø§Ù„Ù…Ø­ØªÙˆÙ‰ ÙˆØªÙ…Ø±ÙŠØ±Ù‡ ÙƒÙ…Ø±ÙÙ‚ (ÙÙŠ Ø­Ø§Ù„ Ø§Ù„Ø¨ÙƒØª Ø®Ø§Øµ)
@app.route("/download/b2")
def download_b2_file():
    # ÙŠØ³Ù…Ø­ Ù„Ù„Ù…Ø¯ÙŠØ± ÙˆØ§Ù„Ù…Ø§Ù„ÙŠØ© ÙÙ‚Ø·
    if session.get("role") not in ["manager", "finance"]:
        return redirect(url_for("login"))
    file_name = request.args.get("file")
    if not file_name:
        abort(400)
    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¨Ù†Ø§Ø¡ Ø±Ø§Ø¨Ø· Ø¹Ø§Ù… Ø¥Ù† ÙƒØ§Ù† Ø§Ù„Ø¨ÙƒØª Ø¹Ø§Ù…Ù‹Ø§
    b2_public = build_b2_public_url(file_name)
    if b2_public:
        # ØªÙ…Ø±ÙŠØ± ØªÙ„Ù…ÙŠØ­ Ø§Ù„ØªÙ†Ø²ÙŠÙ„ Ù„Ù„Ù…Ø³ØªØ®Ø¯Ù… Ø¹Ø¨Ø± Ù‡ÙŠØ¯Ø± Content-Disposition
        try:
            resp = requests.get(b2_public, stream=True, timeout=20)
            resp.raise_for_status()
            from flask import Response
            cd = f"attachment; filename*=UTF-8''{file_name}"
            headers = {
                "Content-Disposition": cd,
                "Content-Type": resp.headers.get("Content-Type", "application/octet-stream"),
            }
            return Response(resp.iter_content(chunk_size=8192), headers=headers)
        except Exception:
            abort(404)
    # Ø¥Ù† Ù„Ù… Ù†Ø³ØªØ·Ø¹ØŒ Ù†ÙØ´Ù„ Ø¨Ø´ÙƒÙ„ Ø¢Ù…Ù†
    abort(404)

# (ØªÙ…Øª Ø¥Ø²Ø§Ù„Ø© Ù…Ø³Ø§Ø±Ø§Øª QR ÙˆØ§Ù„Ø±ÙˆØ§Ø¨Ø· Ø§Ù„Ø¹Ø§Ù…Ø© Ø§Ù„Ù…Ø±ØªØ¨Ø·Ø© Ø¨Ù‡Ø§)

# ---------------- ØªÙƒØ§Ù…Ù„ ØµÙØ­Ø© Ø§Ù„Ø¨Ø§Ø±ÙƒÙˆØ¯ (index.html) ----------------
@app.route("/barcode", endpoint="barcode_page")
def barcode_page():
    """ØªØ®Ø¯Ù… ØµÙØ­Ø© Ø§Ù„Ø¨Ø§Ø±ÙƒÙˆØ¯/QR (index.html) Ù…Ù† Ø¬Ø°Ø± Ø§Ù„Ù…Ø´Ø±ÙˆØ¹.

    ÙŠÙ…ÙƒÙ† ØªÙ…Ø±ÙŠØ± ?hash=<sha256>&print=1 Ù„ÙŠØªÙ… Ø¹Ø±Ø¶ QR ÙˆØ§Ù„Ø·Ø¨Ø§Ø¹Ø© Ù…Ø¨Ø§Ø´Ø±Ø©.
    """
    index_path = os.path.join(app.root_path, "index.html")
    if not os.path.exists(index_path):
        abort(404)
    return send_file(index_path)


@app.route("/verify")
def verify_by_hash():
    """Ø§Ù„ØªØ­Ù‚Ù‚ Ù…Ù† Ø£ØµØ§Ù„Ø© Ø§Ù„ØªÙ‚Ø±ÙŠØ± Ø¹Ø¨Ø± Ø¨ØµÙ…Ø© SHA-256 Ø§Ù„Ù…Ø®Ø²Ù†Ø© ÙÙŠ Ù‚Ø§Ø¹Ø¯Ø© Ø§Ù„Ø¨ÙŠØ§Ù†Ø§Øª.

    Ù…Ø«Ø§Ù„: /verify?hash=<sha256>
    """
    h = request.args.get("hash", type=str)
    if not h:
        return "âŒ Ù„Ø§ ÙŠÙˆØ¬Ø¯ hash Ù„Ù„ØªØ­Ù‚Ù‚", 400

    t = Transaction.query.filter_by(report_sha256=h).first()
    if not t or not t.report_file:
        return "<h2>âŒ Ù‡Ø°Ø§ Ø§Ù„ØªÙ‚Ø±ÙŠØ± ØºÙŠØ± Ø£ØµÙ„ÙŠ Ø£Ùˆ ØªÙ… Ø§Ù„ØªØ¹Ø¯ÙŠÙ„</h2>", 404

    # Ø¥Ù† ØªÙˆÙØ± Ù…Ù„Ù Ø¹Ù„Ù‰ B2 Ù†Ø­Ø§ÙˆÙ„ Ø¥Ù†Ø´Ø§Ø¡ Ø±Ø§Ø¨Ø· Ø¹Ø§Ù… Ø¯Ø§Ø¦Ù…
    if getattr(t, "report_b2_file_name", None):
        b2_url = build_b2_public_url(t.report_b2_file_name)
    else:
        b2_url = None
    file_url = b2_url or url_for("uploaded_file", filename=t.report_file)
    return (
        f"""
        <h2>âœ… Ø§Ù„ØªÙ‚Ø±ÙŠØ± Ø£ØµÙ„ÙŠ</h2>
        <p>Ø±Ù‚Ù… Ø§Ù„ØªÙ‚Ø±ÙŠØ±: {t.report_number or '-'} | Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø©: {t.id}</p>
        <p><a href="{file_url}" target="_blank">ğŸ“„ Ø¹Ø±Ø¶ Ø§Ù„Ù…Ù„Ù</a></p>
        """
    )


@app.route("/file")
def file_by_hash():
    """Ø¥Ø±Ø¬Ø§Ø¹ Ù…Ù„Ù Ø§Ù„ØªÙ‚Ø±ÙŠØ± Ù…Ø¨Ø§Ø´Ø±Ø©Ù‹ Ø¹Ø¨Ø± Ø§Ù„Ø¨ØµÙ…Ø©.

    Ù…Ø«Ø§Ù„: /file?hash=<sha256>
    """
    h = request.args.get("hash", type=str)
    if not h:
        return "âŒ Ù„Ø§ ÙŠÙˆØ¬Ø¯ hash", 400

    t = Transaction.query.filter_by(report_sha256=h).first()
    if not t or not t.report_file:
        return "âŒ Ù„Ù… ÙŠØªÙ… Ø§Ù„Ø¹Ø«ÙˆØ± Ø¹Ù„Ù‰ Ù…Ù„Ù Ù…Ø±ØªØ¨Ø· Ø¨Ù‡Ø°Ø§ Ø§Ù„Ù‡Ø§Ø´", 404

    # Ø¥Ù† ØªÙˆÙØ± Ù…Ù„Ù Ø¹Ù„Ù‰ B2 Ù†Ø­Ø§ÙˆÙ„ Ø§Ù„ØªØ­ÙˆÙŠÙ„ Ø¥Ù„ÙŠÙ‡ Ù…Ø¨Ø§Ø´Ø±Ø©
    if getattr(t, "report_b2_file_name", None):
        b2_url = build_b2_public_url(t.report_b2_file_name)
        if b2_url:
            return redirect(b2_url)
    return send_from_directory(app.config["UPLOAD_FOLDER"], t.report_file)

# ---------------- Ø±ÙØ¹ Ù…Ù„Ù Ø¥Ù„Ù‰ Backblaze B2 ----------------
@app.route("/api/upload", methods=["POST"])
def api_upload_to_b2():
    if session.get("user_id") is None:
        print("âš ï¸ /api/upload unauthorized access: no user_id in session")
        return jsonify({"error": "unauthorized"}), 401

    if "file" not in request.files:
        return jsonify({"error": "no_file"}), 400

    f = request.files["file"]
    if not f or not f.filename:
        return jsonify({"error": "empty_filename"}), 400

    fname = secure_filename(f.filename)
    try:
        bucket = get_b2_bucket()
        # Ù†Ù‚Ø±Ø£ Ø§Ù„Ù…Ø­ØªÙˆÙ‰ Ø¥Ù„Ù‰ Ø§Ù„Ø°Ø§ÙƒØ±Ø© Ù„Ù„Ø¨Ø³Ø§Ø·Ø©. ÙŠÙ…ÙƒÙ† ØªØ­Ø³ÙŠÙ† Ø°Ù„Ùƒ Ø¨ØªØ¯ÙÙ‚ chunked Ø¹Ù†Ø¯ Ø§Ù„Ø­Ø§Ø¬Ø©
        data = f.read()
        # Ù„ØªÙØ§Ø¯ÙŠ Ø§Ù„ØªØ¹Ø§Ø±Ø¶ØŒ Ù†Ø¶ÙŠÙ Ø·Ø§Ø¨Ø¹Ù‹Ø§ Ø²Ù…Ù†ÙŠÙ‹Ø§ Ù„Ùˆ ÙƒØ§Ù† Ø§Ù„Ø§Ø³Ù… Ù…Ø³ØªØ®Ø¯Ù…Ù‹Ø§
        unique_name = f"{int(time.time())}_{fname}"
        uploaded = bucket.upload_bytes(data, file_name=unique_name)
        file_id = uploaded.id_ if hasattr(uploaded, "id_") else getattr(uploaded, "file_id", None)
        return jsonify({
            "status": "ok",
            "bucket_id": app.config.get("B2_BUCKET_ID"),
            "file_name": unique_name,
            "file_id": file_id,
        })
    except Exception as e:
        print(
            f"âš ï¸ /api/upload B2 error: {e} | "
            f"has_key_id={bool(app.config.get('B2_KEY_ID'))}, "
            f"has_application_key={bool(app.config.get('B2_APPLICATION_KEY'))}, "
            f"bucket_id={app.config.get('B2_BUCKET_ID')}")
        return jsonify({"error": str(e)}), 500

# ---------------- ÙØ­Øµ ØµØ­Ø© Ø§Ù„Ø±Ø¨Ø· Ù…Ø¹ Backblaze B2 ----------------
@app.route("/api/b2/health", methods=["GET"])
def api_b2_health():
    # Ù†Ù‚ÙŠÙ‘Ø¯ Ø§Ù„ÙˆØµÙˆÙ„ Ø¹Ù„Ù‰ Ø§Ù„Ù…Ø¯Ø±Ø§Ø¡/Ø§Ù„Ù…Ø¯ÙŠØ± Ø§Ù„Ø¹Ø§Ù… ÙÙ‚Ø·
    if session.get("role") not in ["manager", "admin"]:
        return jsonify({"error": "unauthorized"}), 401

    env_info = {
        "has_key_id": bool(app.config.get("B2_KEY_ID")),
        "has_application_key": bool(app.config.get("B2_APPLICATION_KEY")),
        "bucket_id": app.config.get("B2_BUCKET_ID"),
        "bucket_name": app.config.get("B2_BUCKET_NAME") or app.config.get("B2_BUCKET"),
    }

    try:
        api = get_b2_api()
    except Exception as e:
        return jsonify({
            "status": "error",
            "stage": "authorize_account",
            "message": str(e),
            "env": env_info,
        }), 500

    try:
        bucket = get_b2_bucket()
        result = {
            "status": "ok",
            "env": env_info,
            "bucket": {
                "id": getattr(bucket, "id_", None),
                "name": getattr(bucket, "name", None),
            },
        }

        # Ù…Ø­Ø§ÙˆÙ„Ø© Ù‚Ø±Ø§Ø¡Ø© Ø®ÙÙŠÙØ© (Ø§Ø®ØªÙŠØ§Ø±ÙŠØ©) Ù„Ù„ØªØ£ÙƒØ¯ Ù…Ù† Ø§Ù„ØµÙ„Ø§Ø­ÙŠØ§Øª
        try:
            sample_files = []
            try:
                # Ù‚Ø¯ Ù„Ø§ ØªØªÙˆÙØ± ÙÙŠ ÙƒÙ„ Ø§Ù„Ø¥ØµØ¯Ø§Ø±Ø§ØªØ› Ù…Ù„ÙÙ‘Ø© Ø¯Ø§Ø®Ù„ try
                for item in bucket.ls(max_entries=1):
                    file_info, _ = item if isinstance(item, tuple) else (None, None)
                    if file_info is not None:
                        sample_files.append(getattr(file_info, "file_name", None))
            except Exception:
                # ØªØ¬Ø§Ù‡Ù„ Ø¥Ù† Ù„Ù… ØªØªÙˆÙØ± Ø§Ù„Ø¯Ø§Ù„Ø©
                pass
            result["list_sample"] = sample_files
        except Exception as inner:
            result["list_error"] = str(inner)

        return jsonify(result)
    except Exception as e:
        return jsonify({
            "status": "error",
            "stage": "get_bucket",
            "message": str(e),
            "env": env_info,
        }), 500

# ---------------- ØªÙ†Ø²ÙŠÙ„ Ù…Ù„Ù Ù…Ù† Backblaze B2 ----------------
@app.route("/api/b2/download", methods=["GET"])
def api_b2_download():
    if session.get("user_id") is None:
        return jsonify({"error": "unauthorized"}), 401

    # ÙŠØªÙˆÙ‚Ø¹ query param ?name=path/inside/bucket
    file_name = request.args.get("name", type=str)
    if not file_name:
        return jsonify({"error": "missing_name"}), 400

    try:
        bucket = get_b2_bucket()
        # ØªÙ†Ø²ÙŠÙ„ Ø§Ù„Ù…Ø­ØªÙˆÙ‰ Ø¥Ù„Ù‰ Ø§Ù„Ø°Ø§ÙƒØ±Ø©. ÙŠÙ…ÙƒÙ† Ø§Ù„ØªØ­Ø³ÙŠÙ† Ø¹Ø¨Ø± stream_tmp file Ø¹Ù†Ø¯ Ø§Ù„ÙƒØ¨Ø±.
        downloaded = bucket.download_file_by_name(file_name)
        data = downloaded.response.content

        # ØªØ®Ù…ÙŠÙ† Ù†ÙˆØ¹ Ø§Ù„Ù…Ø­ØªÙˆÙ‰ Ø¨Ø´ÙƒÙ„ Ø¨Ø³ÙŠØ·
        guessed = "application/octet-stream"
        if file_name.lower().endswith(".pdf"):
            guessed = "application/pdf"
        elif file_name.lower().endswith(".png"):
            guessed = "image/png"
        elif file_name.lower().endswith(".jpg") or file_name.lower().endswith(".jpeg"):
            guessed = "image/jpeg"
        elif file_name.lower().endswith(".docx"):
            guessed = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        return Response(data, mimetype=guessed, headers={
            "Content-Disposition": f"attachment; filename=\"{os.path.basename(file_name)}\""
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ---------------- ØµÙØ­Ø© Ø§Ù„ØªÙ‚Ø§Ø±ÙŠØ± Ø§Ù„Ù…Ø´ØªØ±ÙƒØ© ----------------
@app.route("/employee/upload_bank_docs/<int:tid>", methods=["POST"])
def employee_upload_bank_docs(tid):
    if session.get("role") != "employee":
        return redirect(url_for("login"))

    t = Transaction.query.get_or_404(tid)
    user = User.query.get(session["user_id"])
    # Ù…Ù†Ø¹ Ø±ÙØ¹ Ù…Ù„ÙØ§Øª Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ù…Ù† ÙØ±Ø¹ Ø¢Ø®Ø±
    if t.branch_id != user.branch_id:
        flash("â›” Ù„Ø§ ÙŠÙ…ÙƒÙ†Ùƒ Ø±ÙØ¹ Ù…Ø³ØªÙ†Ø¯Ø§Øª Ù„Ù…Ø¹Ø§Ù„Ø¬Ø© Ù…Ù† ÙØ±Ø¹ Ø¢Ø®Ø±", "danger")
        return redirect(url_for("employee_dashboard"))

    # ØªØ£ÙƒØ¯ Ù…Ù† ÙˆØ¬ÙˆØ¯ Ø¨Ù†Ùƒ Ù…Ø±ØªØ¨Ø· Ø¨Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ù‚Ø¨Ù„ Ù‚Ø¨ÙˆÙ„ Ù…Ù„ÙØ§Øª Ø§Ù„Ø¨Ù†Ùƒ
    if not t.bank_id:
        flash("âš ï¸ ÙŠØ¬Ø¨ Ø±Ø¨Ø· Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ø¨Ø¨Ù†Ùƒ Ù‚Ø¨Ù„ Ø±ÙØ¹ Ù…Ø³ØªÙ†Ø¯Ø§Øª Ø§Ù„Ø¨Ù†Ùƒ.", "warning")
        return redirect(url_for("employee_dashboard"))

    uploaded = request.files.getlist("bank_docs")
    saved = []
    for f in uploaded:
        if f and f.filename:
            fname = secure_filename(f.filename)
            f.save(os.path.join(app.config["UPLOAD_FOLDER"], fname))
            saved.append(fname)

    if saved:
        existing = (t.bank_sent_files or "").split(",") if t.bank_sent_files else []
        existing = [x.strip() for x in existing if x.strip()]
        t.bank_sent_files = ",".join(existing + saved)
        db.session.commit()
        flash("âœ… ØªÙ… Ø±ÙØ¹ Ù…Ù„ÙØ§Øª Ø§Ù„Ø¨Ù†Ùƒ ÙˆØ­ÙØ¸Ù‡Ø§", "success")
    else:
        flash("âš ï¸ Ù„Ù… ÙŠØªÙ… Ø§Ø®ØªÙŠØ§Ø± Ø£ÙŠ Ù…Ù„Ù", "warning")

    return redirect(url_for("employee_dashboard"))

# âœ… Ø±ÙØ¹ Ù…Ø³ØªÙ†Ø¯Ø§Øª Ø§Ù„Ø¨Ù†Ùƒ Ø¨Ø§Ù„Ø¨Ø­Ø« Ø¨Ø±Ù‚Ù… Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ø£Ùˆ Ø¨Ø§Ø³Ù… Ø§Ù„Ø¹Ù…ÙŠÙ„
@app.route("/employee/upload_bank_docs_lookup", methods=["POST"])
def employee_upload_bank_docs_lookup():
    if session.get("role") != "employee":
        return redirect(url_for("login"))

    lookup_raw = (request.form.get("lookup") or "").strip()

    # ØªØ­ÙˆÙŠÙ„ Ø§Ù„Ø£Ø±Ù‚Ø§Ù… Ø§Ù„Ø¹Ø±Ø¨ÙŠØ©/Ø§Ù„ÙØ§Ø±Ø³ÙŠØ© Ø¥Ù„Ù‰ Ø£Ø±Ù‚Ø§Ù… Ù„Ø§ØªÙŠÙ†ÙŠØ© Ù„Ø¶Ù…Ø§Ù† Ø§Ù„Ù…Ø·Ø§Ø¨Ù‚Ø© Ø§Ù„ØµØ­ÙŠØ­Ø©
    def normalize_digits(value: str) -> str:
        translation_table = str.maketrans("Ù Ù¡Ù¢Ù£Ù¤Ù¥Ù¦Ù§Ù¨Ù©Û°Û±Û²Û³Û´ÛµÛ¶Û·Û¸Û¹", "01234567890123456789")
        return value.translate(translation_table)

    lookup = normalize_digits(lookup_raw)

    if not lookup:
        flash("âš ï¸ ÙŠØ±Ø¬Ù‰ Ø¥Ø¯Ø®Ø§Ù„ Ø±Ù‚Ù… Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ø£Ùˆ Ø§Ø³Ù… Ø§Ù„Ø¹Ù…ÙŠÙ„", "warning")
        return redirect(url_for("employee_dashboard"))

    # Ù…Ø­Ø§ÙˆÙ„Ø© ØªÙØ³ÙŠØ±Ù‡ ÙƒØ±Ù‚Ù… Ù…Ø¹Ø§Ù…Ù„Ø© Ø£ÙˆÙ„Ù‹Ø§ (Ø¨Ø§Ø³ØªØ®Ø±Ø§Ø¬ Ø§Ù„Ø£Ø±Ù‚Ø§Ù… ÙÙ‚Ø· Ù…Ù† Ø§Ù„Ù†Øµ)
    t = None
    digits_only = "".join(ch for ch in lookup if ch.isdigit())
    if digits_only:
        try:
            tid = int(digits_only)
            t = Transaction.query.get(tid)
        except Exception:
            t = None

    # Ø¥Ù† Ù„Ù… ÙŠÙƒÙ† Ø±Ù‚Ù…ØŒ Ù†Ø¨Ø­Ø« Ø¨Ø§Ù„Ø§Ø³Ù… (ÙŠØ·Ø§Ø¨Ù‚ Ø¬Ø²Ø¦ÙŠÙ‹Ø§ Ø£Ø­Ø¯Ø« Ù…Ø¹Ø§Ù…Ù„Ø©)
    if not t:
        t = (
            Transaction.query
            .filter(Transaction.client.ilike(f"%{lookup}%"))
            .order_by(Transaction.id.desc())
            .first()
        )

    if not t:
        # ØªÙ„Ù…ÙŠØ­ Ù„Ù„Ù…Ø³ØªØ®Ø¯Ù… Ø­ÙˆÙ„ Ø§Ù„ØµÙŠØºØ© Ø§Ù„ØµØ­ÙŠØ­Ø© Ù„Ù„Ø¨Ø­Ø«
        flash("âŒ Ù„Ù… ÙŠØªÙ… Ø§Ù„Ø¹Ø«ÙˆØ± Ø¹Ù„Ù‰ Ù…Ø¹Ø§Ù…Ù„Ø© Ø¨Ù‡Ø°Ø§ Ø§Ù„Ø±Ù‚Ù… Ø£Ùˆ Ø§Ù„Ø§Ø³Ù…. Ø¬Ø±Ù‘Ø¨ Ø¥Ø¯Ø®Ø§Ù„ Ø±Ù‚Ù… Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© ÙÙ‚Ø· Ø£Ùˆ Ø§Ø³Ù… Ø§Ù„Ø¹Ù…ÙŠÙ„ Ø§Ù„ÙƒØ§Ù…Ù„.", "danger")
        return redirect(url_for("employee_dashboard"))

    user = User.query.get(session.get("user_id"))
    if not user:
        return redirect(url_for("login"))

    # Ù…Ù†Ø¹ Ø±ÙØ¹ Ù…Ù„ÙØ§Øª Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ù…Ù† ÙØ±Ø¹ Ø¢Ø®Ø±
    if t.branch_id != user.branch_id:
        flash("â›” Ù„Ø§ ÙŠÙ…ÙƒÙ†Ùƒ Ø±ÙØ¹ Ù…Ø³ØªÙ†Ø¯Ø§Øª Ù„Ù…Ø¹Ø§Ù„Ø¬Ø© Ù…Ù† ÙØ±Ø¹ Ø¢Ø®Ø±", "danger")
        return redirect(url_for("employee_dashboard"))

    # Ø§Ø®ØªÙŠØ§Ø± Ø§Ù„Ø¨Ù†Ùƒ (Ø§Ø®ØªÙŠØ§Ø±ÙŠ) Ù„ØªØ«Ø¨ÙŠØªÙ‡ Ø¹Ù„Ù‰ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ø¥Ù† ÙƒØ§Ù† ÙØ§Ø±ØºÙ‹Ø§
    bank_id_form = request.form.get("bank_id")
    try:
        bank_id_val = int(bank_id_form) if bank_id_form else None
    except Exception:
        bank_id_val = None

    # Ø¥Ø°Ø§ Ù„Ù… ØªÙƒÙ† Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ù…Ø±ØªØ¨Ø·Ø© Ø¨Ø¨Ù†Ùƒ ÙˆÙ„Ù… ÙŠÙØ­Ø¯Ø¯ Ø¨Ù†Ùƒ ÙÙŠ Ø§Ù„Ù†Ù…ÙˆØ°Ø¬: Ù„Ø§ Ù†Ø­ÙØ¸ ÙˆØ«Ø§Ø¦Ù‚ Ù„Ù† ØªØ¸Ù‡Ø± Ø¨Ø£ÙŠ Ø¨Ù†Ùƒ
    if not t.bank_id and not bank_id_val:
        flash("âš ï¸ Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© ØºÙŠØ± Ù…Ø±ØªØ¨Ø·Ø© Ø¨Ø£ÙŠ Ø¨Ù†Ùƒ. ÙŠØ±Ø¬Ù‰ Ø§Ø®ØªÙŠØ§Ø± Ø§Ù„Ø¨Ù†Ùƒ Ù‚Ø¨Ù„ Ø±ÙØ¹ Ø§Ù„Ù…Ø³ØªÙ†Ø¯Ø§Øª.", "warning")
        return redirect(url_for("employee_dashboard"))

    uploaded = request.files.getlist("bank_docs")
    saved = []
    for f in uploaded:
        if f and f.filename:
            fname = secure_filename(f.filename)
            f.save(os.path.join(app.config["UPLOAD_FOLDER"], fname))
            saved.append(fname)

    if saved:
        # ÙÙŠ Ø­Ø§Ù„ ØªÙ… ØªØ­Ø¯ÙŠØ¯ Ø¨Ù†Ùƒ ÙÙŠ Ø§Ù„Ù†Ù…ÙˆØ°Ø¬:
        # - Ø¥Ù† Ù„Ù… ØªÙƒÙ† Ø§Ù„Ù…Ø¹Ø§Ù…Ù„Ø© Ù…Ø±ØªØ¨Ø·Ø© Ø¨Ø¨Ù†Ùƒ Ù†Ø«Ø¨Ù‘Øª Ø§Ù„Ø¨Ù†Ùƒ Ø§Ù„Ù…Ø®ØªØ§Ø±
        # - ÙˆØ¥Ù† ÙƒØ§Ù†Øª Ù…Ø±ØªØ¨Ø·Ø© Ø¨Ø¨Ù†Ùƒ Ù…Ø®ØªÙ„ÙØŒ Ù†Ø­Ø¯Ø« Ø§Ù„Ø±Ø¨Ø· Ù„Ù„Ø¨Ù†Ùƒ Ø§Ù„Ù…Ø®ØªØ§Ø± Ù„Ø¶Ù…Ø§Ù† Ø¸Ù‡ÙˆØ± Ø§Ù„Ù…Ø³ØªÙ†Ø¯Ø§Øª ÙÙŠ ØµÙØ­Ø© Ø§Ù„Ø¨Ù†Ùƒ Ø§Ù„ØµØ­ÙŠØ­Ø©
        if bank_id_val and (not t.bank_id or t.bank_id != bank_id_val):
            t.bank_id = bank_id_val
        existing = (t.bank_sent_files or "").split(",") if t.bank_sent_files else []
        existing = [x.strip() for x in existing if x.strip()]
        t.bank_sent_files = ",".join(existing + saved)
        db.session.commit()
        flash("âœ… ØªÙ… Ø±ÙØ¹ Ù…Ù„ÙØ§Øª Ø§Ù„Ø¨Ù†Ùƒ ÙˆØ­ÙØ¸Ù‡Ø§", "success")
    else:
        flash("âš ï¸ Ù„Ù… ÙŠØªÙ… Ø§Ø®ØªÙŠØ§Ø± Ø£ÙŠ Ù…Ù„Ù", "warning")

    return redirect(url_for("employee_dashboard"))

# âœ… Ø±ÙØ¹ Ù…Ø³ØªÙ†Ø¯/Ø±Ø³Ø§Ù„Ø© Ø¹Ø§Ù…Ø© Ù„Ø¨Ù†Ùƒ (ØºÙŠØ± Ù…Ø±ØªØ¨Ø·Ø© Ø¨Ù…Ø¹Ø§Ù…Ù„Ø©)
@app.route("/employee/bank_documents", methods=["POST"])
def employee_add_bank_document():
    if session.get("role") != "employee":
        return redirect(url_for("login"))

    user = User.query.get(session.get("user_id"))
    if not user:
        return redirect(url_for("login"))

    bank_id = request.form.get("bank_id")
    title = (request.form.get("title") or "").strip()
    message = (request.form.get("message") or "").strip()
    doc_type = (request.form.get("doc_type") or "").strip()
    file = request.files.get("file")

    try:
        bank_id_val = int(bank_id) if bank_id else None
    except Exception:
        bank_id_val = None

    if not bank_id_val:
        flash("âš ï¸ ÙŠØ±Ø¬Ù‰ Ø§Ø®ØªÙŠØ§Ø± Ø§Ù„Ø¨Ù†Ùƒ", "warning")
        return redirect(url_for("employee_dashboard"))
    if not title:
        flash("âš ï¸ Ø§Ù„Ø¹Ù†ÙˆØ§Ù† Ù…Ø·Ù„ÙˆØ¨", "warning")
        return redirect(url_for("employee_dashboard"))

    filename = None
    b2_file_name = None
    b2_file_id = None
    if file and file.filename:
        # Try uploading to Backblaze B2; if fails, fallback to local save
        safe_name = secure_filename(file.filename)
        try:
            bucket = get_b2_bucket()
            data = file.read()
            unique_name = f"{int(time.time())}_{safe_name}"
            uploaded = bucket.upload_bytes(data, file_name=unique_name)
            b2_file_name = unique_name
            b2_file_id = getattr(uploaded, "id_", None) or getattr(uploaded, "file_id", None)
        except Exception:
            try:
                filename = safe_name
                file.stream.seek(0)
                file.save(os.path.join(app.config["UPLOAD_FOLDER"], filename))
            except Exception:
                filename = None

    try:
        doc = BankDocument(
            bank_id=bank_id_val,
            title=title,
            message=message,
            doc_type=doc_type or None,
            file=filename,
            b2_file_name=b2_file_name,
            b2_file_id=b2_file_id,
            created_by=user.id,
            branch_id=user.branch_id,
        )
        db.session.add(doc)
        db.session.commit()
        flash("âœ… ØªÙ… Ø¥Ø±Ø³Ø§Ù„ Ø§Ù„Ù…Ø³ØªÙ†Ø¯/Ø§Ù„Ø±Ø³Ø§Ù„Ø© Ø¥Ù„Ù‰ ØµÙØ­Ø© Ø§Ù„Ø¨Ù†Ùƒ", "success")
    except Exception as e:
        db.session.rollback()
        flash("âŒ ÙØ´Ù„ Ø§Ù„Ø­ÙØ¸", "danger")

    return redirect(url_for("employee_dashboard"))

# ---------------- ØµÙØ­Ø© Ø§Ù„Ø¹Ù…Ù„Ø§Ø¡ (Ø¥Ø¶Ø§ÙØ©/Ù‚Ø§Ø¦Ù…Ø© ÙˆØªØµØ¯ÙŠØ± CSV) ----------------
@app.route("/customers", methods=["GET", "POST"])
def customers_page():
    if session.get("role") not in ["manager", "employee", "finance"]:
        return redirect(url_for("login"))

    if request.method == "POST":
        name = (request.form.get("name") or "").strip()
        phone = (request.form.get("phone") or "").strip()
        if not name or not phone:
            flash("âš ï¸ ÙŠØ±Ø¬Ù‰ Ø¥Ø¯Ø®Ø§Ù„ Ø§Ù„Ø§Ø³Ù… ÙˆØ§Ù„Ø±Ù‚Ù…", "warning")
            return redirect(url_for("customers_page"))
        c = Customer(name=name, phone=phone)
        db.session.add(c)
        db.session.commit()
        flash("âœ… ØªÙ… Ø¥Ø¶Ø§ÙØ© Ø§Ù„Ø¹Ù…ÙŠÙ„", "success")
        return redirect(url_for("customers_page"))

    q = (request.args.get("q") or "").strip()
    query = Customer.query
    if q:
        query = query.filter(or_(Customer.name.ilike(f"%{q}%"), Customer.phone.ilike(f"%{q}%")))
    customers = query.order_by(Customer.id.desc()).all()
    return render_template("customers.html", customers=customers, q=q)


@app.route("/customers/export.csv")
def customers_export_csv():
    if session.get("role") not in ["manager", "employee", "finance"]:
        return redirect(url_for("login"))
    import csv
    from io import StringIO

    q = (request.args.get("q") or "").strip()
    query = Customer.query
    if q:
        query = query.filter(or_(Customer.name.ilike(f"%{q}%"), Customer.phone.ilike(f"%{q}%")))
    customers = query.order_by(Customer.id.desc()).all()

    si = StringIO()
    writer = csv.writer(si)
    writer.writerow(["id", "name", "phone"])  # header
    for c in customers:
        writer.writerow([c.id, c.name, c.phone])

    output = si.getvalue().encode("utf-8-sig")  # with BOM for Excel
    from flask import Response
    return Response(
        output,
        mimetype="text/csv; charset=utf-8",
        headers={
            "Content-Disposition": f"attachment; filename=customers.csv"
        },
    )

# âœ… Ø±ÙØ¹ Ù…Ø³ØªÙ†Ø¯Ø§Øª Ø§Ù„Ø´Ø±ÙƒØ© Ø¨ÙˆØ§Ø³Ø·Ø© Ø§Ù„Ù…ÙˆØ¸Ù Ù„ÙØ±Ø¹Ù‡
@app.route("/employee/branch_documents", methods=["POST"])
def employee_add_branch_document():
    if session.get("role") != "employee":
        return redirect(url_for("login"))

    user = User.query.get(session.get("user_id"))
    if not user or not getattr(user, "branch_id", None):
        flash("â›” Ù„Ø§ ÙŠÙˆØ¬Ø¯ ÙØ±Ø¹ Ù…Ø±ØªØ¨Ø· Ø¨Ø§Ù„Ù…Ø³ØªØ®Ø¯Ù…", "danger")
        return redirect(url_for("employee_dashboard"))

    title = (request.form.get("title") or "").strip()
    doc_type = (request.form.get("doc_type") or "").strip()
    issued_at = request.form.get("issued_at")
    expires_at = request.form.get("expires_at")
    file = request.files.get("file")

    if not title:
        flash("âš ï¸ ÙŠØ¬Ø¨ Ø¥Ø¯Ø®Ø§Ù„ Ø¹Ù†ÙˆØ§Ù† Ø§Ù„Ù…Ø³ØªÙ†Ø¯", "warning")
        return redirect(url_for("employee_dashboard"))

    filename = None
    b2_file_name = None
    b2_file_id = None
    if file and file.filename:
        safe_name = secure_filename(file.filename)
        # Try uploading to B2 first; fallback to local disk
        try:
            bucket = get_b2_bucket()
            data = file.read()
            unique_name = f"{int(time.time())}_{safe_name}"
            uploaded = bucket.upload_bytes(data, file_name=unique_name)
            b2_file_name = unique_name
            b2_file_id = getattr(uploaded, "id_", None) or getattr(uploaded, "file_id", None)
        except Exception:
            try:
                filename = safe_name
                file.stream.seek(0)
                file.save(os.path.join(app.config["UPLOAD_FOLDER"], filename))
            except Exception:
                filename = None

    doc = BranchDocument(
        branch_id=user.branch_id,
        title=title,
        doc_type=doc_type,
        file=filename,
        b2_file_name=b2_file_name,
        b2_file_id=b2_file_id,
        issued_at=datetime.fromisoformat(issued_at) if issued_at else None,
        expires_at=datetime.fromisoformat(expires_at) if expires_at else None,
    )
    db.session.add(doc)
    db.session.commit()
    flash("âœ… ØªÙ… Ø±ÙØ¹ Ù…Ø³ØªÙ†Ø¯ Ø§Ù„ÙØ±Ø¹", "success")
    return redirect(url_for("employee_dashboard"))

# âœ… ØªØ¹Ø¯ÙŠÙ„ Ù…Ø³ØªÙ†Ø¯ ÙØ±Ø¹ (Ù„Ù„Ù…ÙˆØ¸Ù Ø¶Ù…Ù† ÙØ±Ø¹Ù‡)
@app.route("/employee/branch_documents/<int:doc_id>/edit", methods=["POST"])
def employee_edit_branch_document(doc_id):
    if session.get("role") != "employee":
        return redirect(url_for("login"))

    user = User.query.get(session.get("user_id"))
    doc = BranchDocument.query.get_or_404(doc_id)
    if not user or doc.branch_id != user.branch_id:
        flash("â›” ØºÙŠØ± Ù…Ø³Ù…ÙˆØ­ ØªØ¹Ø¯ÙŠÙ„ Ù…Ø³ØªÙ†Ø¯Ø§Øª ÙØ±Ø¹ Ø¢Ø®Ø±", "danger")
        return redirect(url_for("employee_dashboard"))

    title = (request.form.get("title") or "").strip()
    doc_type = (request.form.get("doc_type") or "").strip()
    issued_at = request.form.get("issued_at")
    expires_at = request.form.get("expires_at")
    file = request.files.get("file")

    if title:
        doc.title = title
    doc.doc_type = doc_type
    doc.issued_at = datetime.fromisoformat(issued_at) if issued_at else None
    doc.expires_at = datetime.fromisoformat(expires_at) if expires_at else None

    if file and file.filename:
        safe_name = secure_filename(file.filename)
        # Try to upload replacement to B2; fallback to local save
        new_local = None
        new_b2_name = None
        new_b2_id = None
        try:
            bucket = get_b2_bucket()
            data = file.read()
            unique_name = f"{int(time.time())}_{safe_name}"
            uploaded = bucket.upload_bytes(data, file_name=unique_name)
            new_b2_name = unique_name
            new_b2_id = getattr(uploaded, "id_", None) or getattr(uploaded, "file_id", None)
        except Exception:
            try:
                new_local = safe_name
                file.stream.seek(0)
                file.save(os.path.join(app.config["UPLOAD_FOLDER"], new_local))
            except Exception:
                new_local = None

        if new_b2_name:
            doc.b2_file_name = new_b2_name
            doc.b2_file_id = new_b2_id
            doc.file = None
        elif new_local:
            doc.file = new_local

    db.session.commit()
    flash("âœ… ØªÙ… ØªØ­Ø¯ÙŠØ« Ø§Ù„Ù…Ø³ØªÙ†Ø¯", "success")
    return redirect(url_for("employee_dashboard"))

# âœ… Ø­Ø°Ù Ù…Ø³ØªÙ†Ø¯ ÙØ±Ø¹ (Ù„Ù„Ù…ÙˆØ¸Ù Ø¶Ù…Ù† ÙØ±Ø¹Ù‡)
@app.route("/employee/branch_documents/<int:doc_id>/delete", methods=["POST"])
def employee_delete_branch_document(doc_id):
    if session.get("role") != "employee":
        return redirect(url_for("login"))

    user = User.query.get(session.get("user_id"))
    doc = BranchDocument.query.get_or_404(doc_id)
    if not user or doc.branch_id != user.branch_id:
        flash("â›” ØºÙŠØ± Ù…Ø³Ù…ÙˆØ­ Ø­Ø°Ù Ù…Ø³ØªÙ†Ø¯Ø§Øª ÙØ±Ø¹ Ø¢Ø®Ø±", "danger")
        return redirect(url_for("employee_dashboard"))

    try:
        if doc.file:
            fpath = os.path.join(app.config["UPLOAD_FOLDER"], doc.file)
            if os.path.exists(fpath):
                os.remove(fpath)
    except Exception:
        pass

    db.session.delete(doc)
    db.session.commit()
    flash("âœ… ØªÙ… Ø­Ø°Ù Ø§Ù„Ù…Ø³ØªÙ†Ø¯", "success")
    return redirect(url_for("employee_dashboard"))

@app.route("/reports")
def reports():
    if not session.get("role") in ["employee", "manager", "engineer"]:
        return redirect(url_for("login"))

    reports = Transaction.query.filter_by(status="Ù…Ù†Ø¬Ø²Ø©").order_by(Transaction.id.desc()).all()
    return render_template("reports.html", reports=reports)


# ---------------- Ø§Ù„Ø¨Ø­Ø« Ø¨Ø±Ù‚Ù… Ø§Ù„ØªÙ‚Ø±ÙŠØ± ----------------
@app.route("/reports/search", methods=["GET", "POST"])
def search_report():
    if session.get("role") not in ["manager", "employee", "engineer"]:
        return redirect(url_for("login"))
    results = []
    search_number = None
    if request.method == "POST":
        search_number = (request.form.get("report_number") or "").strip()
        if search_number:
            results = Transaction.query.filter_by(report_number=search_number).all()
    return render_template("reports.html", reports=results, search_number=search_number)

# (ØªÙ…Øª Ø¥Ø²Ø§Ù„Ø© Ù…Ø³Ø§Ø±Ø§Øª Ø§Ù„ØªØ­Ù‚Ù‚ Ø§Ù„Ù…Ø±ØªØ¨Ø·Ø© Ø¨Ø±Ù…ÙˆØ² QR)

# --------- Ø¥Ù†Ø´Ø§Ø¡ Ø§Ù„Ø¬Ø¯Ø§ÙˆÙ„ + ØªØ±Ù‚ÙŠØ¹Ø§Øª Ù…ØªÙˆØ§ÙÙ‚Ø© Ù…Ø¹ Ù‚ÙˆØ§Ø¹Ø¯ Ù‚Ø¯ÙŠÙ…Ø© ---------
with app.app_context():
    db.create_all()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ sent_to_engineer_at Ø¥Ø°Ø§ ÙƒØ§Ù† Ø§Ù„Ø¬Ø¯ÙˆÙ„ Ù‚Ø¯ÙŠÙ…
    try:
        if not column_exists("transaction", "sent_to_engineer_at"):
            db.session.execute(text("ALTER TABLE transaction ADD COLUMN sent_to_engineer_at TIMESTAMP"))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ sent_to_engineer_at")
    except Exception:
        db.session.rollback()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ branch_id Ù„Ø¬Ø¯ÙˆÙ„ payments Ø¥Ø°Ø§ ÙƒØ§Ù† Ø§Ù„Ø¬Ø¯ÙˆÙ„ Ù‚Ø¯ÙŠÙ…
    try:
        if not column_exists("payment", "branch_id"):
            db.session.execute(text("ALTER TABLE payment ADD COLUMN branch_id INTEGER"))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ branch_id Ø¥Ù„Ù‰ payment")
    except Exception:
        db.session.rollback()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ ÙÙˆØ§ØªÙŠØ± Ø§Ù„Ø¨Ù†Ùƒ Ø¥Ø°Ø§ ØºÙŠØ± Ù…ÙˆØ¬ÙˆØ¯
    try:
        db.session.execute(text("SELECT 1 FROM bank_invoice LIMIT 1"))
    except Exception:
        try:
            db.session.execute(text(
                """
                CREATE TABLE IF NOT EXISTS bank_invoice (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    bank_id INTEGER NOT NULL,
                    transaction_id INTEGER,
                    amount FLOAT DEFAULT 0,
                    issued_at TIMESTAMP,
                    delivered_at TIMESTAMP,
                    received_at TIMESTAMP,
                    note VARCHAR(255),
                    invoice_number VARCHAR(50) UNIQUE
                )
                """
            ))
            db.session.commit()
            print("âœ… ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ bank_invoice")
        except Exception:
            db.session.rollback()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ Ø¹Ø±ÙˆØ¶ Ø§Ù„Ø£Ø³Ø¹Ø§Ø± Ø¥Ø°Ø§ ØºÙŠØ± Ù…ÙˆØ¬ÙˆØ¯
    try:
        db.session.execute(text("SELECT 1 FROM quote LIMIT 1"))
    except Exception:
        try:
            db.session.execute(text(
                """
                CREATE TABLE IF NOT EXISTS quote (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    bank_id INTEGER NOT NULL,
                    transaction_id INTEGER,
                    amount FLOAT DEFAULT 0,
                    valid_until TIMESTAMP,
                    note VARCHAR(255),
                    created_at TIMESTAMP,
                    created_by INTEGER
                )
                """
            ))
            db.session.commit()
            print("âœ… ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ quote")
        except Exception:
            db.session.rollback()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ customer_quote Ø¥Ø°Ø§ ØºÙŠØ± Ù…ÙˆØ¬ÙˆØ¯
    try:
        db.session.execute(text("SELECT 1 FROM customer_quote LIMIT 1"))
    except Exception:
        try:
            db.session.execute(text(
                """
                CREATE TABLE IF NOT EXISTS customer_quote (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    customer_name VARCHAR(150) NOT NULL,
                    amount FLOAT DEFAULT 0,
                    valid_until TIMESTAMP,
                    transaction_id INTEGER,
                    note VARCHAR(255),
                    created_at TIMESTAMP,
                    created_by INTEGER
                )
                """
            ))
            db.session.commit()
            print("âœ… ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ customer_quote")
        except Exception:
            db.session.rollback()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ customer_invoice Ø¥Ø°Ø§ ØºÙŠØ± Ù…ÙˆØ¬ÙˆØ¯
    try:
        db.session.execute(text("SELECT 1 FROM customer_invoice LIMIT 1"))
    except Exception:
        try:
            db.session.execute(text(
                """
                CREATE TABLE IF NOT EXISTS customer_invoice (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    customer_name VARCHAR(150) NOT NULL,
                    amount FLOAT DEFAULT 0,
                    issued_at TIMESTAMP,
                    transaction_id INTEGER,
                    note VARCHAR(255),
                    created_by INTEGER,
                    invoice_number VARCHAR(50) UNIQUE
                )
                """
            ))
            db.session.commit()
            print("âœ… ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ customer_invoice")
        except Exception:
            db.session.rollback()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ invoice_sequence Ø¥Ø°Ø§ ØºÙŠØ± Ù…ÙˆØ¬ÙˆØ¯
    try:
        db.session.execute(text("SELECT 1 FROM invoice_sequence LIMIT 1"))
    except Exception:
        try:
            db.session.execute(text(
                """
                CREATE TABLE IF NOT EXISTS invoice_sequence (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    year INTEGER UNIQUE NOT NULL,
                    last_number INTEGER NOT NULL DEFAULT 0
                )
                """
            ))
            db.session.commit()
            print("âœ… ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ invoice_sequence")
        except Exception:
            db.session.rollback()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ project Ø¥Ø°Ø§ ØºÙŠØ± Ù…ÙˆØ¬ÙˆØ¯
    try:
        db.session.execute(text("SELECT 1 FROM project LIMIT 1"))
    except Exception:
        try:
            db.session.execute(text(
                """
                CREATE TABLE IF NOT EXISTS project (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    name VARCHAR(200) UNIQUE NOT NULL,
                    description TEXT,
                    client_id INTEGER,
                    created_at TIMESTAMP,
                    updated_at TIMESTAMP
                )
                """
            ))
            db.session.commit()
            print("âœ… ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ project")
        except Exception:
            db.session.rollback()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ consultation Ø¥Ø°Ø§ ØºÙŠØ± Ù…ÙˆØ¬ÙˆØ¯
    try:
        db.session.execute(text("SELECT 1 FROM consultation LIMIT 1"))
    except Exception:
        try:
            db.session.execute(text(
                """
                CREATE TABLE IF NOT EXISTS consultation (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    project_id INTEGER,
                    client_id INTEGER,
                    consultant_name VARCHAR(150),
                    consultation_type VARCHAR(50) NOT NULL,
                    description TEXT,
                    status VARCHAR(20) NOT NULL,
                    start_date DATE,
                    end_date DATE,
                    cost FLOAT,
                    created_by INTEGER,
                    consultant_id INTEGER,
                    created_at TIMESTAMP,
                    updated_at TIMESTAMP
                )
                """
            ))
            db.session.commit()
            print("âœ… ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ consultation")
        except Exception:
            db.session.rollback()

    # ğŸ†• Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ section_id Ù„Ø¬Ø¯ÙˆÙ„ Ø§Ù„Ù…Ø³ØªØ®Ø¯Ù…ÙŠÙ† Ø¥Ø°Ø§ ÙƒØ§Ù† Ø§Ù„Ø¬Ø¯ÙˆÙ„ Ù‚Ø¯ÙŠÙ…
    try:
        if not column_exists("user", "section_id"):
            db.session.execute(text("ALTER TABLE user ADD COLUMN section_id INTEGER"))
            db.session.commit()
            try:
                db.session.execute(text("CREATE INDEX IF NOT EXISTS ix_user_section_id ON user(section_id)"))
                db.session.commit()
            except Exception:
                db.session.rollback()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ section_id Ø¥Ù„Ù‰ user")
    except Exception:
        db.session.rollback()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ bank_branch Ù„Ù„Ù…Ø¹Ø§Ù…Ù„Ø§Øª Ø¥Ø°Ø§ ÙƒØ§Ù† Ø§Ù„Ø¬Ø¯ÙˆÙ„ Ù‚Ø¯ÙŠÙ…
    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ bank_sent_files Ø¥Ø°Ø§ ÙƒØ§Ù† Ø§Ù„Ø¬Ø¯ÙˆÙ„ Ù‚Ø¯ÙŠÙ…
    try:
        if not column_exists("transaction", "bank_sent_files"):
            db.session.execute(text("ALTER TABLE transaction ADD COLUMN bank_sent_files TEXT"))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ bank_sent_files")
    except Exception:
        db.session.rollback()

    try:
        if not column_exists("bank_invoice", "invoice_number"):
            db.session.execute(text("ALTER TABLE bank_invoice ADD COLUMN invoice_number VARCHAR(50)"))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ invoice_number Ø¥Ù„Ù‰ bank_invoice")
    except Exception:
        db.session.rollback()

    try:
        if not column_exists("customer_invoice", "invoice_number"):
            db.session.execute(text("ALTER TABLE customer_invoice ADD COLUMN invoice_number VARCHAR(50)"))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ invoice_number Ø¥Ù„Ù‰ customer_invoice")
    except Exception:
        db.session.rollback()

    # Ø¥Ø¶Ø§ÙØ© Ù‚ÙŠØ¯ ÙØ±ÙŠØ¯ Ø¥Ø°Ø§ ÙƒØ§Ù† Ù…Ø¯Ø¹ÙˆÙ…Ù‹Ø§ (SQLite Ù„Ø§ ÙŠØ¯Ø¹Ù… Ø¨Ø³Ù‡ÙˆÙ„Ø© ALTER ADD CONSTRAINT)
    # Ù„Ø°Ù„Ùƒ Ù†ÙƒØªÙÙŠ Ø¨ÙÙ‡Ø±Ø³ ÙØ±ÙŠØ¯ Ø¹Ø¨Ø± CREATE UNIQUE INDEX Ø¥Ø°Ø§ Ù„Ù… ÙŠÙˆØ¬Ø¯
    try:
        db.session.execute(text("CREATE UNIQUE INDEX IF NOT EXISTS uq_bank_invoice_number ON bank_invoice(invoice_number)"))
        db.session.execute(text("CREATE UNIQUE INDEX IF NOT EXISTS uq_customer_invoice_number ON customer_invoice(invoice_number)"))
        db.session.commit()
    except Exception:
        db.session.rollback()

    try:
        if not column_exists("transaction", "bank_branch"):
            db.session.execute(text("ALTER TABLE transaction ADD COLUMN bank_branch VARCHAR(120)"))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ bank_branch")
    except Exception:
        db.session.rollback()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ bank_employee_name Ù„Ù„Ù…Ø¹Ø§Ù…Ù„Ø§Øª Ø¥Ø°Ø§ ÙƒØ§Ù† Ø§Ù„Ø¬Ø¯ÙˆÙ„ Ù‚Ø¯ÙŠÙ…
    try:
        if not column_exists("transaction", "bank_employee_name"):
            db.session.execute(text("ALTER TABLE transaction ADD COLUMN bank_employee_name VARCHAR(120)"))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ bank_employee_name")
    except Exception:
        db.session.rollback()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ brought_by Ù„Ù„Ù…Ø¹Ø§Ù…Ù„Ø§Øª Ø¥Ø°Ø§ ÙƒØ§Ù† Ø§Ù„Ø¬Ø¯ÙˆÙ„ Ù‚Ø¯ÙŠÙ…
    try:
        if not column_exists("transaction", "brought_by"):
            db.session.execute(text("ALTER TABLE transaction ADD COLUMN brought_by VARCHAR(120)"))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ brought_by")
    except Exception:
        db.session.rollback()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ visited_by Ù„Ù„Ù…Ø¹Ø§Ù…Ù„Ø§Øª Ø¥Ø°Ø§ ÙƒØ§Ù† Ø§Ù„Ø¬Ø¯ÙˆÙ„ Ù‚Ø¯ÙŠÙ…
    try:
        if not column_exists("transaction", "visited_by"):
            db.session.execute(text("ALTER TABLE transaction ADD COLUMN visited_by VARCHAR(120)"))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ visited_by")
    except Exception:
        db.session.rollback()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ report_template Ø¥Ù† Ù„Ù… ÙŠÙƒÙ† Ù…ÙˆØ¬ÙˆØ¯Ù‹Ø§
    try:
        db.session.execute(text("SELECT 1 FROM report_template LIMIT 1"))
    except Exception:
        try:
            db.session.execute(text(
                """
                CREATE TABLE IF NOT EXISTS report_template (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    template_type VARCHAR(50) NOT NULL,
                    content TEXT,
                    title VARCHAR(150),
                    file VARCHAR(255)
                )
                """
            ))
            db.session.commit()
            print("âœ… ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ report_template")
        except Exception:
            db.session.rollback()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ bank_document Ø¥Ø°Ø§ ØºÙŠØ± Ù…ÙˆØ¬ÙˆØ¯
    try:
        db.session.execute(text("SELECT 1 FROM bank_document LIMIT 1"))
    except Exception:
        try:
            db.session.execute(text(
                """
                CREATE TABLE IF NOT EXISTS bank_document (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    bank_id INTEGER NOT NULL,
                    title VARCHAR(200) NOT NULL,
                    message TEXT,
                    doc_type VARCHAR(100),
                    file VARCHAR(255),
                    b2_file_name VARCHAR(255),
                    b2_file_id VARCHAR(255),
                    created_at TIMESTAMP,
                    created_by INTEGER,
                    branch_id INTEGER
                )
                """
            ))
            db.session.commit()
            print("âœ… ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ bank_document")
        except Exception:
            db.session.rollback()

    # Ø¥Ù†Ø´Ø§Ø¡ Ù…Ø¯ÙŠØ± Ø§ÙØªØ±Ø§Ø¶ÙŠ Ø¥Ù† Ø£Ù…ÙƒÙ† (ØªØ¬Ù†Ø¨ Ø§Ù„Ø£Ø¹Ù…Ø¯Ø© Ø§Ù„Ù†Ø§Ù‚ØµØ©)
    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ branch_document Ø¥Ø°Ø§ ØºÙŠØ± Ù…ÙˆØ¬ÙˆØ¯
    try:
        db.session.execute(text("SELECT 1 FROM branch_document LIMIT 1"))
    except Exception:
        try:
            db.session.execute(text(
                """
                CREATE TABLE IF NOT EXISTS branch_document (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    branch_id INTEGER NOT NULL,
                    title VARCHAR(200) NOT NULL,
                    doc_type VARCHAR(100),
                    file VARCHAR(255),
                    b2_file_name VARCHAR(255),
                    b2_file_id VARCHAR(255),
                    issued_at TIMESTAMP,
                    expires_at TIMESTAMP,
                    created_at TIMESTAMP
                )
                """
            ))
            db.session.commit()
            print("âœ… ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ Ø¬Ø¯ÙˆÙ„ branch_document")
        except Exception:
            db.session.rollback()

    try:
        mgr = User.query.filter_by(role="manager").first()
    except OperationalError:
        mgr = None
    if not mgr:
        admin = User(username="admin", password=generate_password_hash("1234"), role="manager")
        db.session.add(admin)
        db.session.commit()
        print("âœ… ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ Ø­Ø³Ø§Ø¨ Ø§Ù„Ù…Ø¯ÙŠØ± Ø§Ù„Ø§ÙØªØ±Ø§Ø¶ÙŠ (username=admin, password=1234)")

    # âœ… Ø¥Ù†Ø´Ø§Ø¡ Ø­Ø³Ø§Ø¨ Ø§ÙØªØ±Ø§Ø¶ÙŠ Ù„Ù‚Ø³Ù… Ø§Ù„Ù…Ø§Ù„ÙŠØ© Ø¥Ù† Ù„Ù… ÙŠÙƒÙ† Ù…ÙˆØ¬ÙˆØ¯Ù‹Ø§
    try:
        fin = User.query.filter_by(role="finance").first()
    except OperationalError:
        fin = None
    if not fin:
        # Ø±Ø¨Ø·Ù‡ Ø¨Ø£ÙˆÙ„ ÙØ±Ø¹ Ø¥Ù† ÙˆØ¬Ø¯
        first_branch = Branch.query.first()
        finance_user = User(
            username="finance",
            password=generate_password_hash("1234"),
            role="finance",
            branch_id=first_branch.id if first_branch else None
        )
        db.session.add(finance_user)
        db.session.commit()
        print("âœ… ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ Ø­Ø³Ø§Ø¨ Ø§Ù„Ù…Ø§Ù„ÙŠØ© Ø§Ù„Ø§ÙØªØ±Ø§Ø¶ÙŠ (username=finance, password=1234)")

    # (ØªÙ…Øª Ø¥Ø²Ø§Ù„Ø© Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ verification_token)

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ report_sha256 Ø¥Ø°Ø§ ÙƒØ§Ù† Ø§Ù„Ø¬Ø¯ÙˆÙ„ Ù‚Ø¯ÙŠÙ…
    try:
        if not column_exists("transaction", "report_sha256"):
            db.session.execute(text("ALTER TABLE transaction ADD COLUMN report_sha256 VARCHAR(64)"))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ report_sha256")
    except Exception:
        db.session.rollback()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ public_share_token Ø¥Ø°Ø§ ÙƒØ§Ù† Ø§Ù„Ø¬Ø¯ÙˆÙ„ Ù‚Ø¯ÙŠÙ…
    try:
        if not column_exists("transaction", "public_share_token"):
            db.session.execute(text("ALTER TABLE transaction ADD COLUMN public_share_token VARCHAR(128)"))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ public_share_token")
    except Exception:
        db.session.rollback()

    # Ù…Ø­Ø§ÙˆÙ„Ø© Ø¥Ø¶Ø§ÙØ© Ø£Ø¹Ù…Ø¯Ø© Backblaze B2 Ù„Ù…Ù„Ù Ø§Ù„ØªÙ‚Ø±ÙŠØ± Ø¥Ù† ÙƒØ§Ù†Øª Ø§Ù„Ø¬Ø¯Ø§ÙˆÙ„ Ù‚Ø¯ÙŠÙ…Ø©
    try:
        if not column_exists("transaction", "report_b2_file_name"):
            db.session.execute(text('ALTER TABLE "transaction" ADD COLUMN report_b2_file_name VARCHAR(255)'))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ report_b2_file_name")
    except Exception:
        db.session.rollback()

    try:
        if not column_exists("transaction", "report_b2_file_id"):
            db.session.execute(text('ALTER TABLE "transaction" ADD COLUMN report_b2_file_id VARCHAR(255)'))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ report_b2_file_id")
    except Exception:
        db.session.rollback()

    # ØªØ¹Ø¨Ø¦Ø© Ø±Ù…ÙˆØ² Ø§Ù„Ù…Ø´Ø§Ø±ÙƒØ© Ø§Ù„Ø¹Ø§Ù…Ø© Ù„Ù„ØªÙ‚Ø§Ø±ÙŠØ± Ø§Ù„Ù…ÙˆØ¬ÙˆØ¯Ø© Ø¨Ø¯ÙˆÙ† Ø±Ù…Ø²
    try:
        existing_with_files = Transaction.query.filter(
            Transaction.report_file != None,
            or_(Transaction.public_share_token == None, Transaction.public_share_token == "")
        ).all()
        for tx in existing_with_files:
            tx.public_share_token = secrets.token_urlsafe(24)
        if existing_with_files:
            db.session.commit()
            print(f"âœ… ØªÙ… ØªÙˆÙ„ÙŠØ¯ Ø±ÙˆØ§Ø¨Ø· Ø¹Ø§Ù…Ø© Ù„Ù€ {len(existing_with_files)} ØªÙ‚Ø§Ø±ÙŠØ±")
    except Exception:
        db.session.rollback()

# Ø¶Ù…Ø§Ù† ÙˆØ¬ÙˆØ¯ Ø£Ø¹Ù…Ø¯Ø© Backblaze B2 Ø¹Ù†Ø¯ Ø¨Ø¯Ø¡ Ø§Ù„ØªØ´ØºÙŠÙ„ (Ù…ØªÙˆØ§ÙÙ‚ Ù…Ø¹ Flask 3.x)
def ensure_b2_columns_exist():
    try:
        if not column_exists("transaction", "report_b2_file_name"):
            db.session.execute(text('ALTER TABLE "transaction" ADD COLUMN report_b2_file_name VARCHAR(255)'))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ report_b2_file_name")
    except Exception:
        db.session.rollback()
    try:
        if not column_exists("transaction", "report_b2_file_id"):
            db.session.execute(text('ALTER TABLE "transaction" ADD COLUMN report_b2_file_id VARCHAR(255)'))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ report_b2_file_id")
    except Exception:
        db.session.rollback()

    # Ensure Backblaze columns exist for branch_document
    try:
        if not column_exists("branch_document", "b2_file_name"):
            db.session.execute(text('ALTER TABLE "branch_document" ADD COLUMN b2_file_name VARCHAR(255)'))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ b2_file_name Ø¥Ù„Ù‰ branch_document")
    except Exception:
        db.session.rollback()
    try:
        if not column_exists("branch_document", "b2_file_id"):
            db.session.execute(text('ALTER TABLE "branch_document" ADD COLUMN b2_file_id VARCHAR(255)'))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ b2_file_id Ø¥Ù„Ù‰ branch_document")
    except Exception:
        db.session.rollback()

    # Ensure Backblaze columns exist for bank_document
    try:
        if not column_exists("bank_document", "b2_file_name"):
            db.session.execute(text('ALTER TABLE "bank_document" ADD COLUMN b2_file_name VARCHAR(255)'))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ b2_file_name Ø¥Ù„Ù‰ bank_document")
    except Exception:
        db.session.rollback()
    try:
        if not column_exists("bank_document", "b2_file_id"):
            db.session.execute(text('ALTER TABLE "bank_document" ADD COLUMN b2_file_id VARCHAR(255)'))
            db.session.commit()
            print("âœ… ØªÙ…Øª Ø¥Ø¶Ø§ÙØ© Ø¹Ù…ÙˆØ¯ b2_file_id Ø¥Ù„Ù‰ bank_document")
    except Exception:
        db.session.rollback()

# ØªØ´ØºÙŠÙ„ Ù…Ù‡Ù…Ø© Ø§Ù„ØªÙ‡ÙŠØ¦Ø© Ø¹Ù†Ø¯ Ø¨Ø¯Ø¡ Ø§Ù„ØªØ´ØºÙŠÙ„ Ù„Ø¶Ù…Ø§Ù† Ø§Ù„Ø£Ø¹Ù…Ø¯Ø© Ø§Ù„Ù…Ø·Ù„ÙˆØ¨Ø©
try:
    with app.app_context():
        ensure_b2_columns_exist()
except Exception as e:
    print(f"âš ï¸ ÙØ´Ù„ Ø¶Ù…Ø§Ù† Ø£Ø¹Ù…Ø¯Ø© B2 Ø¹Ù†Ø¯ Ø¨Ø¯Ø¡ Ø§Ù„ØªØ´ØºÙŠÙ„: {e}")

# ---------------- ØªÙ‚Ø±ÙŠØ± Ø¯Ø®Ù„ Ù…ÙˆØ¸Ù ----------------
@app.route("/employee_income", methods=["GET", "POST"])
def employee_income():
    if session.get("role") != "manager":
        return redirect(url_for("login"))

    employees = User.query.all()
    income = None
    report_data = []
    selected_emp = None
    start_date = None
    end_date = None

    if request.method == "POST":
        emp_id = request.form.get("employee_id")
        start_date = request.form.get("start_date")
        end_date = request.form.get("end_date")

        employee = User.query.get(emp_id)
        selected_emp = employee.username if employee else None

        query = Transaction.query.filter(Transaction.brought_by == selected_emp)
        if start_date:
            query = query.filter(Transaction.date >= start_date)
        if end_date:
            query = query.filter(Transaction.date <= end_date)

        transactions = query.all()
        income = 0.0

        for t in transactions:
            payments = Payment.query.filter_by(transaction_id=t.id).all()
            paid_amount = sum(p.amount for p in payments)
            status_payment = "Ù…Ø¯ÙÙˆØ¹Ø©" if paid_amount > 0 else "ØºÙŠØ± Ù…Ø¯ÙÙˆØ¹Ø©"
            report_data.append({
                "id": t.id,
                "client": t.client,
                "date": t.date,
                "status": t.status,
                "fee": paid_amount,
                "payment_status": status_payment
            })
            income += paid_amount

    return render_template(
        "employee_income.html",
        employees=employees,
        selected_emp=selected_emp,
        transactions=report_data,
        income=income,
        start_date=start_date,
        end_date=end_date
    )



app.config["BFNeZpjEro8pwFxR1H20twlTd2pL5MZtWrDATu4ME2RcbzhN-PBHcpk_jYrRlDUrn4SUxHJ5TOEF796OXs-NN"] = "ğŸ”‘_Ø¶Ø¹_Ø§Ù„Ù…ÙØªØ§Ø­_Ø§Ù„Ø¹Ø§Ù…"
app.config["Gv_NJwUe_M5R6seQItCoivxv3mTp6JiJQmkcrQmICuk="] = "ğŸ”_Ø¶Ø¹_Ø§Ù„Ù…ÙØªØ§Ø­_Ø§Ù„Ø®Ø§Øµ"
app.config["VAPID_CLAIMS"] = {
    "sub": "mailto:your-email@example.com"
}




if __name__ == "__main__":
    app.run(debug=True)
